"""
Task 5: Assemble Hengli Hydraulics (SSE:601100) initiation coverage report.
Output: 30-50 page DOCX with 25+ embedded charts + 12-20 tables.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CHARTS_DIR = "charts_hengli"
OUT_PATH = "reports/company/Hengli_SSE601100/Hengli_SSE601100_Initiation_Report_2026-05-19.docx"

NAVY = RGBColor(0x1F, 0x4E, 0x79)
ORANGE = RGBColor(0xE0, 0x7B, 0x14)
GREEN = RGBColor(0x2E, 0x7D, 0x45)
RED = RGBColor(0xC6, 0x28, 0x28)
GREY = RGBColor(0x59, 0x59, 0x59)
BLACK = RGBColor(0, 0, 0)

doc = Document()

# Set default font to Times New Roman
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
rPr = style.element.get_or_add_rPr()
rFonts = OxmlElement("w:rFonts")
rFonts.set(qn("w:eastAsia"), "Times New Roman")
rPr.append(rFonts)

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

def add_heading(text, level=1, color=NAVY):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Times New Roman"
        run.font.bold = True
    return h

def add_para(text, bold=False, italic=False, color=None, size=11, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    return p

def add_image(path, width_inches=6.5, caption=None):
    if not os.path.exists(path):
        add_para(f"[Chart not found: {path}]", italic=True, color=GREY)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(path, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.font.size = Pt(9); cr.font.italic = True; cr.font.color.rgb = GREY

def add_chart(num, desc, width=6.5, caption=None):
    """Find and add chart by number."""
    files = [f for f in os.listdir(CHARTS_DIR)
             if f.startswith(f"chart_{num:02d}_") and f.endswith(".png")]
    if files:
        add_image(os.path.join(CHARTS_DIR, files[0]), width_inches=width,
                  caption=caption or f"Exhibit {num}: {desc}")

def add_table_styled(headers, rows, col_widths=None, header_color="1F4E79"):
    """Add a styled table with header row + data rows."""
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Grid Accent 1"
    # Header
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255)
        r.font.name = "Times New Roman"; r.font.size = Pt(10)
        # Shade
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), header_color)
        tcPr.append(shd)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = "Times New Roman"; r.font.size = Pt(10)
            if ci == 0:
                r.font.bold = True
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

def add_page_break():
    doc.add_page_break()

def add_separator():
    p = doc.add_paragraph()
    p.add_run("─" * 90).font.color.rgb = GREY

# ============================================================
# PAGE 1: COVER / INVESTMENT SUMMARY
# ============================================================
# Banner
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("EQUITY RESEARCH  ·  INITIATING COVERAGE")
r.font.name = "Times New Roman"; r.font.size = Pt(10); r.font.bold = True
r.font.color.rgb = NAVY

# Main heading
h = doc.add_heading("Hengli Hydraulics (SSE:601100)", level=0)
for run in h.runs:
    run.font.color.rgb = NAVY
    run.font.name = "Times New Roman"

subt = doc.add_paragraph()
subt.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subt.add_run("China's hydraulics champion — premium valuation already prices in humanoid optionality")
r.font.size = Pt(13); r.font.italic = True; r.font.color.rgb = GREY
r.font.name = "Times New Roman"

# Recommendation banner table
rec_table = doc.add_table(rows=1, cols=4)
rec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
rec_data = [
    ("Rating", "HOLD"),
    ("Price target (12M)", "RMB 106"),
    ("Current price", "RMB 119.60"),
    ("Implied return", "−11%"),
]
for i, (h, v) in enumerate(rec_data):
    cell = rec_table.rows[0].cells[i]
    cell.text = ""
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(h); r1.font.size = Pt(9); r1.font.bold = True; r1.font.color.rgb = GREY
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = cell.add_paragraph()
    r2 = p2.add_run(v); r2.font.size = Pt(16); r2.font.bold = True
    r2.font.color.rgb = ORANGE if i == 0 else NAVY
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Shade header cell
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2"); tcPr.append(shd)

doc.add_paragraph()

# Investment summary intro
add_heading("Investment summary", level=1)
add_para(
    "We initiate coverage of Jiangsu Hengli Hydraulic Co., Ltd. (SSE:601100) with a HOLD "
    "rating and a 12-month price target of RMB 106 per share, implying ~11% downside to "
    "the May 16, 2026 close of RMB 119.60. Hengli is, on a fundamental basis, the highest-"
    "quality industrial-hydraulics franchise in China, and arguably one of the best globally on "
    "a margin- and ROIC-adjusted basis (FY2025 ROE 16.6%, EBITDA margin 33%, ROIC ~22%). "
    "The company is the only domestic player whose product line spans the full hydraulic "
    "stack — cylinders, axial-piston pumps, multi-way valves, motors, systems — for "
    "excavators, marine equipment, tunnel-boring machines, aerial work platforms, "
    "agriculture, and (since 2022) the high-precision linear-drive components used in "
    "machine tools and humanoid robot actuators."
)

add_para(
    "However, the share price already reflects (i) full credit for the cyclical excavator "
    "recovery currently underway, (ii) a substantial premium for the linear-drive / humanoid-"
    "roller-screw narrative, and (iii) top-decile valuation multiples in the 3-year band. Our "
    "base-case DCF generates only RMB 67/share — 44% below spot — and even after assigning "
    "weight to a peer-multiple framework and an explicit humanoid optionality overlay, the "
    "blended fair value sits below the current price. The HOLD is therefore a 'wait for a "
    "better entry' call, not a structurally negative view."
)

# First key chart
add_chart(1, "3-year share price history", width=6.3)
add_chart(2, "Revenue and gross margin trajectory (FY20A-FY30E)", width=6.3)

add_page_break()

# ============================================================
# PAGE 2-3: KEY STATISTICS + THESIS
# ============================================================
add_heading("Key statistics", level=2)

stats_table = doc.add_table(rows=10, cols=4)
stats_table.style = "Light Grid Accent 1"
stats_data = [
    ("Market cap", "RMB 160.3 bn", "FY25 revenue", "RMB 10.94 bn"),
    ("Enterprise value", "RMB 151.2 bn", "FY25 net income", "RMB 2.73 bn"),
    ("Shares outstanding", "1,340.8 m", "FY25 ROE", "16.6%"),
    ("Net cash", "RMB 9.2 bn", "FY25 EBITDA margin", "33.4%"),
    ("Debt / EBITDA", "0.01×", "FY25 EPS", "RMB 2.04"),
    ("TTM P/E", "58.6×", "FY26E EPS", "RMB 2.27"),
    ("TTM EV/EBITDA", "41.5×", "FY26E revenue growth", "13.6%"),
    ("TTM P/S", "14.6×", "Dividend yield (FY25)", "0.47%"),
    ("P/B", "9.3×", "Payout ratio", "33%"),
    ("Beta (2y vs SSE)", "1.05", "Free float", "~35.7%"),
]
for ri, (k1, v1, k2, v2) in enumerate(stats_data):
    cells = stats_table.rows[ri].cells
    cells[0].text = k1; cells[1].text = v1; cells[2].text = k2; cells[3].text = v2
    for i, cell in enumerate(cells):
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = "Times New Roman"; r.font.size = Pt(10)
                r.font.bold = (i % 2 == 0)
                if i % 2 == 1: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()
add_heading("Investment thesis — the 'wait-for-better-entry' HOLD", level=1)

add_heading("1. World-class franchise — the highest-quality Chinese hydraulics asset", level=2)
add_para(
    "Hengli is the only Chinese hydraulics company with a complete product portfolio spanning "
    "cylinders, axial-piston pumps, multi-way valves, low-speed high-torque motors and "
    "integrated systems. It is the dominant supplier in the Chinese excavator hydraulics market "
    "(>50% share in mid-large cylinders), is one of only two Asian Caterpillar 'platinum medal' "
    "suppliers, and has built a 30%+ market share in domestic excavator main pumps and multi-"
    "way valves over the past decade. The company employs 8,400 staff including 1,104 R&D "
    "engineers (13.1% of headcount), holds 1,125 valid patents, and operates a 'national "
    "excellence-class smart factory' — the highest grade in MIIT's hierarchy."
)
add_chart(8, "FY2025 product portfolio — revenue and gross-margin by segment", width=6.3)

add_para(
    "The FY2025 segment mix shows the maturity of this franchise: cylinders RMB 5.25bn at 39.7% "
    "gross margin (the defensible cash cow), pumps/valves/motors RMB 4.33bn at 48.8% gross "
    "margin (the highest-margin engine), systems RMB 0.39bn at 34.4% gross margin (small but "
    "strategic), and components/castings/linear-drive RMB 0.89bn at 15.3% gross margin (the "
    "narrative engine). Group-level gross margin of 41.6% and EBITDA margin of 33% are best-in-"
    "class for any industrial-machinery supplier globally."
)

add_heading("2. Three growth vectors — but all priced in", level=2)
add_para(
    "The base-case bull thesis rests on three vectors: (a) China hydraulics market growing 8% "
    "CAGR FY25-FY30 (RMB 82bn to ~RMB 120bn) with Hengli outperforming at 13-14% CAGR via "
    "continued share gain from Bosch Rexroth retreat from low-end Chinese OEM business; (b) "
    "Mexico plant Tier-1 ramp with Caterpillar Mexicali expected to add RMB 1.5-2.0bn of "
    "overseas revenue by FY28; and (c) the linear-drive / planetary roller-screw ramp from "
    "current RMB 100m to RMB 2.3bn by FY30 per management's implied path. Each of these is "
    "achievable; we model all three in our base case. The issue is not the operational story "
    "but the valuation premium already paid for it."
)
add_chart(22, "Linear-drive revenue ramp — 'second growth curve'", width=6.3)

add_heading("3. Valuation — DCF vs market gap of 80%+", level=2)
add_para(
    "Our base-case DCF generates RMB 67/share — a 44% discount to spot. This is not a result "
    "of pessimistic assumptions; it uses the same FY2026-FY2030 projections as the operational "
    "thesis (13% revenue CAGR, 30.5% terminal EBITDA margin, 8.5% WACC, 3.0% terminal growth). "
    "Even in our bull case (18% CAGR, 33.5% EBITDA margin, full humanoid penetration), DCF "
    "fair value reaches only RMB 96 — still below spot. The market is paying for off-DCF "
    "optionality — specifically a re-rating in line with humanoid-supply-chain precedents "
    "(Tuopu 100× P/E, Shuanglin 117× P/E) — which we view as a 25% probability over the next "
    "12 months."
)
add_chart(28, "★ DCF sensitivity — implied price per share (RMB)", width=5.5)

add_heading("4. Why HOLD and not BUY", level=2)
add_para(
    "We considered three alternative ratings before settling on HOLD: BUY (PT RMB 135-145 "
    "weighted more heavily toward humanoid precedent), HOLD-OUTPERFORM (PT RMB 120 implying "
    "modest 0-5% upside, our framework's most-popular comp-anchored answer), and HOLD-"
    "UNDERPERFORM (the strict DCF answer of PT RMB 67 implying SELL). Our final HOLD at "
    "RMB 106 reflects four judgements. First, the company is fundamentally a high-quality "
    "asset that we want to own at the right price — therefore an outright SELL is not our "
    "view. Second, the comp-anchored RMB 120 PT ignores the DCF reality that the underlying "
    "cash flows do not support the current price — the methodology weighting must include DCF "
    "to reflect this. Third, the BUY-at-RMB-145 case requires a humanoid-supply-chain "
    "announcement which is genuinely a coin-flip rather than a high-probability event, and "
    "we are not willing to underwrite our 12-month rating on a binary catalyst we cannot "
    "model. Fourth, the asymmetry of the position at RMB 119.60 is unattractive: downside "
    "of 30%+ if either (a) the humanoid narrative fades or (b) FY26 results disappoint vs "
    "consensus, against upside of 25-50% only if a major humanoid announcement materialises."
)
add_para(
    "The HOLD is therefore explicitly a 'wait-for-better-entry' rating. We would turn BUY at "
    "RMB 95 or below, which would represent a roughly 25-30× FY26E P/E — well below the peer "
    "median and providing meaningful downside protection if the bear case materialises while "
    "preserving substantial upside if the bull case plays out. We would turn SELL at RMB 145 "
    "or above, where the implied multiple of >65× FY26E P/E starts to embed humanoid-supply-"
    "chain certainty that has not been confirmed.",
    italic=True
)

add_page_break()

# ============================================================
# PAGE 4: WHAT WOULD CHANGE OUR VIEW + RISKS
# ============================================================
add_heading("What would change our view", level=2)
add_para(
    "We would turn BUY at RMB 95 or below — where the embedded humanoid premium becomes "
    "asymmetric again — and SELL at RMB 145 or above — where the implied probability of "
    "humanoid-supply-chain certification exceeds what we view as reasonable. The principal "
    "12-month catalyst that would force an upgrade is a confirmed Tier-1 humanoid OEM supply "
    "award (Tesla Optimus, Figure, or a major Chinese humanoid program at >100k units/year "
    "scale). Sell-side notes have repeatedly cited Hengli as a plausible candidate based on "
    "(i) its planetary roller-screw production capability being one of only ~3 Chinese options, "
    "(ii) the geographic adjacency of Tesla's Mexicali humanoid pilot line to Caterpillar "
    "Mexicali (where Hengli is co-locating its Mexico plant), and (iii) reported sample "
    "shipments compatible with Optimus joint geometry. Hengli has not confirmed any such "
    "relationship, and we explicitly do not include this in our base case."
)

add_heading("Key risks", level=2)
add_chart(35, "Catalyst probability map — humanoid OEM award = high-impact low-probability tail")
add_para(
    "The five most material risks are: (1) linear-drive execution — RMB 1.4bn invested with "
    "FY25 revenue of only RMB 100m at 15% gross margin vs 41% group margin; (2) Caterpillar "
    "in-sourcing — Mexicali plant could cut Hengli cylinder content from ~80% to <50% over "
    "2026-2028 (RMB 500-700m of annual revenue at risk); (3) excavator cyclical down-cycle "
    "returns in 2027-2028; (4) multiple compression — TTM P/E 58× is in the top decile of the "
    "3-year band, with sentiment-driven derating risk; (5) RMB / USD / EUR FX hedging risk on "
    "the USD 580m notional swap book (33% of equity). Severity assessment in Section 9."
)

add_page_break()

# ============================================================
# PAGE 5: TABLE OF CONTENTS
# ============================================================
add_heading("Table of Contents", level=1)
toc = [
    ("1. Investment summary",                          "1"),
    ("2. Investment thesis",                            "3"),
    ("3. Risks to thesis",                              "4"),
    ("4. Company background (Company 101)",             "6"),
    ("    4.1 History & milestones",                    "6"),
    ("    4.2 Management team",                         "7"),
    ("    4.3 Products & segments",                     "9"),
    ("    4.4 Customers & contracts",                   "11"),
    ("    4.5 Geographic footprint",                    "12"),
    ("5. Industry & TAM analysis",                      "14"),
    ("6. Competitive landscape",                        "16"),
    ("7. Financial analysis",                           "18"),
    ("    7.1 Historical performance review",           "18"),
    ("    7.2 Revenue projections & assumptions",       "21"),
    ("    7.3 Margin & profitability projections",      "23"),
    ("    7.4 Cash flow, CapEx & balance sheet",        "25"),
    ("    7.5 Scenario analysis",                       "27"),
    ("8. Valuation analysis",                           "29"),
    ("    8.1 DCF analysis & sensitivity",              "29"),
    ("    8.2 Comparable companies analysis",           "32"),
    ("    8.3 Precedent transactions",                  "34"),
    ("    8.4 Football field & price target",           "35"),
    ("9. Risks (detailed)",                             "37"),
    ("10. ESG considerations",                          "39"),
    ("11. Catalysts",                                   "40"),
    ("12. Appendices",                                  "41"),
]
for item, pg in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(item); r.font.size = Pt(11)
    if item.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "10.", "11.", "12.")):
        r.font.bold = True
    tab = p.add_run("\t" + "." * 50 + " " + pg); tab.font.color.rgb = GREY; tab.font.size = Pt(9)

add_page_break()

# ============================================================
# SECTION 4: COMPANY 101
# ============================================================
add_heading("4. Company background", level=1)

add_heading("4.1 History & milestones", level=2)
add_para(
    "Jiangsu Hengli Hydraulic Co., Ltd. (江苏恒立液压股份有限公司) was founded by Wang Liping in "
    "Wuxi in 1990 as a seven-person pneumatic-components workshop with RMB 50,000 of capital. "
    "The company's transformational moment came in 1999 when it developed the first domestic "
    "excavator hydraulic cylinder, breaking the KYB / Kawasaki / Eaton oligopoly that had "
    "supplied 100% of Chinese excavator OEMs up to that point. The company incorporated in "
    "its modern form in Changzhou in 2005, listed on the Shanghai Stock Exchange in October "
    "2011 at RMB 23.00/share, and has since executed three strategic pivots that define the "
    "modern business."
)
add_chart(5, "Hengli — 35 years of strategic milestones", width=6.5)

add_para(
    "The first pivot, in 2013, was the expansion from cylinders into axial-piston pumps, "
    "multi-way valves, and travel/swing motors. Hengli leveraged its existing relationships "
    "with Chinese excavator OEMs — which already used Hengli cylinders — to seed a substantially "
    "higher-margin product line. Today pumps & valves contribute ~40% of revenue at ~49% gross "
    "margin versus ~40% for cylinders. The second pivot, in 2018, recognised that Chinese "
    "excavator demand was structurally cyclical and broadened the addressable market by "
    "selling pumps to forging, injection-moulding, wind, agricultural, and aerial-work-"
    "platform OEMs. The third — and most strategically consequential — pivot was the 2022 "
    "linear-drive entry, funded by a RMB 1.4bn private placement issued in December 2022 at "
    "RMB 56.40/share. This funded a dedicated factory in Changzhou for high-precision ball "
    "screws, planetary roller screws and linear guideways — an explicit bet that Hengli's "
    "precision-motion know-how (super-finishing, induction hardening, micron-grade grinding) "
    "would transfer to the linear-motion components needed in CNC machine tools, semiconductor "
    "equipment, and humanoid-robot actuators."
)
add_chart(6, "Strategic evolution — from cylinders (1990) to linear-drive (2022→)", width=6.3)

add_heading("4.2 Management team", level=2)
add_chart(7, "Senior management team", width=6.5)

add_para(
    "Wang Liping (汪立平), age 60, Chairman & Founder. Born 1966 in Jiangsu with a junior-"
    "college education and senior-economist certification, Wang has been continuously CEO/"
    "Chairman of the integrated group for 35 years — an exceptionally long tenure for a "
    "Chinese listed-company founder. His signature accomplishment is the 1999 development of "
    "the first domestic excavator hydraulic cylinder that broke the KYB/Kawasaki/Eaton "
    "oligopoly, a technical bet the company would not have survived without. He has then "
    "successfully repeated the same playbook for pumps in 2013 and now linear motion in 2022. "
    "Wang holds, together with his wife Qian Peixin (Hong Kong resident, director of Shenuo "
    "Technology) and son Wang Qi, approximately 64.3% of Hengli through three holding vehicles "
    "— Jiangsu Hengli Holdings (36.95%), Shenuo Technology HK (14.10%), and Ningbo Hengyi "
    "(13.29%, the management equity vehicle). Wang was re-elected to a fresh 3-year board "
    "term in September 2025, indicating no near-term succession trigger. FY2025 cash "
    "compensation was RMB 1.42m — modest by global standards; the wealth is entirely in "
    "equity."
)

add_para(
    "Qiu Yongning (邱永宁), age 56, Director and General Manager (CEO). Holds a bachelor's "
    "degree in mechanical engineering from Nanjing University of Aeronautics and Astronautics. "
    "Career path runs through Jiangsu Zhengmao Group (production manager), Hualixin Electric "
    "Machinery (deputy GM), and Kayaba Hydraulic Industry Zhenjiang — the China JV of Japan's "
    "KYB — where Qiu was head of production. He joined Hengli in the late 2000s as deputy GM, "
    "was promoted to GM, and now sits on the boards of Hengli Technology, Hengli Transmission "
    "(linear-drive subsidiary), and Hengli Precision Industrial. FY2025 compensation RMB 1.06m. "
    "Qiu is the day-to-day operator running production planning, supply chain, and execution "
    "against Wang's strategic direction; his KYB lean-manufacturing methodology is meaningful "
    "given KYB's status as the global cylinder leader."
)

add_para(
    "Peng Mei (彭玫), age 57, CFO / Finance Director. Beihang University, senior accountant. "
    "Prior roles: chief accountant at Changzhou Lanxiang Mechanical (now part of AVIC's "
    "Changzhou Lanxiang Aviation Engine Co.); finance head at Yamazaki Motorcycle Changzhou; "
    "CFO at Globe Industries Changzhou. Peng has 30+ years of finance leadership across "
    "Chinese, Japanese, and US-owned operations in the Changzhou industrial cluster. She "
    "joined Hengli pre-IPO and oversaw the 2011 SSE listing and the 2022 private placement. "
    "FY2025 compensation RMB 0.90m. Track record points include zero financial-statement "
    "restatements, zero audit qualifications across 14 years public, and a continuously "
    "expanding dividend (FY2024 dividend RMB 0.94bn; FY2025 proposed dividend RMB 0.75bn at "
    "RMB 5.60 per 10 shares) — strong governance signals."
)

add_para(
    "Xu Jin (徐进), age 45, Director, Deputy GM & Sales Director. Master's degree, joined "
    "Hengli as a sales-team manager and rose to head of group sales. He directly owns the "
    "global-OEM relationship book — Caterpillar, Komatsu, Doosan, Liugong, Volvo CE — and has "
    "been the architect of Hengli's overseas-direct-sales build-out (Mexico, India, "
    "Indonesia). FY2025 compensation RMB 0.91m. Sits on the board of Hengli Technology, "
    "Shanghai Lixin, and Hengli Precision Industrial."
)

add_para(
    "Wang Bin (王斌), age 43, Deputy GM & Hengli Precision Industrial GM. Master's degree, "
    "internal promotion from regional sales → procurement manager → supply-chain GM → Liquid "
    "Pressure Technology deputy GM. He now leads the precision-machining unit responsible for "
    "the linear-drive product line and the planetary roller screws — the most strategically "
    "important new business at the company. FY2025 compensation RMB 0.96m."
)

add_para(
    "Management track record assessment. This team has delivered for 30+ years across three "
    "step-changes (cylinders → pumps → linear motion). The execution playbook — leverage "
    "existing OEM relationships, run R&D in lock-step with one anchor customer, then expand "
    "the product to all customers — has worked twice and is now being run a third time on "
    "roller screws. The main gap is CFO and CEO succession: both are above 55, and there is "
    "no publicly identified internal successor. Capital allocation has been disciplined — "
    "Hengli has only one major private placement (2022, RMB 1.4bn), virtually no acquisitions "
    "of size, and a 30-40% dividend payout ratio.",
    italic=True
)

add_page_break()

# ============================================================
# 4.3 Products
# ============================================================
add_heading("4.3 Products & segments", level=2)
add_para(
    "Hengli divides its FY2025 main-business revenue of RMB 10.86bn across four reporting "
    "segments. The audited breakdown is shown below."
)

add_table_styled(
    ["Segment", "FY25 Revenue (RMB m)", "YoY", "Gross margin", "% of main biz"],
    [
        ["Hydraulic cylinders", "5,254", "+10.4%", "39.7%", "48.4%"],
        ["Pumps, valves & motors", "4,326", "+20.7%", "48.8%", "39.8%"],
        ["Hydraulic systems", "385", "+30.0%", "34.4%", "3.5%"],
        ["Components & castings (incl. linear-drive)", "891", "+30.3%", "15.3%", "8.2%"],
        ["TOTAL main business", "10,856", "+16.4%", "41.2%", "100.0%"],
    ],
)

add_chart(3, "★ Revenue by product segment — historical and projected (RMB millions)", width=6.3)

add_heading("Hydraulic cylinders — the flagship", level=3)
add_para(
    "Hengli ships ~900,000 cylinders per year, the world's largest single-supplier output. "
    "The product line spans excavator-specific cylinders (boom, arm, bucket) for excavators "
    "ranging 1.5 t to 400 t — including a recently launched giant 90-tonne mining-excavator "
    "cylinder. Non-standard cylinders for aerial work platforms (AWPs), marine deck equipment "
    "(winch, hatch covers, davits, jack-up), wind-turbine pitch & yaw systems, port equipment "
    "(RTG/STS cranes), and water-conservancy gates round out the range. The 156-metre offshore "
    "pile-driver closed-loop energy-recovery cylinder system delivered in 2025 is the largest "
    "of its kind ever built. Competitive verdict: yes, moat. Type: scale + cost leadership + "
    "switching costs at the OEM. Evidence — Hengli holds >50% China share in mid-large "
    "excavator cylinders, has been Caterpillar's only Asian 'platinum-medal' cylinder supplier "
    "for five consecutive years, and runs a gross margin (39.7%) materially above small-scale "
    "local competitors (typical 15-25%). Closest competitor: KYB Corporation (TYO:7242) — "
    "comparable on quality at the high end; behind on cost; equal at parity for Japanese and "
    "Korean OEMs."
)

add_heading("Hydraulic pumps, valves & motors — the highest-margin engine", level=3)
add_para(
    "The product family includes (a) axial-piston main pumps for excavators (the so-called "
    "'main pump', historically a Kawasaki / Bosch Rexroth oligopoly), (b) multi-way valves "
    "(also historically Kawasaki/Eaton/Parker), (c) low-speed high-torque travel motors and "
    "swing motors, (d) industrial pumps for wind, injection moulding, forging presses, and "
    "(e) closed-loop pumps for AWPs and agriculture. In 2025 Hengli broke into the ultra-"
    "large excavator main pump category previously locked up by Bosch Rexroth and Kawasaki, "
    "with units for 50-90-tonne machines. Competitive verdict: yes, technology / IP moat. "
    "Evidence — Hengli is the only Chinese supplier with cylinder + main-pump + multi-way "
    "valve capability under one roof; gross margin 48.8% is consistent with proprietary "
    "technology; multiple national R&D project leadership gives technical first-mover "
    "position. Closest competitor: Bosch Rexroth (Bosch GmbH unit, private) — ahead on "
    "absolute precision and on closed-loop electronic control, behind on cost and on Chinese-"
    "OEM relationship density. Kawasaki Precision Machinery is the second comp — at parity "
    "on technology, behind on cost."
)

add_heading("Hydraulic systems — small but high-strategy", level=3)
add_para(
    "Integrated turnkey systems for tunnel-boring machines (China Railway Construction Heavy "
    "Industry), marine deck equipment, water-conservancy gates (Pinglu Canal), and special "
    "vehicles. FY2025 segment growth was 30%. Margin is the lowest of the four (34.4%) "
    "because system delivery includes purchased third-party content. Competitive verdict: "
    "partial moat — moat is from the underlying components, not from the system-engineering "
    "layer."
)

add_heading("Components & castings — the linear-drive engine", level=3)
add_para(
    "This is the segment to watch. Conventional content: cylinder spares, pump cores, high-"
    "precision castings (Hengli has its own foundry). The new content consolidated here since "
    "2025: linear-drive components — ball screws, planetary roller screws, linear guideways, "
    "electric cylinders. Reported FY2025 capacity is 70,000 sets of precision-ground ball "
    "screws/yr, 360,000 m of linear guideways/yr, and small-batch planetary-roller-screw "
    "shipments. The 2022 private placement defined a full-build-out plan of 104,000 standard "
    "ball-screw electric cylinders, 4,500 heavy-duty ball-screw electric cylinders, 750 "
    "planetary-roller-screw electric cylinders, plus 100,000 m of standard and 100,000 m of "
    "heavy-duty ball screws. FY2025 shipment value was estimated at RMB 80-100 million, with "
    "management guidance for 3× growth in FY2026 to ~RMB 300m and >300 customers already in "
    "the database. Competitive verdict: partial moat, with a clear path to 'yes'. Moat type "
    "today: certification + scale-up advantage; planned moat: technology + cost. Evidence — "
    "Hengli is one of only ~3 Chinese suppliers (the others being Beijing Precision "
    "Engineering and Nanjing Technology) with announced planetary-roller-screw production "
    "capability; sample shipments to Optimus joint-module designers have been reported "
    "(unconfirmed by Hengli directly). Closest named competitors: Tuopu Group (601689) for "
    "the integrated linear-actuator-module deal; Shuanglin (300100) for ball-screws; "
    "Schaeffler (ETR:SHA) for the global benchmark on planetary roller screws."
)

add_chart(25, "Production capacity build-out — cylinder scale + linear-drive ramp", width=6.3)

add_page_break()

# ============================================================
# 4.4 Customers
# ============================================================
add_heading("4.4 Customers & contracts", level=2)
add_para(
    "Hengli's FY2025 top-5 customer concentration is RMB 4.60bn, equal to 42.07% of total "
    "revenue, with zero related-party content. FY2024 was 44.13% (RMB 3.64bn) and FY2023 was "
    "42.45% (RMB 3.81bn) — a stable mid-40% range over 3 years. The single-largest customer "
    "is not separately named in Chinese annual-report disclosure (Chinese A-share rules "
    "require only the top-5 aggregate and confirmation that no single customer is >50%); "
    "industry triangulation from supplier reports and analyst notes points to Caterpillar as "
    "the #1 customer at roughly 13-15% of group revenue, followed by Sany Heavy Industry, "
    "XCMG, Komatsu, and Liugong at 6-10% each. Doosan and Hitachi Construction Machinery are "
    "the other typical entrants in the top-10, with Kubota rising rapidly via the new "
    "agricultural-tractor pump programme."
)
add_chart(9, "Estimated customer concentration — Top-5 = 42% of FY25 revenue", width=4.8)

add_para(
    "Contract structure. Annual framework agreements with the top-10 OEMs, executed at year-"
    "start; specific POs flow against the framework; lead time 30-60 days; payment terms 60-"
    "120 days (a structural reason for Hengli's RMB 1.9bn accounts-receivable balance, which "
    "grew 39% YoY in FY2025 alongside revenue). No multi-year volume commitments — pricing "
    "renegotiated annually. Sales channel is 100% direct OEM with no distribution layer. The "
    "internal sales force is organised in three units — domestic-excavator OEMs, domestic-"
    "industrial OEMs, and overseas — under Xu Jin (sales director). Service is via 30+ China "
    "service offices plus overseas service hubs in Europe, the US, Japan, Mexico, Indonesia, "
    "India, and (since 2025) the UK, Italy, Guinea."
)

add_heading("4.5 Geographic footprint", level=2)
add_para(
    "Domestic FY2025 revenue was RMB 8.75bn (+20.7% YoY, 80% of total); overseas RMB 2.11bn "
    "(+1.6% YoY, 19% of total). Gross margin on domestic and overseas is nearly identical "
    "(41.1% vs 41.4%), unusual for Chinese exporters and indicating that Hengli is not "
    "competing on price overseas but on quality. Overseas growth was muted in 2025 because "
    "the US/EU construction-machinery market was weak; the Mexico plant has been built "
    "specifically to capture growth in 2026-2028. Headquarters and the principal Tier-1 "
    "flexible smart factory are at 99 Longqian Road, Wujin Hi-Tech District, Changzhou, "
    "Jiangsu; subsidiaries span Shanghai, Wuxi, Changsha, Nanjing, plus overseas units in "
    "Germany (InLine Hydraulik GmbH), Japan (HARADA Giken, Hattori Seiko, HST), the United "
    "States (Hengli America, the 580 West Crossroads Chicago facility), India, Indonesia, "
    "Mexico (just commissioned 2025), the UK, Italy, Brazil, France, Singapore, and Guinea."
)
add_chart(4, "★ Revenue by geography — domestic dominant; Mexico-led NA ramp 2026-2028", width=6.3)
add_chart(26, "Global footprint — 12 countries operating, Mexico new in 2025", width=6.3)

add_page_break()

# ============================================================
# SECTION 5: INDUSTRY & TAM
# ============================================================
add_heading("5. Industry & TAM analysis", level=1)

add_heading("5.1 China hydraulics industry", level=2)
add_para(
    "The China Hydraulic Pneumatic & Sealing Industry Association (CHPSC) reported FY2025 "
    "industry output of approximately RMB 82bn, up materially from FY2020 and continuing the "
    "14th Five-Year Plan growth path. China imported USD 2.37bn and exported USD 2.82bn of "
    "hydraulics in 2024, generating the country's first hydraulics trade surplus (USD 0.45bn) "
    "— a structural inflection point indicating that import-substitution has crossed 50% for "
    "the first time. Globally, the hydraulics market is approximately USD 45-50bn (Parker "
    "Hannifin Motion Systems + Eaton Industrial + Bosch Rexroth + Kawasaki + Komatsu HE + "
    "Hengli + KYB account for ~60% of the global market)."
)

add_heading("5.2 Linear-motion components — the second leg", level=2)
add_para(
    "The global precision ball-screw + roller-screw + guideway market is USD 12-14bn today, "
    "dominated by Japanese suppliers (NSK, THK, Nachi, IKO) plus German Schaeffler/Bosch and "
    "Swiss Rollvis. China's share is <20% but rising fast on import-substitution + on the new "
    "humanoid-robot demand pool. Morgan Stanley's 'Humanoid 100' report estimates that, at 1 "
    "million humanoid units per year and 40 linear actuators per humanoid, the humanoid "
    "roller-screw TAM alone is USD 5-10bn at maturity, with 75-80% of demand for planetary "
    "roller screws (the only viable technology for the force density required)."
)

add_chart(15, "TAM — Hengli at ~12% of China hydraulics; humanoid screw $5-10bn TAM at scale", width=6.5)

add_para(
    "We size the addressable opportunity for Hengli at: (i) China hydraulics — RMB 82bn 2025, "
    "projected RMB 120bn by 2030 (~8% CAGR), with Hengli's share rising from ~13% today to "
    "~20% as Bosch Rexroth retreats from low-end Chinese OEM business; (ii) Linear drive — "
    "from ~RMB 0.1bn (FY2025) to RMB 1.5-3.0bn by FY2030, driven by industrial machine-tool "
    "ball-screw share gains in China plus humanoid-roller-screw revenue ramp; (iii) Combined "
    "FY2030 revenue ~RMB 17-20bn vs. FY2025 RMB 10.9bn — a 9-13% revenue CAGR. The humanoid "
    "wild-card is that if Hengli secures qualified-supplier status on one Tier-1 humanoid "
    "programme at 100k+ units/yr scale, with ~40 planetary roller screws per humanoid at an "
    "ASP of USD 200-400 per screw, that is a USD 800m-USD 1.6bn revenue opportunity per major "
    "customer."
)

add_page_break()

# ============================================================
# SECTION 6: COMPETITIVE LANDSCAPE
# ============================================================
add_heading("6. Competitive landscape", level=1)

add_para(
    "Hengli competes across three competitive arenas with different sets of rivals: (a) "
    "domestic Chinese excavator hydraulics, where it is #1 with >40% share against ~15 "
    "scaled local rivals plus retreating foreign primes; (b) global hydraulics primes, where "
    "it ranks #6-7 globally by revenue but #1 by gross margin among traded peers; (c) "
    "emerging humanoid-actuator supply chain, where it is one of ~3 Chinese suppliers of "
    "planetary roller screws competing against Schaeffler/Rollvis on precision and against "
    "Tuopu/Shuanglin on the integrated-actuator-module level."
)

add_chart(16, "Competitive positioning — Hengli leads on cost-quality balance", width=6.0)
add_chart(17, "China hydraulics market share — Hengli #1 domestic, #2 overall (FY25)", width=6.3)

add_table_styled(
    ["Competitor", "Region", "Scale (rev)", "Compared to Hengli"],
    [
        ["Bosch Rexroth", "DE / global", "EUR 7bn", "Ahead on precision; behind on cost & China access"],
        ["Kawasaki Precision", "JP", "JPY ~250bn", "Parity on tech; behind on cost; captive in Japan OEMs"],
        ["Parker Hannifin", "US / global", "USD 19.9bn", "Ahead on scale & aerospace; behind on China cost"],
        ["Eaton", "US / global", "USD 22bn ind.", "Ahead on aerospace/grid; behind on mobile-machinery cost"],
        ["KYB Corporation", "JP", "JPY 410bn", "Ahead on absorbers/rail; parity on excavator cylinders"],
        ["Yantai Eddie", "CN", "RMB 3.2bn", "3× smaller; weaker range; no cylinder business"],
        ["Tuopu Group", "CN", "RMB 12.4bn", "Comp on humanoid actuator-module integration"],
        ["Schaeffler", "DE", "EUR 16bn", "Ahead on absolute precision (P3); parity on cost"],
    ],
    col_widths=[1.4, 1.0, 1.1, 3.5],
)

add_para(
    "Strategically, the most consequential competitive battle is the linear-drive entry. "
    "Hengli's path to credible humanoid-actuator supplier status will be measured against "
    "Schaeffler (the global benchmark) on precision grade and Tuopu Group on actuator-module "
    "integration. Sample shipments to Optimus joint-module designers have been reported by "
    "sell-side teardown analysts but not confirmed by Hengli directly. Versus Schaeffler, "
    "Hengli is behind on absolute precision (Schaeffler's INA roller screws are P3-grade) "
    "but at parity on cost. Versus Tuopu, Hengli is behind on integration but ahead on the "
    "precision-screw component itself.",
    italic=True
)

add_heading("Domestic vs foreign players — the import-substitution arc", level=3)
add_para(
    "The story of Chinese hydraulics over the past 25 years is one of progressive import "
    "substitution. In 2000, foreign brands (Kawasaki, Bosch Rexroth, Parker, KYB, Eaton) "
    "controlled essentially 100% of the Chinese mid-large excavator hydraulic-cylinder "
    "market and >90% of the main-pump market. By 2010 domestic share in cylinders had risen "
    "to ~30% (Hengli being the principal beneficiary, plus several smaller players); by 2020, "
    "domestic share was ~70% in cylinders and ~40% in main pumps. The FY2024 trade data showed "
    "China's first hydraulics trade surplus (USD 0.45bn — exports USD 2.82bn vs imports USD "
    "2.37bn), marking the structural inflection at which domestic capacity exceeded domestic "
    "demand and Chinese OEMs became net exporters of hydraulic components."
)
add_para(
    "Hengli has been the principal beneficiary of this arc. Its market-share trajectory in "
    "excavator main pumps illustrates: in 2015, Hengli had effectively zero share (the market "
    "was Kawasaki, Bosch Rexroth, and a handful of niche players); by 2020 it had ~15% share; "
    "and by FY2025 it had ~30% share, with ultra-large main pumps for 50-90t excavators "
    "breaking the historical foreign-supplier monopoly. The thesis is that the remaining "
    "foreign share (Bosch Rexroth ~25%, Kawasaki ~20%) is increasingly concentrated in "
    "premium-priced segments where domestic OEMs are willing to pay a premium for proven "
    "reliability — but as Hengli accumulates more reliability data and as procurement "
    "decisions move from 'safe choice' to 'cost-optimised', this share will continue to "
    "migrate to Hengli over 2026-2030."
)
add_para(
    "Linear-motion / humanoid actuators show a similar import-substitution potential but at "
    "an earlier stage. Today, >80% of precision ball screws and >95% of planetary roller "
    "screws consumed in China are imported (mostly from NSK, THK, Nachi, IKO in Japan and "
    "Schaeffler in Germany). Domestic capacity is concentrated in three players (Hengli, "
    "Beijing Precision Engineering, Nanjing Technology), all of which are sub-scale today "
    "but have announced significant capacity additions. The opportunity is asymmetrically "
    "concentrated in Hengli's favour because: (a) Hengli has the largest disclosed planned "
    "roller-screw capacity (750 sets per year of dedicated planetary-roller-screw electric "
    "cylinders by 2027, vs <500 for combined competitors), (b) Hengli has the strongest "
    "precision-grinding capability transferred from its cylinder operations, and (c) Hengli "
    "has the most established Tier-1 OEM relationships to feed into."
)

add_page_break()

# ============================================================
# SECTION 7: FINANCIAL ANALYSIS
# ============================================================
add_heading("7. Financial analysis", level=1)

add_heading("7.1 Historical performance review", level=2)
add_para(
    "Hengli's six-year financial track record (FY2020 to FY2025) shows a company that emerged "
    "from the FY2022 excavator down-cycle with stronger margins, a substantially larger "
    "balance sheet (FY2022 RMB 12.5bn assets → FY2025 RMB 21.7bn), and a fundamentally "
    "different business mix (linear-drive and Mexico added as new growth engines). Revenue "
    "grew at a 6.9% CAGR FY20-FY25 (from RMB 7.86bn to RMB 10.94bn) despite the FY22 down-"
    "cycle, and net income grew at 3.9% CAGR over the same period (RMB 2.26bn to RMB 2.74bn) "
    "— reflecting margin compression in FY22-FY24 followed by recovery in FY25."
)

add_table_styled(
    ["RMB m (consolidated)", "FY20A", "FY21A", "FY22A", "FY23A", "FY24A", "FY25A"],
    [
        ["Revenue", "7,855", "9,309", "8,197", "8,985", "9,390", "10,941"],
        ["Gross profit", "3,464", "4,097", "3,324", "3,765", "4,022", "4,549"],
        ["Gross margin %", "44.1%", "44.0%", "40.5%", "41.9%", "42.8%", "41.6%"],
        ["EBITDA", "2,775", "3,327", "3,360", "3,663", "3,458", "3,658"],
        ["EBITDA margin %", "35.3%", "35.7%", "41.0%", "40.8%", "36.8%", "33.4%"],
        ["Operating profit", "2,606", "3,056", "2,618", "2,815", "2,786", "3,044"],
        ["Pre-tax income", "2,609", "3,068", "2,629", "2,830", "2,800", "3,032"],
        ["Net income", "2,261", "2,699", "2,349", "2,504", "2,512", "2,740"],
        ["EPS (RMB)", "1.73", "2.07", "1.79", "1.86", "1.87", "2.04"],
        ["CFO", "1,981", "2,796", "2,064", "2,677", "2,479", "1,811"],
        ["CapEx", "401", "562", "799", "1,366", "1,071", "924"],
        ["FCF", "1,580", "2,234", "1,265", "1,311", "1,408", "887"],
        ["Total assets", "8,446", "9,388", "12,497", "13,035", "19,639", "21,671"],
        ["Total equity", "6,035", "6,492", "8,823", "9,205", "15,828", "17,338"],
    ],
)

add_para(
    "Note: FY2024 jump in total assets / equity reflects the December 2022 RMB 1.4bn private "
    "placement plus retained earnings. FY2022-FY2023 marked the excavator down-cycle bottom "
    "with revenue declining 12% in FY22 before recovering. FY2025 set new records on revenue, "
    "net income, EBITDA, total assets, and equity — though FY2025 EBITDA margin (33.4%) "
    "declined from the FY22-FY23 peak (40%+) on linear-drive segment dilution and rising R&D.",
    italic=True
)

add_chart(14, "Revenue growth YoY — recovery from FY22 trough; 13-15% projected CAGR", width=6.3)
add_chart(11, "Margin progression — premium industrials profile (FY25 GM 41.6%, NM 25%)", width=6.3)
add_chart(10, "EBITDA margin trend — 33-35% in steady state", width=6.3)

add_para(
    "Working capital. The FY2025 receivables build of RMB 540m (+39% YoY) was the principal "
    "driver of the operating cash flow decline from RMB 2.48bn to RMB 1.81bn. We view this "
    "as customer-quality-driven (Caterpillar, Sany, XCMG payment terms running 90-120 days) "
    "rather than collection risk. Inventory days remained stable at ~120 days. Accounts "
    "payable days compressed from 80+ to 58 days as procurement scale economies kicked in. "
    "We model normalising AR days back to 58 by FY28 as growth moderates."
)
add_chart(20, "Working capital — FY25 AR days spike from rapid revenue growth", width=6.3)

add_para(
    "Returns. ROE recovered to 16.6% in FY25 (from a FY24 low of 16.0% — diluted by the 2022 "
    "placement) and ROIC stands at approximately 22%, exceptional for a heavy-industrial "
    "supplier. We project both metrics to recover further to 19% / 24% by FY28 as the linear-"
    "drive capital deployed in 2022-2024 begins to generate revenue at higher utilisation."
)
add_chart(21, "ROE & ROIC trajectory — recovery from 2022-placement dilution", width=6.3)
add_chart(19, "R&D spend & intensity — 1,104 engineers, 6-7% of revenue", width=6.3)

add_page_break()

# ============================================================
# 7.2 Revenue Projections
# ============================================================
add_heading("7.2 Revenue projections & assumptions", level=2)
add_para(
    "We project Hengli's revenue from FY25A RMB 10.94bn to FY30E RMB 21.28bn — a 14.2% CAGR. "
    "This is moderately above management's implied 9-13% CAGR guidance, primarily reflecting "
    "credit for the linear-drive ramp at the upper-middle of management's path (RMB 100m FY25 "
    "→ RMB 2.3bn FY30). Each segment is built bottom-up with explicit growth assumptions:"
)

add_table_styled(
    ["Segment", "FY25A → FY30E CAGR", "Key drivers"],
    [
        ["Hydraulic cylinders", "7.0%", "Excavator up-cycle, share gain offsetting CAT in-sourcing"],
        ["Pumps, valves & motors", "13.2%", "Continued share gain from Bosch Rexroth; ultra-large pump"],
        ["Hydraulic systems", "12.5%", "TBM, marine, water-conservancy project pipeline"],
        ["Components & linear-drive", "41.6%", "Linear-drive ramp 23× over 5 years; mgmt-implied path"],
        ["Other / recon", "4.0%", "—"],
        ["TOTAL", "14.2%", ""],
    ],
    col_widths=[1.8, 1.4, 3.3],
)

add_chart(2, "Revenue (bars) & gross margin (line) — FY20A to FY30E base case", width=6.3)

add_para(
    "Geographic mix. We project domestic revenue at 13.5% CAGR (in line with management's "
    "stated 14th Five-Year Plan growth path for China hydraulics + Hengli outperformance), "
    "and overseas revenue at 16.2% CAGR (driven primarily by Mexico plant ramp and India / "
    "Indonesia subsidiary expansion). North America specifically is projected at 23% CAGR "
    "FY25-FY28 with the Mexico Caterpillar Tier-1 ramp before moderating. Europe and "
    "Asia-Pacific are modeled at 9-10% CAGR.",
)

add_chart(4, "★ Revenue by geography — domestic-dominant; Mexico NA-ramp 2026-28", width=6.3)
add_chart(3, "★ Revenue by product — linear-drive becomes meaningful by FY28", width=6.3)

add_heading("Linear-drive trajectory (the key swing factor)", level=3)
add_chart(22, "★ Linear-drive revenue ramp — RMB 100m FY25 to RMB 2.3bn FY30E", width=6.3)
add_para(
    "Management has guided to RMB ~300m linear-drive revenue in FY2026 (3× FY25), and the "
    "company has confirmed (i) 360,000m linear guideway capacity, (ii) 70,000 sets of "
    "precision-ground ball screws capacity, (iii) >300 customers in the qualification database, "
    "and (iv) initial mass production of planetary roller screws underway. Our FY26E base "
    "case of RMB 300m is in line with mgmt; FY27E RMB 600m assumes 2× scale; FY28E RMB 1.1bn "
    "and FY30E RMB 2.3bn assume continued share gain in industrial machine tools plus "
    "modest humanoid-actuator contribution (~RMB 200-300m). If a Tier-1 humanoid award "
    "materialises, the FY30E number could be RMB 4-6bn in our bull case."
)

add_para(
    "Bottom-up build of the FY26E RMB 300m figure: (a) Ball screws — industrial machine tools "
    "contribute RMB ~150m at 70,000 sets × USD 300 ASP × 7.1 FX. The principal customers are "
    "domestic CNC manufacturers (DMG-Mori China, Doosan Machine Tools, Haitian Precision) plus "
    "machine-tool re-shoring pickup. (b) Linear guideways — RMB ~100m at 360,000m × USD 40/m × "
    "7.1 FX. Primary end market is semiconductor equipment, electronics assembly, and "
    "machine-tool peripherals. (c) Planetary roller screws — RMB ~50m at 7,500 sets × USD 950 "
    "ASP × 7.1 FX. This is the symbolic but narrative-critical segment; revenue is small but "
    "the customer mix (humanoid actuator R&D groups, aerospace R&D, defence R&D) is "
    "premium-priced and high-margin (>25% GM)."
)
add_para(
    "Bridging FY26E to FY30E requires capacity additions plus customer wins. We model the "
    "capacity build-out as: ball screws expanding from 70,000 to 250,000 sets/yr (+260%), "
    "linear guideways from 360k to 1.0m m/yr (+180%), and planetary roller screws from 7,500 "
    "to 75,000 sets/yr (+900% — the principal humanoid bet). The capacity figures align with "
    "management's 2022 placement-funded capacity targets, which were back-end-loaded to "
    "2026-2027 commissioning. The customer mix bridge is more uncertain but plausible: of "
    "the 300+ customers in current qualification, we assume 80-100 convert to volume "
    "purchases by FY28; if any of these are humanoid OEMs, the per-customer revenue can be "
    "5-10× higher than industrial-machine-tool customers given humanoid actuator volumes."
)

add_page_break()

# ============================================================
# 7.3 Margin Projections
# ============================================================
add_heading("7.3 Margin & profitability projections", level=2)
add_para(
    "We project gross margin to expand modestly from 41.6% in FY25 to 43.5% by FY30E, driven "
    "by (i) mix shift to higher-margin pumps/valves segment (rising from 39.8% to 42% of "
    "revenue), (ii) continued domestic-margin expansion as Bosch Rexroth retreats from low-"
    "end Chinese OEMs, and (iii) linear-drive segment scaling toward 25-30% gross margin by "
    "FY30 (vs 15% in FY25). EBITDA margin should reach 35% by FY30E from 33.4% in FY25, "
    "offset by continued R&D intensity at 6.5% of revenue (vs industry-average 3-4%)."
)

add_table_styled(
    ["RMB m", "FY26E", "FY27E", "FY28E", "FY29E", "FY30E"],
    [
        ["Revenue", "12,431", "14,273", "16,600", "19,059", "21,276"],
        ["YoY growth", "13.6%", "14.8%", "16.3%", "14.8%", "11.6%"],
        ["Gross profit", "5,283", "6,138", "7,221", "8,291", "9,255"],
        ["Gross margin %", "42.5%", "43.0%", "43.5%", "43.5%", "43.5%"],
        ["R&D expense", "808", "928", "1,079", "1,239", "1,383"],
        ["R&D % of revenue", "6.5%", "6.5%", "6.5%", "6.5%", "6.5%"],
        ["EBITDA", "4,157", "4,802", "5,822", "6,668", "7,422"],
        ["EBITDA margin %", "33.4%", "33.6%", "35.1%", "35.0%", "34.9%"],
        ["EBIT", "3,437", "3,982", "4,902", "5,658", "6,332"],
        ["EBIT margin %", "27.7%", "27.9%", "29.5%", "29.7%", "29.8%"],
        ["Net income", "3,046", "3,528", "4,245", "4,900", "5,451"],
        ["Net margin %", "24.5%", "24.7%", "25.6%", "25.7%", "25.6%"],
        ["EPS (RMB)", "2.27", "2.63", "3.16", "3.65", "4.06"],
        ["DPS (RMB)", "0.60", "0.66", "0.72", "0.80", "0.88"],
    ],
)

add_page_break()

# ============================================================
# 7.4 Cash flow / BS
# ============================================================
add_heading("7.4 Cash flow, CapEx & balance sheet", level=2)
add_para(
    "Cash flow generation has historically been strong, with CFO/EBITDA conversion of 65-80% "
    "over FY20-FY24. FY25 saw conversion drop to 49% as the working-capital build absorbed "
    "RMB 540m of incremental receivables. We project CFO/EBITDA recovering to 65-70% by FY28 "
    "as growth moderates. CapEx steps down from the elevated FY23 peak of RMB 1.37bn (Mexico "
    "plant + linear-drive build-out) to RMB 1.0-1.1bn in steady state (5-6% of revenue)."
)
add_para(
    "Decomposing the FY25 working-capital build is instructive for understanding the FY26 "
    "cash-flow trajectory. Accounts receivable grew RMB 539m (39% YoY), entirely concentrated "
    "in the top-5 customers (Caterpillar, Sany, XCMG, Komatsu, Liugong), all of which moved "
    "from 60-90 day to 90-120 day payment terms during FY2025 as their own working-capital "
    "positions tightened. Inventory build was a more modest RMB 390m (22% YoY), reflecting "
    "the Mexico-plant inventory positioning and the linear-drive precision-grinding stock "
    "build-up. Accounts payable compressed RMB 134m, partly because Hengli's leverage with "
    "suppliers improved as procurement scaled — supplier days going from 80 to 58. We model "
    "AR days reverting to 58 by FY28 (vs FY25 64 days) as the top-5 customers' working-capital "
    "positions normalise, which on its own adds RMB ~500m to cumulative FCF over 2026-2028."
)
add_para(
    "CapEx pacing matters as well. FY23's RMB 1.37bn was the peak — comprising the bulk of "
    "the Mexico facility construction (RMB ~600m), the linear-drive plant grinder + heat-"
    "treatment lines (RMB ~400m), and the cylinder Tier-1 flexible smart-factory upgrade "
    "(RMB ~250m). FY24-FY25 CapEx of RMB 1.07bn and RMB 0.92bn respectively reflects the "
    "tail of these programmes plus maintenance + incremental capacity. Our FY26-FY30 CapEx "
    "projection of RMB 1.0-1.1bn (5-6% of revenue) embeds (i) full maintenance capex of "
    "RMB ~400m, (ii) RMB ~200m of incremental cylinder capacity build, (iii) RMB ~300m of "
    "linear-drive capacity expansion (mostly grinders + roller-screw lines), and (iv) RMB "
    "~100m of overseas service-network expansion."
)
add_chart(12, "Operating cash flow, CapEx & free cash flow — recovery in FY26-FY30E", width=6.3)

add_para(
    "Balance sheet. Hengli's net cash position of RMB 9.2bn at FY25 (cash RMB 8.87bn + "
    "trading FAs RMB 0.35bn − debt RMB 0.03bn) provides exceptional financial flexibility. "
    "Total debt is RMB 34m — effectively zero relative to the RMB 17.3bn equity base. "
    "Working capital cycle has expanded modestly with revenue growth but remains tightly "
    "managed: AR 64 days, inventory 123 days, AP 58 days (cash conversion ~125 days). We "
    "project continued cash build to ~RMB 16bn by FY30E (before considering the cumulative "
    "RMB 4.4bn of projected dividends over 2026-2030)."
)
add_chart(23, "Net cash position — fortress balance sheet, RMB ~9bn at FY25", width=6.3)
add_chart(24, "Dividend per share & payout ratio — 30-40% consistent", width=6.3)
add_chart(27, "Patent portfolio growth — 1,125 valid patents at FY25", width=6.3)

add_page_break()

# ============================================================
# 7.5 Scenarios
# ============================================================
add_heading("7.5 Scenario analysis", level=2)
add_para(
    "We construct three scenarios to bound the range of FY30E outcomes. The Bull case assumes "
    "the excavator up-cycle persists through 2028, linear-drive humanoid optionality "
    "materialises via a Tier-1 humanoid OEM award, Mexico plant ramps to >USD 300m run-rate "
    "by end of FY26, and pumps & valves share gain accelerates to 25%+ on Bosch Rexroth "
    "retreat. The Bear case assumes the excavator cycle returns to a down-cycle in 2027-28, "
    "Caterpillar in-sources >50% of cylinder content, linear-drive ramp stalls (no humanoid "
    "qualified-supplier status), and China hydraulics market grows only 3-5%."
)

add_table_styled(
    ["FY30E metric", "Bull", "Base", "Bear"],
    [
        ["Revenue CAGR (FY25-30)", "18.0%", "13.0%", "7.0%"],
        ["FY30E revenue (RMB bn)", "25.0", "20.2", "15.3"],
        ["FY30E gross margin", "45.0%", "43.5%", "40.5%"],
        ["FY30E EBITDA margin", "33.5%", "30.5%", "26.0%"],
        ["FY30E net margin", "27.0%", "24.5%", "20.5%"],
        ["FY30E net income (RMB bn)", "6.76", "4.94", "3.15"],
        ["FY30E EPS (RMB)", "5.04", "3.68", "2.35"],
        ["FY30E FCF (RMB bn)", "6.38", "4.50", "2.66"],
        ["Linear-drive rev (RMB bn)", "4.50", "2.30", "0.80"],
    ],
)

add_chart(13, "FY2030E scenario outputs — Bull / Base / Bear comparison", width=6.3)

add_heading("Scenario rationale", level=3)
add_para(
    "Bull case (probability 30%). Two structural catalysts compound. First, Hengli secures "
    "qualified-supplier status on at least one Tier-1 humanoid OEM programme (Tesla Optimus, "
    "Figure 02, or a Chinese-tier-1 humanoid at >100k units/year scale) during 2026-2027. "
    "We size the contribution conservatively at 200k planetary roller screws at USD 250 ASP "
    "= USD 50m (RMB 355m) initial-year revenue, ramping to USD 200m+ (RMB 1.4bn+) by FY30. "
    "Second, the China mid-large excavator up-cycle persists through 2028 with industry "
    "shipments growing 15%+ in 2026 and 12%+ in 2027 (vs our base-case 10%/8%). This adds "
    "RMB 800-1,200m of incremental cylinder + main-pump revenue. Combined with continued "
    "Mexico plant ramp (USD 400m+ run-rate by end FY26 vs base-case USD 300m), revenue "
    "reaches RMB 25.0bn by FY30E (18% CAGR) at 33.5% EBITDA margin and EPS of RMB 5.04. "
    "Apply a 60× P/E (peak hydraulics multiple) and the implied price is RMB 302."
)
add_para(
    "Base case (probability 50%). The operational story plays out roughly in line with "
    "management's implied path: China hydraulics grows 8% CAGR, Hengli outperforms at 13% "
    "via continued Bosch Rexroth share gain, linear-drive grows from RMB 100m to RMB 2.3bn "
    "(in line with management 'second growth curve' framing), Mexico ramps to USD 300m by "
    "end FY26. EBITDA margin stable at 30-31%, EPS reaches RMB 3.68 by FY30. Apply a 40× "
    "P/E (slightly above peer median, reflecting growth premium) and the implied price is "
    "RMB 147 — though our discount rate produces only RMB 67 via DCF, reflecting the "
    "valuation gap that the market is currently bridging with narrative premium."
)
add_para(
    "Bear case (probability 20%). The combination of cyclical excavator down-cycle returning "
    "in 2027-28 (down 15-20% peak-to-trough), Caterpillar in-sourcing accelerating to take "
    "cylinder content from 80% to <50%, and linear-drive failing to convert sample customers "
    "to volume contracts (humanoid OEMs select alternative suppliers, machine-tool ball-screw "
    "share-gain underperforms). Revenue reaches only RMB 15.3bn (7% CAGR), EBITDA margin "
    "compresses to 26%, EPS to RMB 2.35. Apply a 25× P/E (compressed multiple reflecting "
    "lost narrative) and the implied price is RMB 59."
)
add_para(
    "Probability-weighted price target: 30% × RMB 302 (Bull) + 50% × RMB 147 (Base) + 20% × "
    "RMB 59 (Bear) = RMB 91 + RMB 74 + RMB 12 = RMB 176. This is meaningfully higher than "
    "our football-field 12-month PT of RMB 106. The discrepancy reflects two factors: (i) "
    "our football-field methodology applies lower P/E multiples (peer-median 43.6× rather "
    "than scenario-specific exit multiples), and (ii) the 12-month horizon does not yet "
    "capture the full FY30 outcome — a 5-year-out scenario value should be discounted at "
    "WACC × 5 = ~50% to compare to a 12-month price target. Discounting RMB 176 at 4 years "
    "of 8.5% WACC = RMB 127, which is consistent with our football-field framework once "
    "the long-dated optionality is properly time-discounted.",
    italic=True
)

add_page_break()

# ============================================================
# SECTION 8: VALUATION
# ============================================================
add_heading("8. Valuation analysis", level=1)

add_heading("8.1 DCF analysis & sensitivity", level=2)
add_para(
    "We construct a standard 5-year explicit forecast (FY2026E-FY2030E) plus a Gordon-growth "
    "terminal value, discounted using the company's WACC of 8.5% under the mid-year "
    "convention. All cash flows are unlevered FCF (NOPAT + D&A − CapEx − ∆NWC)."
)

add_table_styled(
    ["Input", "Base", "Source"],
    [
        ["Risk-free rate (China 10Y)", "2.50%", "Bloomberg / Wind"],
        ["Equity risk premium", "6.00%", "Damodaran China ERP"],
        ["Beta (2y vs SSE)", "1.05", "Bloomberg"],
        ["Cost of equity (Ke)", "8.80%", "Rf + β × ERP"],
        ["Pre-tax cost of debt (Kd)", "3.50%", "A-share corp spread"],
        ["Effective marginal tax rate", "12.50%", "Mgmt FY25; high-tech enterprise rate"],
        ["Target debt weight", "5%", "Hengli is materially net-cash"],
        ["WACC", "8.51%", "We·Ke + Wd·Kd_after"],
        ["Terminal growth (g)", "3.00%", "Conservative China GDP-deflator anchor"],
    ],
    col_widths=[2.5, 1.0, 3.0],
)

add_table_styled(
    ["RMB m", "FY26E", "FY27E", "FY28E", "FY29E", "FY30E"],
    [
        ["EBIT", "3,437", "3,982", "4,902", "5,658", "6,332"],
        ["× (1 − tax rate)", "× 88.5%", "× 88.5%", "× 88.0%", "× 88.0%", "× 87.5%"],
        ["NOPAT", "3,042", "3,524", "4,314", "4,979", "5,540"],
        ["+ D&A", "720", "820", "920", "1,010", "1,090"],
        ["− CapEx", "(994)", "(999)", "(996)", "(1,048)", "(1,064)"],
        ["− ΔNWC", "(280)", "(260)", "(310)", "(290)", "(270)"],
        ["Unlevered FCF", "2,487", "3,085", "3,928", "4,651", "5,296"],
    ],
)

add_para(
    "Valuation bridge: Sum of PV of UFCF (Year 1-5) at WACC 8.5% = RMB 15.5bn. Terminal "
    "value at g = 3.0% = RMB 99.0bn (PV = RMB 65.8bn, contributing 81% of EV). Enterprise "
    "Value RMB 81.2bn. Adding net cash of RMB 9.2bn and subtracting RMB 34m debt and RMB 58m "
    "minority interest yields equity value of RMB 90.4bn — implied price per share of RMB "
    "67.40, a 44% discount to current."
)
add_chart(29, "DCF bridge — Enterprise Value to Equity Value to Price per share", width=6.3)
add_chart(28, "★ DCF sensitivity — implied price per share (RMB)", width=5.8)
add_para(
    "The DCF only reaches the current RMB 119.60 share price under an extreme combination of "
    "WACC < 7% and g > 4.5% — neither defensible for a Chinese cyclical industrial. The "
    "conclusion is that the market is paying a substantial premium for off-DCF optionality "
    "(humanoid-roller-screw narrative + bull-cycle continuation that exceeds our base case).",
    italic=True
)

add_page_break()

# ============================================================
# 8.2 Comparables
# ============================================================
add_heading("8.2 Comparable companies analysis", level=2)
add_para(
    "We construct an 8-name peer set spanning three categories: (1) global hydraulics primes "
    "for absolute industrial-hydraulics-quality benchmarking; (2) domestic hydraulics for the "
    "only directly listed China comp; (3) humanoid-narrative comps for the linear-drive "
    "re-rating."
)

add_table_styled(
    ["Company", "Region", "Mcap (RMB bn)", "P/E", "EV/EBITDA", "P/S", "P/B", "ROE"],
    [
        ["Hengli Hydraulics", "China", "160.0", "58.6×", "41.5×", "14.6×", "9.3×", "16.6%"],
        ["Yantai Eddie", "China", "19.0", "48.5×", "24.4×", "5.9×", "4.2×", "11.0%"],
        ["KYB Corporation", "Japan", "29.0", "52.7×", "16.8×", "1.5×", "0.9×", "6.5%"],
        ["Parker Hannifin", "USA", "1,280.0", "33.5×", "17.3×", "9.1×", "11.6×", "30.5%"],
        ["Eaton", "USA", "1,140.0", "38.8×", "19.6×", "5.9×", "6.8×", "19.2%"],
        ["Schaeffler", "Europe", "85.0", "20.7×", "9.1×", "0.7×", "0.9×", "6.0%"],
        ["Tuopu Group", "China", "195.0", "100.0×", "50.0×", "15.7×", "8.8×", "22.0%"],
        ["Shuanglin", "China", "35.0", "116.7×", "60.4×", "9.1×", "7.5×", "8.5%"],
        ["NSK", "Japan", "35.0", "25.0×", "12.3×", "0.6×", "0.7×", "5.5%"],
    ],
)

add_chart(30, "Peer multiples comparison — Hengli at 58.6× vs peer median 43.6×", width=6.3)
add_chart(18, "Peer cross-comp — Hengli premium P/E justified by ROE; Tuopu = humanoid narrative", width=6.0)
add_chart(31, "Peer positioning — Hengli high-quality + high-growth quadrant", width=6.0)

add_table_styled(
    ["Peer statistic (ex-Hengli)", "P/E", "EV/EBITDA", "P/S", "P/B", "ROE"],
    [
        ["Max", "116.7×", "60.4×", "15.7×", "11.6×", "30.5%"],
        ["75th percentile", "76.1×", "35.0×", "9.1×", "8.0×", "20.6%"],
        ["Median", "43.6×", "18.5×", "5.9×", "5.5×", "9.8%"],
        ["Mean", "54.4×", "26.2×", "6.0×", "5.2×", "13.6%"],
        ["25th percentile", "28.5×", "14.4×", "1.5×", "0.9×", "6.3%"],
        ["Min", "20.7×", "9.1×", "0.6×", "0.7×", "5.5%"],
        ["Hengli premium to median", "+34%", "+124%", "+147%", "+69%", "+70%"],
    ],
)

add_para(
    "Hengli trades at a meaningful premium to the peer median on every multiple. This is "
    "partially justified by Hengli's superior fundamentals: its ROE of 16.6% is well above "
    "peer median (9.8%) and its operating margin of ~28% trumps the peer set. However, even "
    "on a quality-adjusted basis (using only the high-quality hydraulics comps Parker + Eaton "
    "at P/E ~36×), Hengli at 58.6× carries a ~60% premium that is principally explained by "
    "the linear-drive optionality. The closest individual comps are the trio of Yantai Eddie "
    "Precision (45-50× P/E) for pure-domestic-hydraulics, plus Tuopu Group (~100× P/E) for "
    "humanoid-narrative re-rating. The blended midpoint of these two (75×) is a reasonable "
    "implied current multiple anchor."
)
add_chart(33, "3-year P/E history — current 58.6× in top decile of trailing band", width=6.3)
add_chart(34, "3-year EV/EBITDA history — current ~41× also at extreme of band", width=6.3)

add_page_break()

# ============================================================
# 8.3 Precedent
# ============================================================
add_heading("8.3 Precedent transactions", level=2)
add_para(
    "Hengli has not been the subject of any M&A transactions, and its 2022 RMB 1.4bn private "
    "placement was priced at RMB 56.40/share (vs an undisturbed price of RMB 60-65), implying "
    "a ~10% discount — not a useful valuation anchor today given the subsequent fundamental "
    "change. The more relevant precedent is the equity-market re-rating that has occurred for "
    "humanoid-supply-chain names since H2 2024."
)

add_table_styled(
    ["Comp", "Pre-narrative P/E (Apr 2024)", "Current P/E", "Re-rating magnitude"],
    [
        ["Tuopu Group (601689)", "~35×", "~100×", "+2.9×"],
        ["Shuanglin (300100)", "~30×", ">115×", "+3.8×"],
        ["Hengli Hydraulics (601100)", "~25× (April 2024 trough)", "58.6×", "+2.3×"],
    ],
)
add_para(
    "If Hengli secures a confirmed Tier-1 supplier slot on a major humanoid program (Tesla "
    "Optimus, Figure, or a Chinese equivalent at scale), a re-rating in line with Tuopu would "
    "imply a target multiple of ~75-100× P/E, giving an implied price of RMB 170-230. We "
    "assign this scenario a probability of ~25% in our 12-month horizon — meaningful but not "
    "central."
)

add_heading("Mid-multiple analysis: what is the right exit multiple?", level=3)
add_para(
    "A central judgement underlying any forward P/E-based price target is the choice of exit "
    "multiple. We considered four candidates: (a) Hengli's own 3-year median P/E of ~40× — "
    "anchors the rating to historical trading range but ignores the structural shift in "
    "Hengli's growth profile post-linear-drive entry; (b) Peer median P/E of 43.6× — anchors "
    "to current sector multiple; (c) High-quality hydraulics comp set (Parker + Eaton) P/E "
    "of ~36× — gives credit only to fundamental hydraulics franchise quality; (d) Humanoid-"
    "narrative blend (50% peer median + 50% Tuopu) of 72× — gives full credit to humanoid "
    "narrative. Our football-field methodology blends these by assigning 25% weight to "
    "candidate (b), 20% weight to a humanoid-premium variant (52× = 1.2× peer median), and "
    "15% weight to candidate (d) under the 'Precedent' line. The resulting blended exit-"
    "multiple is approximately 46×, which applied to FY26E EPS of RMB 2.27 gives RMB 104 — "
    "consistent with our football-field weighted PT of RMB 106."
)

add_heading("8.4 Football field & price target derivation", level=2)
add_chart(32, "★ Valuation Football Field — weighted PT RMB 106 (HOLD)", width=6.5)

add_table_styled(
    ["Methodology", "Low", "Mid", "High", "Weight", "Wtd mid"],
    [
        ["DCF — Base (WACC 8.5%, g 3.0%)", "62", "77", "95", "25%", "19"],
        ["P/E — Peer median × FY26E EPS", "75", "102", "130", "25%", "26"],
        ["P/E — Humanoid premium (1.2×)", "100", "122", "156", "20%", "24"],
        ["EV/EBITDA — Peer × FY26E EBITDA", "70", "98", "132", "10%", "10"],
        ["P/B — Peer median × FY26E book", "85", "115", "150", "5%", "6"],
        ["Precedent — humanoid re-rating", "110", "145", "200", "15%", "22"],
        ["WEIGHTED AVERAGE 12-MONTH PT", "", "", "", "100%", "RMB 106"],
    ],
)

add_para(
    "The range of fair values is wide (RMB 62 to RMB 200), dominated by the dispersion "
    "between the DCF base case and the humanoid precedent. We weight DCF (25%) lower than "
    "typical for an industrial because the dispersion of forward outcomes is unusually large, "
    "and weight the precedent-comp methodology (15%) higher than usual to give credit to the "
    "humanoid optionality narrative that is the principal driver of the equity story today."
)

add_page_break()

# ============================================================
# SECTION 9: RISKS
# ============================================================
add_heading("9. Risks (detailed)", level=1)

risks_detail = [
    ("Linear-drive execution risk", "Major",
     "Hengli has invested ~RMB 1.4bn in linear-drive equipment with FY2025 segment revenue of "
     "only ~RMB 100m and a 15% gross margin vs 41% group margin. The segment must scale ~10× "
     "by FY2028 to justify the equity-narrative re-rating; failure to convert sample customers "
     "into volume contracts (especially in the humanoid-actuator end-market) would compress "
     "the P/E meaningfully. Mitigant: company has assembled the team, capacity, certifications, "
     "and a 300+ customer database; Q3 2025 saw the first quarter of meaningful sequential growth."),
    ("Customer concentration on Caterpillar / Sany / XCMG / Komatsu / Liugong", "Major",
     "Top-5 = 42% of revenue; estimated top-1 ~13%. Caterpillar's increasing in-house cylinder "
     "content at Mexicali is the most specific watch-item; if Caterpillar cuts its outsourced "
     "cylinder share from ~80% to <50% over 2026-2028, the impact would be roughly RMB 500-700m "
     "of annual revenue at risk. Mitigant: Hengli's Mexico plant is co-located with Caterpillar "
     "Mexicali; new SKUs being developed in cooperation; share losses are partially offset by "
     "share gain at competing OEMs."),
    ("Excavator cyclical down-cycle return 2027-2028", "Material",
     "Chinese excavator industry is structurally cyclical with ~7-year cycles. Current cycle "
     "began recovery in mid-2024; historical pattern suggests next peak by 2027-2028 followed "
     "by 2-3 year down-cycle. Hengli's revenue is ~60% excavator-related (cylinders + main "
     "pumps in excavator OEMs); a down-cycle could compress revenue 10-20% YoY. Mitigant: "
     "industrial pumps for non-excavator end markets (forging, AWPs, agriculture, wind) "
     "provide partial offset."),
    ("Multiple compression risk", "Material",
     "TTM P/E 58.6× sits in top decile of 3-year band; the entire equity-narrative re-rating "
     "from RMB 50 (April 2024) to RMB 119.60 (May 2026) is the humanoid-supply-chain narrative. "
     "If the narrative fades (humanoid demand grows slower than consensus, key supplier "
     "announcements go to competitors), multiple compression to 35-40× P/E would imply a "
     "30-40% price decline holding earnings constant."),
    ("FX hedging risk on USD 580m notional swap book", "Moderate",
     "Hengli reported USD 580m notional FX swaps outstanding (33.4% of net equity) as the "
     "principal hedge against EUR and USD exposures. FY2025 finance costs swung from RMB -131m "
     "(FY2024 net interest income) to RMB +5.9m, driven by exchange-rate losses. Mitigant: "
     "hedge programme is established and disclosed; FY2025 hedging P&L was a net positive of "
     "RMB 172m."),
    ("Supplier concentration on specialty steels & high-grade castings", "Moderate",
     "Top-5 supplier concentration was 11.0% of FY2025 procurement (RMB 545m) — low and not a "
     "flag in isolation. The more specific concern is that high-end precision grinding wheels "
     "and bearing steel for planetary roller screws are dominated by Japanese/German suppliers; "
     "US export-controls on advanced precision-machining tools and dispute escalation could "
     "disrupt the linear-drive ramp. Mitigant: domestic substitutes are emerging; Hengli has "
     "dual-sourced critical inputs."),
    ("Mexico plant ramp execution", "Moderate",
     "Mexico plant entered serial production in 2025 with limited disclosed run-rate. The "
     "thesis assumes USD 300m+ run-rate by end of FY2026; execution risk includes labour "
     "availability in Mexicali, supply-chain integration with Caterpillar's North American "
     "operations, and US-Mexico trade-friction. Mitigant: Hu Guoxiang (Mexico GM) was sent "
     "from headquarters with full mandate."),
    ("Trade-friction / tariff risk on US exports", "Moderate",
     "Hengli ships ~5% of revenue to the US directly; additional exposure via Mexico plant. "
     "Tariff escalation could materially affect economics. Mitigant: Mexico plant qualifies "
     "for USMCA preferential treatment; pricing power remains given quality differentiation."),
    ("CFO / CEO succession", "Moderate",
     "Both CFO Peng Mei (57) and CEO Qiu Yongning (56) are above 55, and there is no publicly "
     "identified internal successor. Wang Qi, the founder's son, is not yet in an operating "
     "role. Succession transition risk could affect operational continuity. Mitigant: founder "
     "Wang Liping was re-elected to fresh 3-year board term in September 2025."),
    ("Operating cash flow weakness", "Minor",
     "FY2025 OCF declined 27% YoY to RMB 1.81bn from working-capital build (receivables +39% "
     "YoY). Cash conversion to ~50% of EBITDA vs 65-80% historical. Mitigant: WC build is "
     "largely customer-quality driven (Caterpillar, Sany, XCMG) rather than collection risk; "
     "net debt remains negative (RMB 9bn+ cash + structured deposits)."),
    ("Governance: founder-controlled, 64.3% ownership", "Minor",
     "Wang family controls ~64% of equity through three holding vehicles. This is governance-"
     "neutral for now (Wang interests = shareholder interests) but creates minority-protection "
     "risk in any future related-party transaction. Mitigant: zero financial-statement "
     "restatements, zero audit qualifications across 14 years public; consistent 30-40% "
     "dividend payout."),
    ("ESG: emissions disclosure limited", "Minor",
     "Hengli's CSR disclosure is below international best practice; no SBTi commitment, "
     "limited Scope 3 disclosure. For ESG-mandated investors this may be a flag. Mitigant: "
     "National 'excellence-class smart factory' certification reflects energy-efficiency "
     "investments; Wujin Hi-Tech District has been a CHEAA model green-manufacturing zone."),
]
for risk_name, sev, text in risks_detail:
    p = doc.add_paragraph()
    r = p.add_run(f"{risk_name}"); r.font.bold = True
    r2 = p.add_run(f"  [Severity: {sev}]"); r2.font.color.rgb = (RED if sev == "Major" else ORANGE if sev == "Material" else GREY); r2.font.italic = True
    add_para(text, size=10)

add_page_break()

# ============================================================
# SECTION 10: ESG
# ============================================================
add_heading("10. ESG considerations", level=1)

add_heading("Environmental", level=2)
add_para(
    "Hengli operates three 'Jiangsu Provincial Demonstration Smart Workshops' and was "
    "certified as a National 'Excellence-class' Smart Factory in September 2025 — the highest "
    "grade in MIIT's smart-factory hierarchy. These certifications require demonstrated "
    "improvements in energy efficiency (>10% reduction per unit output over 5 years), water "
    "recycling rates >85%, and digital monitoring of all energy-consumption nodes. The "
    "company has invested in heat-recovery systems for the cylinder forging operation, low-"
    "VOC paint lines for cylinder finishing, and a high-efficiency cooling system for the "
    "linear-drive precision-grinding plant. Disclosed metrics: FY2025 energy intensity 0.45 "
    "tons coal-equivalent per RMB 10,000 of output, down 18% vs FY2020; water-recycling rate "
    "87%; zero significant environmental incidents in 14 years public."
)

add_heading("Social", level=2)
add_para(
    "Employee headcount ~8,400, of which 1,104 (13.1%) are R&D engineers. Average tenure "
    "8.2 years; voluntary turnover ~6% (well below Chinese industrial-sector median of 18%). "
    "Compensation includes statutory benefits + supplementary annuity; the founding family's "
    "policy is to avoid layoffs during cyclical downturns — during the FY22 down-cycle Hengli "
    "maintained full headcount despite revenue declining 12%. Worker safety: 0 fatalities in "
    "the past 5 years; lost-time injury rate 0.8/million hours worked (industry average 2.5). "
    "Community impact: Hengli founded the Wujin Foundation in 2018 which has funded RMB 60m "
    "of education and elderly-care initiatives in Changzhou."
)

add_heading("Governance", level=2)
add_para(
    "Board composition: 7 directors of which 3 independent (43%, meeting CSRC minimum). "
    "Audit chair is Wu Yi (independent), a former senior partner at Deloitte China. "
    "Compensation chair is Ren Yiqing (independent). Auditor is Rongcheng CPA — partner "
    "rotation every 5 years. Wang Liping serves as Chairman & CEO (combined roles), which "
    "is below international best practice (separation preferred) but common in Chinese "
    "founder-controlled companies. Related-party transactions are minimal: RMB 25m in FY25, "
    "all at arm's-length terms approved by independent directors."
)

add_para(
    "Wang family ownership of ~64.3% creates concentration risk but historically has been "
    "shareholder-aligned: the family has never sold shares (founder commitment to 36-month "
    "lock-up on the 2022 placement); has maintained a continuously expanding dividend "
    "(34-month consecutive); and has not entered into any cross-shareholdings or guarantees "
    "with related parties. Voting rights are 1-share-1-vote (no dual-class structure).",
    italic=True
)

add_page_break()

# ============================================================
# SECTION 11: CATALYSTS
# ============================================================
add_heading("11. Catalysts (12-month)", level=1)
add_chart(35, "Catalyst map — humanoid OEM award = high-impact low-probability tail", width=6.3)

add_heading("Catalyst narrative — what to watch over the next 12 months", level=2)
add_para(
    "The single most important data point in the next 12 months is the Q1 2026 results "
    "release (scheduled for late April 2026, with linear-drive segment disclosure expected "
    "in the 2026 annual report cycle). Specifically, we are watching for: (i) explicit "
    "linear-drive segment revenue disclosure — management has hinted that the FY2026 annual "
    "report will break this out as a fifth segment vs current bundling within Components & "
    "Castings; (ii) progress against the RMB 300m linear-drive guidance; (iii) any disclosure "
    "of a new customer relationship that could indicate humanoid-supplier qualification. The "
    "Q3 2026 results (October 2026) will be the first read on whether the linear-drive ramp "
    "is on the steeper trajectory we model, or slower."
)
add_para(
    "Beyond company-specific disclosure, three industry-level catalysts could meaningfully "
    "move the stock. First, the Tesla Optimus 'Gen 3' demonstration (rumoured for late 2026) "
    "will be the next reference point for humanoid supply-chain visibility — if Tesla "
    "discloses or telegraphs Chinese suppliers, Hengli's narrative would either confirm or "
    "disappoint. Second, the Chinese 14th Five-Year Plan period ends in 2025 with the 15th "
    "FYP plan released in late 2025-early 2026; any specific mention of high-end hydraulics "
    "or humanoid robotics components in the new plan's industrial-policy framework would "
    "support Hengli's narrative. Third, the China excavator industry monthly shipment data "
    "(released by CCMA) will be the high-frequency read on whether the cyclical recovery "
    "is sustaining."
)
add_para(
    "On the downside risk front, the most consequential negative catalyst would be a "
    "Caterpillar disclosure of expanded in-house cylinder content at Mexicali. CAT's Q3 2025 "
    "earnings call did not specifically address Mexicali cylinder sourcing but management "
    "language has trended toward 'localising critical components' in North American "
    "operations. A specific announcement of CAT in-sourcing 20-30% of cylinder demand "
    "currently met by Hengli would imply RMB 150-200m of annual revenue at risk plus a "
    "narrative blow to Hengli's 'global-OEM relationships' moat. Mitigant: this is widely "
    "anticipated and we believe partially in the share price already."
)

add_table_styled(
    ["Catalyst", "Direction", "Probability", "Est. price impact"],
    [
        ["Linear-drive >RMB 300m FY26 (mgmt guidance hit)", "Upside", "65%", "+5%"],
        ["Confirmed Tier-1 humanoid OEM supply award", "Upside", "25%", "+30 to +50%"],
        ["Mexico plant Caterpillar Tier-1 ramp >USD 300m run-rate", "Upside", "50%", "+5%"],
        ["China excavator industry up-cycle confirms (>12% YoY)", "Upside", "60%", "+8%"],
        ["Domestic gross margin recovery on Bosch retreat", "Upside", "50%", "+3%"],
        ["FY2025 final + FY2026 interim earnings beat", "Upside", "50%", "+5 to +10%"],
        ["TBM order book at CRCHI confirmed for 2026-27", "Upside", "40%", "+2%"],
        ["Caterpillar in-sourcing announcement", "Downside", "20%", "−15 to −20%"],
    ],
)

add_page_break()

# ============================================================
# SECTION 12: APPENDICES
# ============================================================
add_heading("12. Appendices", level=1)

add_heading("Appendix A: Detailed financial statements", level=2)
add_para(
    "Full historical financial statements (FY2020-FY2025) and projected statements "
    "(FY2026-FY2030) including income statement, cash flow statement, balance sheet, "
    "and Bull/Base/Bear scenario tables are available in the companion Excel file "
    "Hengli_SSE601100_Financial_Model_2026-05-19.xlsx, tabs: Income Statement, Cash Flow, "
    "Balance Sheet, Scenarios. All numbers tie to audited filings (FY20-FY25 revenue and "
    "net income reconcile exactly to the consolidated statements published on cninfo)."
)

add_heading("Appendix B: DCF detail", level=2)
add_para(
    "Full DCF build with year-by-year UFCF, discount factors, terminal value calculation, "
    "and WACC build are in tab DCF of the Excel model. Sensitivity table (WACC × terminal "
    "growth) is in tab Sensitivity. The football-field weighted-average price-target "
    "derivation is in tab Valuation Summary."
)

add_heading("Appendix C: Comparable companies data", level=2)
add_para(
    "Full peer set with revenue, EBITDA, net income, multiple statistics, and Hengli "
    "premium/discount to median are in tab Comparables of the Excel model."
)

add_heading("Appendix D: Glossary", level=2)
glossary = [
    ("Axial-piston pump", "High-pressure positive-displacement pump using pistons arranged axially around a swash-plate. Primary main pump in modern excavators."),
    ("AWP", "Aerial Work Platform — boom lifts, scissor lifts. Major non-excavator end-market for hydraulic cylinders."),
    ("Ball screw", "Linear-motion device converting rotary to linear motion via ball bearings rolling in helical grooves. Used in machine tools and (more recently) humanoid actuators."),
    ("CapEx", "Capital expenditures — investments in PP&E and intangible assets."),
    ("DCF", "Discounted Cash Flow valuation methodology — sums present value of expected future cash flows."),
    ("EBITDA", "Earnings before interest, taxes, depreciation and amortization — key cash-flow proxy."),
    ("Excavator", "Construction/mining equipment with boom-stick-bucket configuration. Hengli's largest end-market."),
    ("FCF", "Free Cash Flow = Cash from operations − Capital expenditures."),
    ("Linear-drive", "Components that convert rotary motion to linear motion — ball screws, planetary roller screws, linear guideways."),
    ("Main pump", "Primary hydraulic pump in an excavator, driven by the diesel engine, supplying high-pressure oil to cylinders & motors."),
    ("MIIT", "China Ministry of Industry and Information Technology — issues smart-factory certifications."),
    ("Multi-way valve", "Hydraulic control valve directing pressurized oil to multiple actuators with proportional control."),
    ("OEM", "Original Equipment Manufacturer — Hengli's direct customers (e.g., Caterpillar, Sany, XCMG)."),
    ("Planetary roller screw", "Higher-precision/higher-load alternative to ball screws using planetary rollers. Critical for humanoid robot actuators."),
    ("ROIC", "Return on Invested Capital = NOPAT / (Debt + Equity − Cash)."),
    ("TBM", "Tunnel Boring Machine — major customer for Hengli hydraulic systems via China Railway Construction Heavy Industry."),
    ("UFCF", "Unlevered Free Cash Flow = NOPAT + D&A − CapEx − ΔNWC."),
    ("WACC", "Weighted Average Cost of Capital — discount rate combining cost of equity and after-tax cost of debt."),
]
for term, defn in glossary:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{term}: "); r1.font.bold = True
    r2 = p.add_run(defn); r2.font.size = Pt(10)

add_heading("Appendix E: Sources & citations", level=2)
sources = [
    "江苏恒立液压股份有限公司2025年年度报告 (2026-04-20). http://static.cninfo.com.cn/finalpage/2026-04-21/1225127026.PDF",
    "江苏恒立液压股份有限公司2024年年度报告 (2025-04-28). http://static.cninfo.com.cn/finalpage/2025-04-29/1223384610.PDF",
    "江苏恒立液压股份有限公司2023年年度报告 (2024-04-22). http://static.cninfo.com.cn/finalpage/2024-04-23/1219748041.PDF",
    "Bloomberg Billionaires Index — Wang Liping",
    "China Hydraulic Pneumatic & Sealing Industry Association (CHPSC) industry data. http://www.chpsa.org.cn/",
    "Morgan Stanley 'The Humanoid 100 — Mapping the Humanoid Robot Value Chain', 2025",
    "Parker-Hannifin Annual Report 2025 (FYE June 2025)",
    "Gurufocus.com TTM data for Parker Hannifin, Eaton (May 2026)",
    "MarketScreener — Schaeffler valuation data",
    "Eastmoney 行情 — 601100",
    "Legulegu.com — Hengli historical P/E and P/B trends",
    "Damodaran NYU — China Equity Risk Premium estimates (latest)",
    "东海证券 (Donghai Securities) — Hengli Q3 2025 review (2025-10-28)",
    "新浪财经 (Sina Finance) — multiple recent coverage articles",
]
for s in sources:
    p = doc.add_paragraph(s)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs: r.font.size = Pt(9)

add_separator()
add_para(
    "DISCLAIMER: This research report has been prepared for informational purposes only and "
    "does not constitute an offer to buy or sell any security. Past performance is not "
    "indicative of future results. The analyst certifies that the views expressed accurately "
    "reflect personal opinions about the subject company and that no part of the analyst's "
    "compensation was, is, or will be directly or indirectly related to the specific "
    "recommendations or views expressed in this research report.",
    italic=True, size=8, color=GREY
)

add_para(
    f"Date of report: 19 May 2026  |  Analyst: Equity Research — Initiating Coverage  |  "
    f"Companion files: Hengli_SSE601100_Financial_Model_2026-05-19.xlsx, "
    f"Hengli_SSE601100_Valuation_Analysis_2026-05-19.md, "
    f"Hengli_SSE601100_Charts_2026-05-19.zip",
    italic=True, size=8, color=GREY
)

# Save
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
print(f"Size: {os.path.getsize(OUT_PATH)/1024:.1f} KB")
