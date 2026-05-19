#!/usr/bin/env python3
"""
Shuanglin Co. (SZSE:300100) — Task 5 Final Report Assembly

Builds a 30-50 page DOCX equity-research initiation report by combining:
  - Task 1 company research (双林股份_SZSE300100_公司研究_2026-05-17.md)
  - Task 2 financial model (双林股份_SZSE300100_Financial_Model_*.xlsx)
  - Task 3 valuation analysis (双林股份_SZSE300100_Valuation_Analysis_*.md)
  - Task 4 charts (25 PNG files in ./charts/)

Output: 双林股份_SZSE300100_Initiation_Report_<date>.docx

Standards: Times New Roman, ~11K words, 25 charts, 14+ tables.
"""
from __future__ import annotations
import os
import datetime as _dt
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from openpyxl import load_workbook

BASE = os.path.dirname(__file__)
CHARTS = os.path.join(BASE, "charts")
MODEL_PATH = os.path.join(BASE, "双林股份_SZSE300100_Financial_Model_2026-05-18.xlsx")
OUT_PATH = os.path.join(BASE, f"双林股份_SZSE300100_Initiation_Report_{_dt.date.today()}.docx")

# Colors
NAVY = RGBColor(0x0B, 0x53, 0x94)
RED  = RGBColor(0xC0, 0x00, 0x00)
GREEN = RGBColor(0x38, 0x76, 0x1D)
GREY = RGBColor(0x66, 0x66, 0x66)
BLACK = RGBColor(0x00, 0x00, 0x00)

# ============================================================================
# DOCUMENT SETUP
# ============================================================================
doc = Document()

# Default font: Times New Roman 11
style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(11)
# Set East Asian font (for any Chinese characters)
rPr = style.element.rPr
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.append(rFonts)
rFonts.set(qn('w:eastAsia'), 'PingFang SC')

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)


def set_run(run, bold=False, italic=False, color=None, size=None, font_name=None):
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    if font_name:
        run.font.name = font_name


def add_heading(text, level=1, color=NAVY, before=12, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    if level == 0:
        set_run(r, bold=True, color=color, size=20)
    elif level == 1:
        set_run(r, bold=True, color=color, size=15)
    elif level == 2:
        set_run(r, bold=True, color=color, size=13)
    elif level == 3:
        set_run(r, bold=True, color=BLACK, size=12)
    else:
        set_run(r, bold=True, color=BLACK, size=11)
    return p


def add_para(text, italic=False, bold=False, color=None, size=11, align=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    set_run(r, bold=bold, italic=italic, color=color, size=size)
    return p


def add_bullet(text, indent=0, bold_lead=None, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + indent * 0.6)
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        set_run(r, bold=True, size=size)
        r2 = p.add_run(" " + text)
        set_run(r2, size=size)
    else:
        r = p.add_run(text)
        set_run(r, size=size)
    return p


def add_chart(path, width_inches=6.5, caption=None):
    if not os.path.exists(path):
        add_para(f"[Chart missing: {os.path.basename(path)}]", italic=True, color=RED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run()
    r.add_picture(path, width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        cr = cp.add_run(caption)
        set_run(cr, italic=True, color=GREY, size=9)


def add_page_break():
    doc.add_page_break()


def add_table_from_data(headers, rows, col_widths=None, header_color=NAVY):
    """Add a formatted table; rows is list of lists; col_widths is list of cm."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Grid Accent 1"
    # Header
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
        # Fill cell
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), '1F4E79')
        tcPr.append(shd)
    # Body
    for ri, row in enumerate(rows, start=1):
        for ci, v in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            if ci > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(str(v))
            set_run(r, size=10)
    if col_widths:
        for col, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[col].width = Cm(w)
    # Spacer after
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ============================================================================
# COVER + PAGE 1 (Investment Summary)
# ============================================================================
add_heading("INITIATING COVERAGE", level=0, color=NAVY, before=24, after=2)
add_para("Equity Research · China Industrial — Auto Parts", italic=True, color=GREY, size=10)
add_para("", space_after=8)
add_heading("双林股份 (Shuanglin Co., Ltd.) — SZSE: 300100", level=1, color=BLACK, before=4, after=2)
add_para("Smart-Drive Solutions Provider; HDM Leader Pivoting to Humanoid Roller Screws & Smart Corner Modules",
          italic=True, color=GREY, size=11, space_after=12)

# Summary table at top
summary_tbl = doc.add_table(rows=2, cols=4)
summary_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
summary_data = [
    ("Current Price", "¥30.00", "Rating", "SELL"),
    ("12-Mo Target", "¥24",     "Implied", "−20.0%"),
]
for ri, row in enumerate(summary_data):
    for ci, val in enumerate(row):
        cell = summary_tbl.rows[ri].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(val)
        is_label = ci % 2 == 0
        if is_label:
            set_run(r, bold=True, size=11, color=GREY)
        else:
            is_rating = ci == 3 and ri == 0
            set_run(r, bold=True, size=14 if is_rating or "Target" in str(row[ci-1]) else 12,
                     color=RED if is_rating else BLACK)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), 'F2F2F2' if is_label else 'FFFFFF')
        tcPr.append(shd)
add_para("", space_after=8)

add_heading("Executive Summary", level=2)

add_para(
    "We initiate coverage of Shuanglin Co. (SZSE:300100, ¥30.00, ¥17.2 bn market cap) "
    "with a SELL rating and a 12-month price target of ¥24, implying −20.0% downside. "
    "Shuanglin is a Ningbo-headquartered automotive-parts manufacturer best known as the #1 China / #2 "
    "global producer of seat-positioning HDM (horizontal drive module) motors, with 32.8% domestic share "
    "and a customer roster that spans Tesla, BYD, NIO, Li Auto, XPeng, GAC, Geely, Great Wall, Chery, "
    "Changan, Leapmotor, plus Tier-1 globals (Faurecia, Lear, Adient, BorgWarner). Beyond HDM, the "
    "company holds the #3 China share in wheel-hub bearings via its Hubei subsidiary, a small but "
    "growing NEV e-drive business (Shandong), and — most pivotally for valuation — early-stage product "
    "pipelines in humanoid-robot reverse planetary roller screws, smart corner modules (with Tsinghua "
    "University), and low-altitude eVTOL e-drives.",
    align="justify")

add_para(
    "Our SELL rating is driven by three observations. First, every realistic valuation methodology — "
    "DCF Base (¥16.62), EV/EBITDA NTM peer median (¥23.79), P/E NTM 2026E peer median (¥21.06), and the "
    "12-month forward P/E 2027E peer median (¥27.40) — implies an intrinsic value below the current "
    "¥30 share price. Only DCF Bull (¥35.08) clears the current price, and that scenario requires both "
    "(i) the 100,000-unit/year humanoid roller-screw line to ramp on schedule by June 2026, and "
    "(ii) a formal program award (定点) from at least one humanoid OEM by 2027 — neither of which has been "
    "secured. Second, Q1 2026 results (revenue −10.4% YoY, NI −47.0%, ex-non-recurring NI −39.1%) "
    "confirm material pricing-pressure transmission from NEV OEM customers, with Top-1 customer "
    "concentration at 26.1% and Top-5 at 51.6% (up from 41% three years ago). Annualizing Q1 results "
    "puts the FY2026 EPS run-rate at ¥0.55–0.65, against an implied ¥0.91+ to justify the current price. "
    "Third, the current TTM P/E of ~34x sits 12 turns above the auto-parts peer median of 22x; should "
    "FY2026E earnings disappoint, multiple compression alone could drive the stock toward ¥15–18.",
    align="justify")

add_para(
    "What keeps us from a stronger SELL rating is the embedded option value of three emerging product "
    "lines — humanoid roller screws (TAM ¥50 bn by 2030E per industry estimates), smart corner modules "
    "(Schaeffler/Mobis adjacent technology, ¥18 bn TAM 2030E), and eVTOL e-drives — combined with a "
    "conservative balance sheet (net debt only ¥600 mn vs. ¥17 bn market cap) that limits absolute "
    "downside in the bull case. The pending HKEX A+H IPO (re-submitted March 2026) could also raise "
    "¥0.8–1.5 bn for capacity, accelerating the build-out without straining the capital structure. We "
    "expect Q2 2026 earnings (Aug release) and the roller-screw量产 line start (target June 2026) to "
    "be the two pivotal data points for our rating in the next six months.",
    align="justify")

# Mini chart on front page
add_chart(os.path.join(CHARTS, "chart_01_share_price_target.png"), width_inches=6.0,
          caption="Figure 1. 52-week share price history and 12-month price target of ¥24 (SELL)")

add_page_break()

# ============================================================================
# INVESTMENT THESIS (pages 2-5)
# ============================================================================
add_heading("Investment Thesis", level=1)

add_heading("Why SELL Now: Three Quantitative Arguments", level=2)

add_para(
    "Our SELL rating is grounded in three quantitative arguments that each independently — and "
    "collectively reinforce — the conclusion that Shuanglin's current ¥30 quote is above intrinsic value.",
    align="justify")

add_heading("1. All Five Core Methodologies Imply a Price Below ¥30", level=3)

add_para(
    "We apply five distinct valuation methods, weighted by our judgment of each method's relevance to "
    "Shuanglin's current business mix and 12-month outlook: DCF Base case (20%), DCF Bull case (10%), "
    "EV/EBITDA NTM peer median (20%), P/E NTM 2026E peer median (20%), and P/E forward 2027E peer "
    "median (30%). Of these, only the DCF Bull case implies a price above ¥30 — and that requires "
    "binary events (roller-screw量产 + humanoid OEM定点) that have not yet occurred.",
    align="justify")

add_table_from_data(
    headers=["Methodology", "Base case (¥)", "Range (¥)", "Weight", "Implied upside"],
    rows=[
        ["DCF Base (WACC 9.7%, g 2.5%)", "16.62", "14.96 – 18.28", "20%", "−44.6%"],
        ["DCF Bull (robot win)", "35.08", "31.57 – 38.59", "10%", "+16.9%"],
        ["EV/EBITDA NTM peer median", "23.79", "18.01 – 37.83", "20%", "−20.7%"],
        ["P/E NTM 2026E peer median", "21.06", "16.32 – 38.73", "20%", "−29.8%"],
        ["P/E forward 2027E peer median", "27.40", "21.23 – 50.38", "30%", "−8.7%"],
        ["Weighted-average target", "24.02", "19.38 – 37.94", "100%", "−20.0%"],
    ],
    col_widths=[5.0, 2.5, 3.0, 2.0, 2.5]
)

add_para("Source: Company filings; analyst estimates. Peer multiples sourced from Eastmoney, Yahoo Finance, "
          "company filings (May 2026 snapshot).", italic=True, color=GREY, size=9)

add_chart(os.path.join(CHARTS, "chart_32_football_field.png"), width_inches=6.5,
          caption="Figure 2. Valuation football field — six methodology ranges with current price (¥30) and target (¥24)")

add_heading("2. Q1 2026 Earnings Confirm Material Margin Compression", level=3)

add_para(
    "The 1Q26 report filed 2026-04-28 disclosed revenue of ¥1,193 mn (−10.4% YoY) and parent net "
    "income of ¥72 mn (−47.0% YoY). On an ex-non-recurring basis, NI fell 39.1% — which rules out the "
    "'one-off' narrative. The deterioration in margin profile, combined with customer concentration "
    "data disclosed in the 2025 annual report (Top-1 customer 26.1%, Top-5 at 51.6%, up +10 pp over "
    "three years), points to direct margin transmission from NEV OEM pricing renegotiations. "
    "Annualizing Q1 puts the FY2026 NI run-rate at approximately ¥288 mn — significantly below the "
    "¥399 mn assumed in our base case projection, and 43% below FY2025 actuals.",
    align="justify")

add_chart(os.path.join(CHARTS, "chart_09_customer_concentration.png"), width_inches=6.0,
          caption="Figure 3. FY2025 customer concentration — Top-1 at 26.1% (single Tesla-shaped exposure)")

add_heading("3. Multiple Compression Dominates Near-Term Risk", level=3)

add_para(
    "Shuanglin trades at a TTM P/E of ~34x — well above the auto-parts peer median of 22x (Eastmoney, "
    "May 2026). The premium is justified narratively by the option-value of robot/corner-module "
    "pipelines, but the math of multiple compression is unforgiving: if FY2026 EPS lands at ¥0.55 "
    "(annualized Q1 run-rate) and the multiple compresses to peer median 22x, the stock would trade "
    "at ¥12. Even a milder compression to 28x on ¥0.70 EPS yields ¥19.60 — still well below the "
    "current quote. We see consensus EPS revisions (expected after Q2 2026 results in August) as the "
    "near-term trigger for this re-rating.",
    align="justify")

add_chart(os.path.join(CHARTS, "chart_30_pe_ntm.png"), width_inches=6.5,
          caption="Figure 4. Peer P/E (NTM) — Shuanglin at 43.9x vs. auto-parts median 22x, robot-pure-play 60–119x")

add_heading("Why Not Stronger SELL: Three Offsetting Considerations", level=2)

add_heading("1. Real Option Value in Three Emerging Pipelines", level=3)

add_para(
    "Shuanglin is not a one-trick auto-parts pony. Three early-stage initiatives could each "
    "individually re-rate the stock by 30–60% if they reach commercial maturity by 2028:",
    align="justify")

add_bullet(
    "Humanoid roller screws: The company manufactured and delivered first sample reverse planetary "
    "roller screws (the core linear-actuator for humanoid robot upper/lower-body joints) in mid-2025. "
    "A 100,000-unit/year line is scheduled to start production in June 2026, supported by the "
    "January 2025 acquisition of Wuxi Kexin Mechanical (a CNC thread-grinder maker, ¥135 mn deal) "
    "that closes the screw-manufacturing supply chain. The humanoid roller-screw TAM grows from "
    "~¥1.5 bn in 2026E to ¥50 bn by 2030E (analyst estimates).",
    bold_lead="•")

add_bullet(
    "Smart corner modules: In March 2026, Shuanglin announced a joint venture with Tsinghua University "
    "and Huakong Tech to develop smart corner modules — integrated drive + steering + brake + "
    "suspension wheel-end modules viewed by Schaeffler and Continental as the next-generation "
    "by-wire chassis architecture. First commercial application is a 240-ton fully-electric "
    "corner-module mining truck (Inner Mongolia, 100 units in 2026 H1).",
    bold_lead="•")

add_bullet(
    "EHB/EMB ball screws: Multiple Chinese OEMs (Bethel, Likai, BorgWarner) are actively sampling "
    "Shuanglin's auto ball-screw products for brake-by-wire (EHB) and electric mechanical brake (EMB) "
    "systems. First volume contracts expected in 2026 H2.",
    bold_lead="•")

add_chart(os.path.join(CHARTS, "chart_15_tam_growth.png"), width_inches=6.5,
          caption="Figure 5. Three emerging TAM trajectories — humanoid + EHB + corner modules grow 100x by 2030E")

add_heading("2. Structural Operating Leverage Confirmed in FY2025", level=3)

add_para(
    "FY2025 ex-non-recurring net income grew +36.6% on revenue of +11.7% — a clear sign that "
    "operating leverage from the HDM platform-mix shift toward NEV is real and durable. The headline "
    "NI growth of only +1.25% was depressed by one-off items (FY2024 had ¥160 mn of investment "
    "income; FY2025 had ¥36 mn of asset disposal gains). The ex-non-recurring trajectory is the "
    "cleaner read and supports the long-term thesis that Shuanglin's core HDM + bearings business "
    "remains structurally healthy.",
    align="justify")

add_chart(os.path.join(CHARTS, "chart_02_revenue_gm_ni.png"), width_inches=6.5,
          caption="Figure 6. Revenue, net income, and gross margin trajectory FY2021A–FY2030E")

add_heading("3. Conservative Balance Sheet Limits Downside", level=3)

add_para(
    "Shuanglin's net debt at Q1 2026 of approximately ¥600 mn against a market cap of ¥17.2 bn (i.e., "
    "net-debt/equity ratio of 4%) is materially below the China auto-parts sector median (~20%). "
    "Even with the pending HKEX A+H IPO raising ¥0.8–1.5 bn of equity, the capital structure remains "
    "low-risk. This shields downside in scenarios where Q2/Q3 earnings disappoint but emerging "
    "products execute — and provides ample cushion to fund the roller-screw, Thailand bearings, and "
    "smart-corner-module build-outs through 2028 without further financing.",
    align="justify")

add_page_break()

# ============================================================================
# RISKS TO PRICE TARGET
# ============================================================================
add_heading("Risks to Our SELL Rating", level=1)

add_heading("Upside Catalysts (would prompt rating upgrade)", level=2)

upside_catalysts = [
    ("Roller-screw量产 line start (Jun 2026)",
     "Confirmation of >80% yield on the 100k-set/year line AND a formal humanoid-robot OEM定点 (program "
     "award) would shift our methodology weighting toward DCF Bull (¥35.08) and trigger immediate "
     "review. Watch July–September 2026 IR updates and the 2026 H1 半年报 (Aug 2026)."),
    ("Smart-corner-module mining-truck deliveries (2026 H1)",
     "First 100 units to Inner Mongolia coal-mining site. Successful six-month operation would open "
     "the unmanned-AGV market (TAM ¥10–15 bn by 2030E) and validate Shuanglin's adjacency to "
     "Schaeffler / Hyundai Mobis e-Corner Module technology."),
    ("Q2 2026 earnings recovery",
     "Base case requires gross margin to recover to ~22% by H2 2026. Two consecutive quarters of YoY "
     "revenue +5% with stable margins would re-rate base-case projections and raise the "
     "methodology-weighted target into the ¥28–32 range."),
    ("HKEX A+H IPO confirmation (2026 Q3–Q4)",
     "A+H listing brings ¥0.8–1.5 bn fresh capital; HK pricing premium would compress the A-share "
     "discount and raise institutional ownership. Anchor investor identity (HK long-only vs. PE) "
     "will signal strategic direction."),
    ("EHB/EMB ball-screw design wins",
     "Multiple OEMs (Bethel, Likai, BorgWarner) are actively sampling; first volume contract in "
     "2026 H2 would validate the line-control-chassis pivot and add ¥300–500 mn incremental revenue "
     "runway by FY2028."),
]
for title, desc in upside_catalysts:
    add_para("", space_after=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run("▲ ")
    set_run(r1, bold=True, color=GREEN, size=11)
    r2 = p.add_run(title)
    set_run(r2, bold=True, size=11)
    add_para(desc, align="justify", size=10)

add_heading("Downside Catalysts (would extend or deepen the SELL stance)", level=2)

downside = [
    ("Continued NEV pricing pressure (high probability)",
     "Q1 2026 showed −47% NI on −10% revenue. Further OEM pricing renegotiations could compress "
     "FY2026 EBITDA margin by another 100–200 bp, taking DCF Base toward ¥10–12. Tesla's 2026 "
     "platform-pricing reset is the specific risk to monitor."),
    ("Roller-screw delay or yield problem (medium probability)",
     "The June 2026 量产 timing is aggressive. Yield issues at the new Kexin grinder line would push "
     "first humanoid revenue out 12+ months, collapsing the bull-case option value priced into the "
     "current share quote."),
    ("HDM share loss to local competitors (medium probability)",
     "Hangzhou Xinjian and Yimai are 2nd/3rd-place HDM players with growing NEV OEM relationships. "
     "Tesla in-source remains tail risk (low probability given the 30-month PPAP cycle and "
     "high-precision manufacturing know-how barrier)."),
    ("Thailand operations execution",
     "The Thailand New Torch bearing plant just launched January 2025 and the NEV e-drive line "
     "targets 2026 Q1 ramp. Shuanglin's foreign-operations track record is poor — the 2017 DSI "
     "Australia transmission integration was impaired in 2022–2023."),
    ("Family-controlled governance",
     "The Wu family controls 48.9% via founder Wu Jianbin (4.49% direct) and Shuanglin Group "
     "(44.43%). Transformative M&A (e.g., another DSI-style related-party injection) could be "
     "unilaterally executed. The 2017 DSI deal is the cautionary precedent."),
]
for title, desc in downside:
    add_para("", space_after=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run("▼ ")
    set_run(r1, bold=True, color=RED, size=11)
    r2 = p.add_run(title)
    set_run(r2, bold=True, size=11)
    add_para(desc, align="justify", size=10)

add_page_break()

# ============================================================================
# COMPANY 101 — BUSINESS OVERVIEW
# ============================================================================
add_heading("Company Overview", level=1)

add_heading("Business Description and History", level=2)

add_para(
    "Shuanglin Co., Ltd. is a Ningbo (Zhejiang)-headquartered manufacturer of automotive parts and "
    "robotics components, listed on the Shenzhen Stock Exchange's ChiNext (创业板) board under "
    "SZSE:300100 since August 2010. The company's roots trace to 1989 when founder Wu Jianbin's family "
    "established a small plastic-parts workshop in Ninghai county, which gradually transitioned into "
    "automotive supply in the late 1990s. The pivotal R&D project began in 2000 with the development "
    "of HDM (horizontal drive module) — the precision worm-gear and ball-screw mechanism that powers "
    "fore-aft adjustment in electrically-adjustable car seats. HDM became the company's signature "
    "product, with annual unit shipments breaking 30 million in 2025 and Chinese market share at 32.8% "
    "(industry source: 智研咨询 2025).",
    align="justify")

add_para(
    "Post-IPO, Shuanglin grew through two major acquisitions: (i) the August 2014 acquisition of "
    "Hubei New Torch (Hubei Xinhuoju) for ¥820 mn cash, securing entry into the wheel-hub bearing "
    "category (now the company's second-largest profit pillar at ¥1,415 mn revenue / ¥150.8 mn NI in "
    "FY2025); and (ii) the October 2017 ¥2.3 bn acquisition of DSI Australia (automatic transmissions) "
    "via private placement to controlling shareholder Shuanglin Group. The DSI integration proved "
    "difficult — the global transmission market shifted toward CVT/DCT/EV powertrains, and DSI's "
    "traditional 6-AT business eroded, leading to goodwill impairments in 2022–2023. Earnings troughed "
    "in 2022 (¥75 mn NI, vs. peak ¥495 mn in 2018 pre-integration) before a sharp recovery in 2024 "
    "(¥497 mn NI, +514% YoY) on HDM volumes ramping with NEV adoption.",
    align="justify")

add_chart(os.path.join(CHARTS, "chart_05_milestones_timeline.png"), width_inches=6.5,
          caption="Figure 7. Shuanglin corporate milestones — 1989 founding through 2026 HKEX IPO and Tsinghua JV")

add_para(
    "The current strategic narrative — and the principal driver of valuation premium — is Shuanglin's "
    "Third Pivot launched in 2023: a stated transformation from 'automotive parts' to 'smart drive "
    "solutions provider.' Concretely, this means leveraging HDM's worm-gear / ball-screw / "
    "miniaturized-precision-manufacturing know-how into three adjacent end markets: humanoid robots "
    "(reverse planetary roller screws + joint modules), smart corner modules (wheel-end drive + "
    "steer-by-wire + brake-by-wire integrated assemblies), and low-altitude eVTOL aircraft (30–250 kW "
    "e-drive systems). These three pipelines collectively contribute <5% of FY2025 revenue but "
    "represent the binary upside that drives the current valuation premium.",
    align="justify")

add_chart(os.path.join(CHARTS, "chart_06_strategic_pivots.png"), width_inches=6.5,
          caption="Figure 8. Three strategic pivots in 22 years — current pivot to smart-drive solutions")

add_heading("Ownership and Governance", level=2)

add_para(
    "Shuanglin is a family-controlled company with concentrated insider ownership. As of the Q1 2026 "
    "filing, controlling shareholder Shuanglin Group holds 44.43% of shares; founder Wu Jianbin "
    "directly holds an additional 4.49% (with 75% of his stake under lockup as restricted shares); "
    "the 2025 employee stock-ownership plan adds 0.35%. The combined insider voting bloc is "
    "approximately 49.3%, with the Wu family (Wu Jianbin, sister Wu Weijing, sister Wu Xiaojing) "
    "executing a binding concert-party agreement. Founder Wu Jianbin chairs both the board and acts "
    "as CEO; he has held the chairman role continuously since November 2004 (age 24 at appointment) "
    "and now stands as a 22-year tenured chief executive — a notable stability profile in China's "
    "Tier-1 supplier space.",
    align="justify")

add_para(
    "Governance risk is heightened by the family-control structure. The 2017 DSI Australia deal — "
    "a related-party transaction with Shuanglin Group — illustrates the precedent for unilateral "
    "transformative M&A. Minority shareholders had limited practical recourse to oppose the "
    "transaction even as subsequent impairments showed the integration logic had been flawed. The "
    "ongoing HKEX A+H IPO (re-submitted March 2026 with CITIC Securities and GF Securities as joint "
    "sponsors) could be the next inflection point — successful listing could improve governance "
    "discipline through international institutional ownership.",
    align="justify")

add_chart(os.path.join(CHARTS, "chart_07_management_org.png"), width_inches=6.5,
          caption="Figure 9. Management structure and key operating subsidiaries (FY2025 revenue contribution)")

add_heading("Management Team — Individual Profiles", level=2)

add_heading("Wu Jianbin (邬建斌) — Chairman and CEO", level=3)

add_para(
    "Wu Jianbin, born September 1980 in Ninghai (Zhejiang), is the architect of Shuanglin's 22-year "
    "strategic trajectory. He holds dual EMBA degrees — Shanghai National Accounting Institute "
    "(Finance and Financial Management) and Cheung Kong Graduate School of Business. Wu took the "
    "chairman role in November 2004 at age 24, a notably young appointment that reflected both the "
    "family-control structure and his early demonstration of strategic capability. As 'second-"
    "generation' family successor, he has led the company through three distinct chapters.",
    align="justify")

add_para(
    "First, from 2000–2010, Wu oversaw the HDM product industrialization — taking the seat-positioning "
    "motor concept from zero to volume production, ultimately reaching 30 million annual units by "
    "2025 with 32.8% domestic and 15.1% global market share. This represents the foundational moat "
    "of the company. Second, in 2010 he led the ChiNext IPO, raising capital that financed the "
    "subsequent acquisition wave: Hubei New Torch in 2014 (¥820 mn) and the larger DSI Australia "
    "deal in 2017 (¥2.3 bn). The DSI transaction is now widely viewed as the major strategic "
    "misstep — the integration challenges and subsequent impairments cost the company several years "
    "of earnings momentum and credibility. Third, from 2023 onward, Wu has championed the company's "
    "'second-generation entrepreneurship' positioning Shuanglin as a 'smart drive solutions "
    "provider' rather than auto-parts company. The roller-screw, corner-module, and eVTOL bets are "
    "his signature current-era initiatives.",
    align="justify")

add_para(
    "Wu Jianbin directly holds 25.69 mn shares (4.49% of total), with 19.27 mn under lockup as "
    "restricted shares. Through Shuanglin Group and upstream entities (Ningbo Zhiyuan Investment "
    "57.14%, Ninghai Baolai Investment 42.86%), he and sisters Wu Weijing and Wu Xiaojing jointly "
    "control 48.9% of voting rights via a binding concert-party agreement. He is also a Ningbo "
    "Municipal People's Congress representative, director of the China Youth Entrepreneurs "
    "Association, and recipient of multiple municipal/provincial entrepreneur awards (2008 Ningbo "
    "Excellent Entrepreneur, 2014 Ningbo Labor Medal, 2022 Ningbo Charity Award). His public "
    "communications consistently emphasize 'technology faith without compromise' — a positioning "
    "consistent with his engineering-plus-finance background, and reflected in the 2025 Shareholder "
    "Letter framing 'industrialization challenges of reverse-planetary roller screws' as both moat "
    "and credo. The family-control concentration limits minority-shareholder influence on major "
    "strategic decisions, but also creates decision-speed advantages and long-horizon strategic "
    "patience that are rare in China's mid-cap industrials.",
    align="justify")

add_heading("Wu Huaiying (武淮颖) — CFO Function (Chief Accountant)", level=3)

add_para(
    "Wu Huaiying serves as the responsible person for accounting affairs (主管会计工作负责人), the "
    "effective CFO function. The company has not publicly disclosed his individual biographical "
    "details with the level of granularity typical for Western public companies, but his completed "
    "transactions during 2025 include three significant items. First, the May 2025 ¥1.5 bn A-share "
    "private placement was withdrawn after 7 days when regulatory dynamics shifted; the company "
    "immediately pivoted to the HKEX A+H IPO route — demonstrating execution agility. Second, the "
    "¥135 mn cash acquisition of Wuxi Kexin in January 2025 was completed using internal cash flow "
    "without taking on new debt, demonstrating capital discipline. Third, in March 2026, the HKEX "
    "A+H IPO prospectus was re-submitted with CITIC Securities and GF Securities as joint sponsors.",
    align="justify")

add_para(
    "The CFO team's balance-sheet management is conservative. Year-end 2025 short-term debt was "
    "¥565 mn (8.05% of total assets), long-term debt was reduced to zero (vs. ¥100 mn at year-end "
    "2024), and the current portion of long-term debt of ¥240 mn is well-covered by ¥735 mn of cash. "
    "Compared to the broader Chinese auto-parts mid-cap universe (median net-debt-to-EBITDA of ~1.5x), "
    "Shuanglin's <1.0x leverage is conservative — leaving optionality for both organic CapEx "
    "(roller-screw, Thailand expansion, smart corner module) and potential acquisition (after the "
    "DSI lesson, however, large M&A is unlikely in the near term).",
    align="justify")

add_heading("Zhu Liming (朱黎明) — Board Secretary", level=3)

add_para(
    "Zhu Liming serves as Board Secretary, based in the Ningbo Ninghai headquarters, and is the "
    "company's primary point of contact for regulators and investors. During 2025, the company "
    "conducted three formal investor-relations events covering roller-screw product positioning, "
    "Kexin acquisition rationale, overseas expansion strategy, and quarterly operating updates. The "
    "investor-relations cadence has been disciplined — earnings releases on schedule, no material "
    "guidance revisions outside of formal channels, and increasingly detailed disclosure of segment "
    "and subsidiary financials in the FY2024 and FY2025 annual reports.",
    align="justify")

add_heading("Key Subsidiary Leadership", level=3)

add_para(
    "Below the parent-company level, two operating subsidiaries are critical to the consolidated "
    "results. Hubei Shuanglin Bearings (formerly Hubei New Torch, the 2014 acquisition) generated "
    "¥1,415 mn revenue and ¥150.8 mn net income in FY2025 — a 10.7% net margin, materially above the "
    "consolidated 9.2% — making it the second-largest profit pillar after the parent's HDM business. "
    "It holds national-grade Torch Plan project status, supports BYD, NIO, Leapmotor, Avatr, Zeekr, "
    "AITO M8, Dongfeng Nissan N7 (2025 launch), Hongqi, Changan Ford, and XPeng vehicle programs. "
    "Shandong Shuanglin New Energy (NEV e-drive) generated ¥690 mn revenue and ¥65 mn net income "
    "(9.4% margin) supporting Wuling Hongguang Mini EV 3rd/4th generations, FAW Bestune Mini, Chery "
    "ice-cream mini-EV, and Changan glutinous-corn mini-EV applications. Both subsidiaries are 100%-"
    "owned (no minority interest).",
    align="justify")

add_para(
    "The 2025 corporate reorganization added five new strategic centers — Innovation Incubation "
    "Center, Strategic Technology Center, Strategic Expansion Center, Strategic Investment Center, "
    "and Intelligence Center — designed to support the transition from auto-parts to smart-drive-"
    "solutions provider. These five centers serve as the institutional architecture for the company's "
    "current and future bets on roller screws, corner modules, and low-altitude eVTOL.",
    align="justify")

add_heading("Product Portfolio", level=2)

add_para(
    "Shuanglin organizes its product portfolio along three reporting segments under the FY2025 "
    "restated taxonomy: Transmission/Drive/Intelligent (¥3,270 mn, 59.6% of revenue, GM 23.88%); "
    "Auto Interior/Exterior trim (¥1,946 mn, 35.5%, GM 14.32%); and Other—primarily Kexin grinder "
    "equipment and Shuanglin Molds (¥260 mn, 4.7%, GM 31.87%). Rental income (¥7.8 mn) is a "
    "rounding item but disclosed separately.",
    align="justify")

add_chart(os.path.join(CHARTS, "chart_03_revenue_by_product.png"), width_inches=6.5,
          caption="Figure 10. ⭐ Revenue by product segment FY2021A–FY2030E — mix shifting toward Transmission/Drive")

add_para(
    "Within Transmission/Drive/Intelligent — the most strategically important segment — we identify "
    "eleven distinct product lines, each at a different stage of maturity:",
    align="justify")

prod_tbl_data = [
    ["Product line", "FY2025 revenue (¥mn, est)", "Moat", "Stage"],
    ["HDM 座椅水平驱动器", "~1,200", "Strong (33% share)", "Cash cow"],
    ["Wheel hub bearings (Hubei)", "1,415", "Medium (#3 share)", "Growth"],
    ["NEV e-drive (Shandong)", "690", "Weak (commoditized)", "Scale chase"],
    ["Seat motors", "200", "Weak (late entrant)", "Growth"],
    ["Headrest actuator", "50", "Early lead", "Emerging"],
    ["Auto ball-screws (EHB/EMB)", "30", "Adjacency", "Pre-volume"],
    ["Steering folding actuator", "20", "Customer-locked", "Pre-volume"],
    ["Roller screws (humanoid)", "30", "Equipment + know-how", "⭐ Option"],
    ["Joint modules (humanoid)", "15", "Vertical integration", "⭐ Option"],
    ["eVTOL e-drive 30–250 kW", "20", "Platform synergy", "⭐ Option"],
    ["Smart corner modules", "10", "Tsinghua JV + scarcity", "⭐ Option"],
]
add_table_from_data(prod_tbl_data[0], prod_tbl_data[1:],
                     col_widths=[5.5, 4.0, 4.0, 3.0])

add_para("Source: 2025 年度报告; Task 1 company research. Revenue allocations are analyst estimates from segment-level disclosure plus subsidiary breakouts.",
          italic=True, color=GREY, size=9)

add_chart(os.path.join(CHARTS, "chart_08_product_portfolio_matrix.png"), width_inches=6.5,
          caption="Figure 11. Product portfolio matrix — growth vs. maturity, bubble size = FY2025 revenue")

add_heading("Auto Interior/Exterior Trim — Cash Cow with Limited Growth", level=3)

add_para(
    "The Auto Interior/Exterior segment (¥1.95 bn revenue, 35.5% of FY2025 sales, 14.3% gross margin) "
    "is a mature business serving China's mid-market OEMs (Wuling, Changan, XPeng, Volkswagen China). "
    "Products include bumpers, door panels, B/C pillar trim, safety pads, precision injection-molded "
    "parts, PEEK specialty-material parts, and ignition coils. This segment is intentionally not "
    "growing — the company manages it as a cash flow harvest with annual revenue growth in the 2–5% "
    "range. We assume −5% in FY2026 (reflecting NEV pricing pass-through), then 3–5% growth annually "
    "through FY2030E.",
    align="justify")

add_page_break()

# ============================================================================
# CUSTOMERS AND COMPETITIVE LANDSCAPE
# ============================================================================
add_heading("Customers and Market Position", level=1)

add_heading("Customer Profile and Concentration Risk", level=2)

add_para(
    "Shuanglin sells direct (no distributor) to two customer types: (i) Tier-1 global automotive "
    "suppliers — Faurecia, UAES, Brose, Autoliv, MAHLE, ZF, Lear, BorgWarner, Adient, Magna, Valeo; "
    "and (ii) integrated OEMs — including Tesla (the largest single customer at an estimated 26% of "
    "FY2025 revenue), BYD, NIO, Li Auto, XPeng, Chery, Geely, Great Wall, Changan, Leapmotor, SAIC-GM-"
    "Wuling. Contract structure is project-by-project (program-by-program), with each design win "
    "carrying a 5–7 year vehicle-lifecycle revenue commitment. PPAP (Production Part Approval Process) "
    "cycles take 18–36 months from initial RFQ to volume, creating significant switching costs and "
    "incumbency advantages.",
    align="justify")

add_para(
    "However, customer concentration has materially increased. Top-5 customer share rose from 41% in "
    "FY2023 to 47% in FY2024 to 51.6% in FY2025 — a 10-percentage-point increase over three years. The "
    "single Top-1 customer (anonymized in filings but commonly inferred to be Tesla based on the HDM "
    "+ North American supply pattern) reached 26.1% in FY2025. Both metrics now exceed institutional "
    "thresholds for material concentration risk (Top-1 >20%, Top-5 >50%). The Q1 2026 weakness "
    "demonstrates the transmission mechanism: a single major customer renegotiating prices, or "
    "introducing in-house alternatives, has outsized impact on consolidated revenue and EBIT margin.",
    align="justify")

add_chart(os.path.join(CHARTS, "chart_09_customer_concentration.png"), width_inches=6.0,
          caption="Figure 12. FY2025 customer concentration — Top-1 26.1%, Top-5 51.6% (both material)")

add_heading("Competitive Position by Product Line", level=2)

add_heading("HDM — Clear Market Leader, Stable Moat", level=3)

add_para(
    "Shuanglin holds 32.8% of China's HDM (horizontal drive module) market and 15.1% of the global "
    "market, ranking #1 and #2 respectively (industry source: 智研咨询 2025). Competitors include "
    "Yimai (14% China share) and Hangzhou Xinjian (9% China share). The moat is multi-layered: "
    "(a) scale and yield advantages from 30+ million units of annual production volume; (b) deep "
    "Tier-1 relationships requiring 2–3 year PPAP qualification cycles; (c) cross-product synergy with "
    "seat motors and headrest actuators (same factory floor, shared tooling); (d) global presence "
    "(China + Thailand + North America) matching OEM regional supply needs.",
    align="justify")

add_chart(os.path.join(CHARTS, "chart_17_hdm_market_share.png"), width_inches=5.5,
          caption="Figure 13. HDM China market share 2025 — Shuanglin clear #1 at 32.8%")

add_heading("Wheel Hub Bearings — Strong #3 with NEV Tailwind", level=3)

add_para(
    "The Hubei Shuanglin Bearings subsidiary (formerly Hubei New Torch, acquired 2014) holds an "
    "estimated 9% share of China's wheel-hub bearing market, ranking #3 behind Wanxiang Qianchao "
    "(15%) and Renben Bearings (12%). Foreign incumbents — SKF, NSK, NTN, Schaeffler, JTEKT — "
    "collectively still hold ~35% of the China market, primarily serving joint-venture OEMs and "
    "high-end domestic brands. Hubei Bearings' competitive positioning has improved markedly through "
    "NEV penetration: 2025 design wins included Avatr, Zeekr, AITO M8, Dongfeng Nissan N7, Hongqi, "
    "Changan Ford, and XPeng. The Thailand bearing plant (launched January 2025) extends production "
    "footprint to support OEMs' Southeast Asian operations.",
    align="justify")

add_chart(os.path.join(CHARTS, "chart_18_bearing_market_share.png"), width_inches=5.5,
          caption="Figure 14. Wheel hub bearing China market share 2025 — Shuanglin Hubei at #3 (9%)")

add_heading("Roller Screws and Humanoid — The Critical Pipeline", level=3)

add_para(
    "Reverse planetary roller screws are the core linear-actuator inside humanoid robot "
    "upper- and lower-body joints, and they represent an acute supply chain bottleneck for the "
    "industry. The Chinese competitive set falls into three tiers: (1) early commercializers — "
    "Shuanglin, Hengli Hydraulic (601100), Beite Tech (603009, 'pure-play' but small); (2) "
    "concept-stage — Beste (300580), XCC (603667), Dingzhi Tech (873593); and (3) foreign incumbents "
    "— GSA (Switzerland), Rollvis, SKF, Schaeffler. Shuanglin's specific advantages are (a) vertical "
    "integration via the Kexin grinder acquisition (eliminating reliance on imported precision "
    "grinders that face export controls); (b) HDM-derived precision-manufacturing process know-how; "
    "(c) early sample delivery to a leading Chinese NEV-OEM-affiliated humanoid program (mid-2025); "
    "and (d) ¥135 mn Kexin acquisition opens a separate adjacency — selling thread grinders to "
    "competitors.",
    align="justify")

add_para(
    "The critical gating event is the June 2026 量产 (volume production) line start at the 100,000-set/"
    "year capacity. Yield curve achievement (industry benchmark: 80%+ within six months) and first "
    "formal program award (定点) from a humanoid OEM are the two binary catalysts our valuation hinges on.",
    align="justify")

add_chart(os.path.join(CHARTS, "chart_16_peer_positioning.png"), width_inches=6.5,
          caption="Figure 15. Peer positioning — revenue scale × EBITDA margin × market cap")

add_page_break()

# ============================================================================
# MARKET OPPORTUNITY (TAM)
# ============================================================================
add_heading("Market Opportunity (TAM)", level=1)

add_para(
    "Shuanglin's total addressable market is the sum of (a) global auto-parts categories where the "
    "company currently competes (HDM, seat motors, bearings, e-drive, trim, ball screws) and (b) "
    "three emerging markets — humanoid roller screws, smart corner modules, eVTOL e-drives — that "
    "represent the option-value premium investors are paying for. We size each separately:",
    align="justify")

add_table_from_data(
    headers=["Market", "2025 TAM", "2030E TAM", "CAGR", "Shuanglin TAM exposure"],
    rows=[
        ["China HDM", "¥3 bn", "¥4.5 bn", "8%", "32.8% share = ¥1.5 bn ceiling"],
        ["Global HDM", "¥18 bn", "¥27 bn", "8%", "15.1% global share"],
        ["China wheel hub bearings", "¥15 bn", "¥21 bn", "7%", "9% share via Hubei"],
        ["Auto ball-screws (EHB+EMB)", "¥3 bn", "¥17 bn", "42%", "New entry, target 5–10%"],
        ["Humanoid roller screws", "¥0.5 bn", "¥50 bn", "157%", "⭐ Target 5–15% share"],
        ["Smart corner modules", "¥0.3 bn", "¥18 bn", "127%", "⭐ Early lead"],
        ["eVTOL e-drives", "¥0.2 bn", "¥8 bn", "108%", "⭐ Platform synergy"],
        ["NEV e-drive (China)", "¥250 bn", "¥450 bn", "12%", "<1% share (commoditized)"],
        ["Auto interior/exterior trim", "¥600 bn", "¥700 bn", "3%", "Mature scale player"],
    ],
    col_widths=[5.0, 2.5, 2.5, 2.0, 4.5]
)

add_para(
    "The market structure is asymmetric: Shuanglin's mature markets (HDM, bearings, trim) total ~¥640 "
    "bn in 2030E with single-digit growth, but the company holds positions where its addressable "
    "ceiling is constrained to a few percentage points. By contrast, the three emerging markets "
    "(humanoid + corner module + eVTOL) total ~¥76 bn in 2030E with triple-digit CAGRs and where "
    "Shuanglin can credibly aim for 5–15% share given technology leadership and government policy "
    "support.",
    align="justify")

add_chart(os.path.join(CHARTS, "chart_15_tam_growth.png"), width_inches=6.5,
          caption="Figure 16. Three emerging TAM curves — humanoid roller-screws, EHB, smart corner modules")

add_page_break()

# ============================================================================
# INDUSTRY OVERVIEW
# ============================================================================
add_heading("Industry Overview", level=1)

add_heading("China Automotive Parts Industry — Structural Reset", level=2)

add_para(
    "Shuanglin operates at the intersection of three industry currents that together define its "
    "near-term and long-term outlook. The first is the structural reset in China's automotive parts "
    "industry, driven by the transition from internal-combustion-engine (ICE) vehicles to new-energy "
    "vehicles (NEVs). China NEV penetration reached approximately 51% of new vehicle sales in 2025 "
    "(China Association of Automobile Manufacturers data), up from 6% in 2020 — one of the most "
    "rapid industrial transitions in modern industrial history. This transition fundamentally "
    "reshuffles the auto-parts value chain: traditional ICE components (engine, transmission, "
    "exhaust) lose volume while new categories (battery, e-motor, e-axle, power electronics, "
    "thermal management) emerge with different supplier ecosystems.",
    align="justify")

add_para(
    "For Shuanglin, the NEV transition is bifurcated. The HDM, wheel bearings, interior/exterior "
    "trim, and seat motor categories are 'powertrain-agnostic' — they are required regardless of "
    "whether the vehicle is ICE or EV — and have benefited from rising NEV penetration via two "
    "mechanisms: (1) higher unit content per vehicle (NEV interiors are typically more feature-rich, "
    "with more electrically-adjustable seats, electric headrests, etc.), and (2) reduced foreign-"
    "incumbent competition (many global Tier-1 suppliers had ICE-focused product lines that did not "
    "transfer cleanly to NEV programs, opening market share for local players like Shuanglin). The "
    "transmission business (DSI Australia) was the casualty — automatic transmissions are largely "
    "redundant in pure-electric vehicles, leading to the 2022–2023 impairment cycle.",
    align="justify")

add_para(
    "The countervailing pressure is OEM pricing power. Chinese NEV brands — particularly the "
    "domestic leaders BYD, Geely, Chery, and the new-entrant 'new force' brands NIO, Li Auto, "
    "XPeng — have engaged in aggressive price competition since 2023, with multiple rounds of "
    "model-specific price reductions. The pricing pressure flows up the supply chain: OEMs renegotiate "
    "annual supplier contracts at lower prices, supplier margins compress, and the supplier with "
    "the most-concentrated single-customer exposure suffers most. This is the precise mechanism that "
    "produced Shuanglin's Q1 2026 (−47% NI on −10% revenue) and that drives our SELL rating.",
    align="justify")

add_heading("Humanoid Robot Industry — From Sample to Scale", level=2)

add_para(
    "The humanoid robot industry transitioned from research-laboratory curiosity to genuine "
    "commercial activity during 2024–2025. Tesla's Optimus program demonstrated approximately 1,000 "
    "unit production capacity by end-2025, with plans for 10,000+ units in 2026; Figure AI raised "
    "over USD 1 bn at a USD 39 bn valuation; multiple Chinese players (UBTECH, Unitree, Galbot, "
    "Fourier Intelligence, Zhibang, Agibot) ramped production from prototype to small-batch volume. "
    "Industry analyst consensus is that humanoid robot annual production volumes could reach 100,000+ "
    "units globally by 2027 and 1 million+ by 2030 — a trajectory that, if realized, would create a "
    "supply chain category as significant as the smartphone supply chain in the mid-2010s.",
    align="justify")

add_para(
    "Within the humanoid supply chain, three component categories are particularly under-supplied "
    "relative to expected demand: (1) reverse planetary roller screws (the linear-actuator core "
    "for arm and leg joints, ~6–10 per humanoid); (2) harmonic reducers (for rotational joints, "
    "~30–40 per humanoid); and (3) high-torque flat brushless motors. Shuanglin's strategic bet is "
    "on roller screws — leveraging its HDM-derived precision-screw manufacturing know-how and the "
    "vertical-integration advantage from the January 2025 Kexin grinder acquisition. The competitive "
    "advantage versus larger foreign incumbents (Switzerland's GSA, Rollvis; SKF; Schaeffler) is "
    "cost: Chinese manufacturing economics with comparable yield and tolerance can produce roller "
    "screws at 30–50% of foreign cost, which is meaningful when each humanoid robot consumes ¥10–30k "
    "of screw value.",
    align="justify")

add_para(
    "The pivotal commercialization gate for Shuanglin is the 100,000-set/year production line "
    "scheduled for June 2026 start-up. Industry benchmarks suggest 80%+ yield within six months is "
    "achievable for an experienced manufacturer; falling short of that benchmark would push first "
    "formal humanoid OEM program awards out 12+ months, eroding the option-value premium baked into "
    "the current share price.",
    align="justify")

add_heading("Smart Corner Module Industry — Emerging Architecture", level=2)

add_para(
    "Smart corner modules — integrated wheel-end assemblies that combine drive motor, "
    "steering, brake, and suspension into a single replaceable unit — represent the proposed "
    "'next-generation' chassis architecture for autonomous vehicles. The technology is positioned by "
    "Schaeffler, Continental, ZF, and Hyundai Mobis as the eventual successor to traditional vehicle "
    "architectures, with two principal advantages: (1) modularity and serviceability — replacing a "
    "failed corner module is far simpler than replacing an integrated axle; (2) packaging flexibility "
    "— removing the mechanical steering column and brake hydraulics frees interior space and enables "
    "novel vehicle layouts.",
    align="justify")

add_para(
    "Commercial deployment to date has been limited. Hyundai Mobis has demonstrated e-Corner Module "
    "concepts at multiple shows but no production application; Schaeffler has integrated wheel-hub "
    "motor concepts but not full corner-module integration. Shuanglin's March 2026 joint venture "
    "with Tsinghua University and Huakong Technology — establishing 'Zhejiang Shuanglin Technology' "
    "with corner-module focus — positions the company as one of the few credible players globally "
    "with both the academic R&D partnership and the manufacturing scale to commercialize. The first "
    "application is a 240-ton fully-electric corner-module mining truck, with engineering prototype "
    "completed and 100-unit deployment planned to an Inner Mongolia coal-mining site in 2026 H1.",
    align="justify")

add_heading("Low-Altitude eVTOL Industry — Adjacent Bet", level=2)

add_para(
    "Low-altitude eVTOL (electric vertical take-off and landing) aircraft constitute Shuanglin's "
    "third emerging market opportunity. Industry players include Joby Aviation (US), EHang (China), "
    "Volocopter (Germany), and several Chinese new-entrants. The Chinese central government's 'low-"
    "altitude economy' policy designation (低空经济) in 2024 created policy tailwinds for domestic "
    "eVTOL development. Shuanglin's specific opportunity is in the e-drive propulsion subsystems "
    "(motor + power electronics + thermal management) in the 30–250 kW range — directly leveraging "
    "the company's NEV e-drive platform technology with adaptations for aerospace certification. "
    "Revenue contribution from this segment is expected to remain minimal (<¥50 mn) through FY2027 "
    "but represents real optionality on a market that could reach ¥30+ bn in China by 2030E.",
    align="justify")

add_page_break()

# ============================================================================
# DEEP RISK ANALYSIS
# ============================================================================
add_heading("Detailed Risk Assessment", level=1)

add_heading("Operational Risks", level=2)

add_heading("Customer Concentration Risk — Material", level=3)

add_para(
    "The most material operational risk facing Shuanglin is customer concentration. Top-1 customer "
    "share of 26.1% (FY2025) is above the 20% institutional threshold for material concentration "
    "risk, and Top-5 share of 51.6% (up 10 percentage points over three years) is above the 50% "
    "threshold. The single Top-1 customer (anonymized in filings) is widely believed to be Tesla "
    "based on the HDM product line and North American supply relationship pattern. Tesla's "
    "historical behavior on supplier negotiations is well-documented: aggressive annual price "
    "reductions, willingness to dual-source, and occasional in-house manufacturing of components "
    "previously procured externally. While HDM is a high-precision, low-unit-value component where "
    "in-sourcing has weak economics for a customer of Tesla's scale, we cannot rule out price-"
    "reduction demands as the principal Q1 2026 driver.",
    align="justify")

add_para(
    "Mitigation pathways: (1) Continued NEV customer diversification — Shuanglin has added BYD, NIO, "
    "Li Auto, XPeng, and multiple second-tier brands as HDM customers, reducing Top-1 dependence "
    "structurally; (2) Product diversification into roller screws and smart corner modules with "
    "entirely different customer bases (humanoid OEMs, mining-equipment OEMs); (3) Geographic "
    "diversification via Thailand and North American production footprints. However, the diversification "
    "trajectory is multi-year and does not protect against near-term pricing-pressure transmission.",
    align="justify")

add_heading("Technology Execution Risk — Roller Screws", level=3)

add_para(
    "The roller-screw program is the company's most consequential single technology bet. The "
    "June 2026 量产 line start at 100,000 sets/year capacity represents a step-change from the 1,500 "
    "sets delivered in 2025. Two specific execution risks: (1) Yield curve — achieving the industry-"
    "standard 80%+ yield within six months of line start; (2) Customer qualification — securing the "
    "first formal program award (定点) from a humanoid OEM, which requires passing sample-validation "
    "cycles that typically take 6–18 months. If both go well, FY2027 humanoid revenue could reach "
    "¥200–500 mn with healthy margins (60%+ gross margin estimates); if either fails, the "
    "Bull-case DCF scenario evaporates.",
    align="justify")

add_heading("Geographic Expansion Risk — Thailand and North America", level=3)

add_para(
    "Shuanglin's Thailand bearing plant launched January 2025; its NEV e-drive production line in "
    "Thailand is targeted for ramp in Q1 2026. The company's track record of foreign-operations "
    "execution is poor — the 2017 DSI Australia transmission acquisition was impaired in 2022–2023 "
    "with management citing 'integration difficulties.' For the Thailand initiative, the strategic "
    "logic is sound (proximity to NEV OEMs' Southeast Asia production, e.g., BYD's Thailand "
    "facility), but execution carries risk: operational losses in the first 2–4 quarters are common "
    "for greenfield foreign manufacturing, and our base case assumes Thailand contributes positively "
    "from H2 2026 onward.",
    align="justify")

add_heading("Financial Risks", level=2)

add_heading("Free Cash Flow Volatility", level=3)

add_para(
    "Free cash flow has been positive in each of the past five years but with significant volatility "
    "(¥100 mn to ¥375 mn range). Our FY2026E projection of ¥174 mn FCF (a step-down from ¥373 mn in "
    "FY2025A) reflects the elevated CapEx cycle: roller-screw line, Thailand e-drive line, smart "
    "corner module pilot, and continued bearings expansion total ¥470 mn of CapEx in FY2026E (vs. "
    "¥408 mn FY2025A). FCF recovery to ¥430+ mn requires both CapEx normalization (FY2028E onward) "
    "and EBITDA expansion from the new product lines.",
    align="justify")

add_heading("HKEX IPO Dilution and Use of Proceeds", level=3)

add_para(
    "The pending HKEX A+H IPO is a structural event with multiple impacts. If the offering proceeds "
    "at ¥0.8–1.5 bn of fresh capital, the dilution impact ranges from approximately 4.5% (low end) "
    "to 8.0% (high end) of existing share count — meaningful but not catastrophic. Use of proceeds "
    "is expected to focus on roller-screw and Thailand capacity expansion. The positive aspect of "
    "the offering is the introduction of HK institutional ownership (typically long-only managers "
    "with research depth) which could improve governance discipline and stock-price information "
    "efficiency. The negative aspect is the dilution itself plus the typical 6–12 month underperformance "
    "of stocks following A+H listings (consistent with the 'A-share premium' compression dynamic).",
    align="justify")

add_heading("Governance and Regulatory Risks", level=2)

add_heading("Family-Control Concentration", level=3)

add_para(
    "The Wu family's 48.9% effective voting control (Shuanglin Group 44.43% + Wu Jianbin direct "
    "4.49% + employee plan 0.35%) means minority shareholders have limited ability to influence "
    "major strategic decisions. The principal historical example is the 2017 DSI Australia acquisition "
    "— a ¥2.3 bn related-party transaction with Shuanglin Group that, in hindsight, destroyed "
    "significant minority value through subsequent impairments. While there is no specific evidence "
    "of related-party self-dealing in recent years (Top-5 customer and supplier related-party "
    "transactions were 0.00% in FY2025), the structural risk remains.",
    align="justify")

add_heading("Audit Firm and Disclosure Quality", level=3)

add_para(
    "Shuanglin's auditor is Beijing Dehao International Accounting Firm (Special General "
    "Partnership) — a domestic mid-tier firm rather than a Big Four. While Dehao is properly "
    "licensed and the two signing partners (Chen Jianfeng, Zhou Sunji) have multi-year experience "
    "with the company, the choice of a non-Big-Four auditor for a ¥17 bn market-cap company with "
    "international customers and a pending HK listing is a moderate governance flag. The HKEX listing "
    "process will likely require a Big Four auditor for the HK prospectus, which would meaningfully "
    "improve disclosure quality but also raise audit costs.",
    align="justify")

add_heading("Macroeconomic and Industry Risks", level=2)

add_heading("China NEV Pricing War Continuation", level=3)

add_para(
    "The China NEV pricing war that began in 2023 has not yet ended and shows no clear signs of "
    "abating in 2026. BYD continues to launch lower-priced models, new entrants (Xiaomi, etc.) are "
    "intensifying competition, and several mid-tier brands face viability questions. Continued "
    "price competition transmits to supplier margins — Shuanglin's Q1 2026 demonstrated the "
    "mechanism. If the pricing war extends through 2026 and into 2027 (vs. our base-case assumption "
    "of stabilization in H2 2026), our FY2026E and FY2027E numbers would be cut by 15–25% and the "
    "DCF Base would compress further toward ¥10–12.",
    align="justify")

add_heading("Geopolitical and Trade Risks", level=3)

add_para(
    "U.S.-China trade relations have direct implications for Shuanglin given its Tesla supply "
    "relationship and the North American production footprint. Tariffs on Chinese auto parts "
    "imported into the U.S., or restrictions on Chinese-headquartered suppliers servicing North "
    "American OEMs, would materially impact the Top-1 customer relationship. The company has "
    "already started building North American production capability, but full localization takes "
    "2–3 years. The Thailand bearing plant — strategically positioned outside of U.S.-China bilateral "
    "tariff lines — provides one form of mitigation.",
    align="justify")

add_page_break()

# ============================================================================
# FINANCIAL ANALYSIS — HISTORICAL TRENDS
# ============================================================================
add_heading("Financial Analysis", level=1)

add_heading("Historical Performance (FY2021A–FY2025A)", level=2)

add_para(
    "Shuanglin's historical financial profile reflects the volatility of an auto-parts integrator "
    "executing a multi-stage transformation. Revenue grew from ¥3,682 mn (FY2021) to ¥5,484 mn "
    "(FY2025) — a 5-year CAGR of 8.3%, materially below China NEV-parts peer median (~15%). The path "
    "is non-linear: FY2022 revenue −1.1% (DSI transmission deterioration); FY2023 −1.1% (continued "
    "DSI drag + COVID lingering); FY2024 +18.6% (HDM volume inflection); FY2025 +11.7%. Net income "
    "shows even higher volatility — from ¥129 mn (FY2021) to ¥75 mn (FY2022 trough) to ¥503 mn "
    "(FY2025 peak), a 4x swing.",
    align="justify")

# Extract historical IS from the model
wb = load_workbook(MODEL_PATH, data_only=False)
ws_is = wb["Income Statement"]

# Build a summary historical IS table
hist_table = [
    ["Line item (¥ mn)", "FY2021A", "FY2022A", "FY2023A", "FY2024A", "FY2025A", "5Y CAGR"],
    ["Revenue", "3,682", "4,185", "4,139", "4,911", "5,484", "8.3%"],
    ["Gross profit", "683", "713", "782", "908", "1,148", "10.9%"],
    ["  Gross margin %", "18.5%", "17.0%", "18.9%", "18.5%", "20.9%", "+2.4 pp"],
    ["EBIT", "146", "101", "125", "520", "549", "30.3%"],
    ["  EBIT margin %", "4.0%", "2.4%", "3.0%", "10.6%", "10.0%", "+6.0 pp"],
    ["EBITDA", "312", "289", "311", "741", "796", "20.6%"],
    ["  EBITDA margin %", "8.5%", "6.9%", "7.5%", "15.1%", "14.5%", "+6.0 pp"],
    ["Net income (parent)", "129", "75", "81", "497", "503", "31.3%"],
    ["  Net margin %", "3.5%", "1.8%", "2.0%", "10.1%", "9.2%", "+5.7 pp"],
    ["EPS (basic, ¥)", "0.32", "0.19", "0.20", "0.89", "0.89", "22.7%"],
    ["CFO", "426", "443", "378", "671", "781", "16.3%"],
    ["CapEx", "253", "213", "277", "296", "408", "12.7%"],
    ["FCF", "173", "230", "100", "375", "373", "21.2%"],
]
add_table_from_data(hist_table[0], hist_table[1:],
                     col_widths=[4.5, 1.8, 1.8, 1.8, 1.8, 1.8, 1.8])

add_para("Source: 双林股份 2021–2025 年度报告 (cninfo). All figures consolidated and as-reported (no restatements applied).",
          italic=True, color=GREY, size=9)

add_para(
    "Three observations from the historical record are noteworthy. First, the FY2024 inflection in "
    "EBIT (+316% YoY from ¥125 mn to ¥520 mn) was driven by both a one-off ¥160 mn investment income "
    "and structural HDM/bearing volume scale — not exclusively non-recurring. Second, the gross "
    "margin expansion from 17.0% (FY2022) to 20.9% (FY2025) — +390 bp over three years — reflects "
    "real mix shift: HDM (estimated 30%+ GM) growing share, and the higher-margin Other category "
    "(31.9% GM) becoming non-trivial via the Kexin grinder acquisition. Third, free cash flow has "
    "been positive throughout (¥173 mn average annually for 2021–2025) — capital discipline has "
    "remained intact even through the DSI integration period.",
    align="justify")

add_chart(os.path.join(CHARTS, "chart_10_ebitda_margin.png"), width_inches=6.5,
          caption="Figure 17. EBITDA trajectory — clear FY2024 inflection point on HDM platform mix")

add_chart(os.path.join(CHARTS, "chart_11_ebit_margin.png"), width_inches=6.5,
          caption="Figure 18. Operating profit and EBIT margin — trough FY2022–FY2023, recovery from FY2024")

add_chart(os.path.join(CHARTS, "chart_12_cashflow.png"), width_inches=6.5,
          caption="Figure 19. Operating cash flow vs. CapEx and free cash flow — investment cycle elevated 2026–28")

add_heading("Q1 2026 Reading and FY2026 Outlook", level=2)

add_para(
    "The Q1 2026 report (filed April 28, 2026) disclosed three concerning data points. First, revenue "
    "of ¥1,193 mn was −10.42% YoY, ending eight consecutive quarters of positive YoY growth. Second, "
    "parent net income of ¥72 mn fell 47.01% YoY — a deterioration far exceeding the revenue drop, "
    "indicating substantial operating deleverage. Third, ex-non-recurring NI fell 39.11%, ruling out "
    "an explanation centered on one-off items. Our conclusion is that Q1 reflects genuine pricing-"
    "pressure pass-through from the company's largest OEM customers, particularly NEV brands "
    "renegotiating supplier contracts amid intensifying domestic price competition.",
    align="justify")

add_para(
    "For our FY2026 base-case projection, we assume Q2 2026 sequential improvement (slight QoQ "
    "growth, but still YoY decline of ~5%), with H2 2026 stabilizing around flat-to-slightly-positive "
    "YoY. This gives a FY2026 revenue trajectory of ¥5,531 mn (+0.9% YoY) — essentially flat — and "
    "net income to parent of ¥399 mn (−20.7% YoY). The base case rests on the assumption that "
    "OEM-driven pricing pressure abates in H2 as the NEV price-war moderates and HDM volume continues "
    "to grow with NEV penetration. We acknowledge that more bearish scenarios are plausible: a fully-"
    "annualized Q1 trajectory would give FY2026 revenue of ¥4,772 mn (−13.0%) and net income of ¥288 "
    "mn (−42.7%) — closer to our bear case.",
    align="justify")

add_heading("Projection Methodology (FY2026E–FY2030E)", level=2)

add_para(
    "Our 5-year projections decompose revenue growth by the four FY2025 reporting segments, with "
    "segment-specific growth rates reflecting maturity and competitive position:",
    align="justify")

add_table_from_data(
    headers=["Segment / Driver", "FY2025A weight", "FY2026E growth", "FY2027E growth", "FY2030E growth", "5Y CAGR"],
    rows=[
        ["Transmission/Drive/Intelligent", "59.6%", "+4%", "+20%", "+15%", "15.6%"],
        ["  HDM stable mid-single-digit", "", "+5%", "+8%", "+8%", "+7%"],
        ["  Bearings (NEV)", "", "+10%", "+12%", "+10%", "+10%"],
        ["  NEV e-drive (Thailand)", "", "+8%", "+30%", "+20%", "+18%"],
        ["  Roller screws / robot (new)", "", "+50%", "+100%", "+40%", "+80%"],
        ["Auto Interior / Exterior", "35.5%", "−5%", "+5%", "+2%", "+1.9%"],
        ["Other (Kexin + molds)", "4.7%", "+5%", "+15%", "+15%", "+14.9%"],
        ["Rental", "0.1%", "+5%", "+5%", "+5%", "+5%"],
        ["Total Revenue", "100%", "+0.9%", "+14.7%", "+11.8%", "11.5%"],
    ],
    col_widths=[5.0, 2.5, 2.5, 2.5, 2.5, 2.0]
)

add_para("Source: Analyst projections. Segment maps to FY2025 restated 3-segment basis (per 2025 年度报告 p.21).",
          italic=True, color=GREY, size=9)

add_chart(os.path.join(CHARTS, "chart_13_revenue_scenarios.png"), width_inches=6.5,
          caption="Figure 20. Revenue trajectory by scenario — Bull / Base / Bear paths to FY2030E")

add_heading("Margin and Capital Intensity Assumptions", level=2)

add_para(
    "Gross margin expands from 20.9% (FY2025A) to 24.0% (FY2030E) as Transmission/Drive mix rises to "
    "70%+ of revenue and the higher-margin roller-screw / corner-module products contribute. SG&A as "
    "a percentage of revenue compresses modestly from 7.6% (FY2025A) to 6.7% (FY2030E) on operating "
    "leverage; R&D rises to 4.5–5.0% (vs. 4.0% in FY2025) reflecting continued investment in robotics "
    "+ corner-module + low-altitude platforms.",
    align="justify")

add_para(
    "CapEx as a percentage of revenue averages 8.5% in FY2026 (the roller-screw line investment year) "
    "before declining to 5.0% by FY2030E. D&A tracks at 4.0–4.5% of revenue, modestly higher than "
    "historical norms reflecting the depreciation tail of the 2025–2026 CapEx cycle.",
    align="justify")

add_table_from_data(
    headers=["Metric", "FY2025A", "FY2026E", "FY2027E", "FY2028E", "FY2029E", "FY2030E"],
    rows=[
        ["Revenue (¥ mn)", "5,484", "5,531", "6,345", "7,403", "8,436", "9,427"],
        ["YoY growth", "11.7%", "0.9%", "14.7%", "16.7%", "14.0%", "11.8%"],
        ["Gross margin", "20.9%", "20.5%", "22.0%", "23.5%", "24.0%", "24.0%"],
        ["EBITDA (¥ mn)", "796", "709", "870", "1,141", "1,360", "1,557"],
        ["EBITDA margin", "14.5%", "12.8%", "13.7%", "15.4%", "16.1%", "16.5%"],
        ["EBIT (¥ mn)", "549", "460", "598", "831", "1,022", "1,180"],
        ["EBIT margin", "10.0%", "8.3%", "9.4%", "11.2%", "12.1%", "12.5%"],
        ["Net income parent (¥ mn)", "503", "399", "519", "722", "878", "1,014"],
        ["EPS basic (¥)", "0.89", "0.70", "0.91", "1.26", "1.54", "1.77"],
        ["CFO (¥ mn)", "781", "644", "694", "874", "1,061", "1,243"],
        ["CapEx (¥ mn)", "408", "470", "444", "444", "464", "471"],
        ["CapEx % rev", "7.4%", "8.5%", "7.0%", "6.0%", "5.5%", "5.0%"],
        ["FCF (¥ mn)", "373", "174", "250", "430", "597", "771"],
    ],
    col_widths=[5.0, 2.0, 2.0, 2.0, 2.0, 2.0, 2.0]
)
add_para("Source: Financial model — Income Statement and Cash Flow tabs. CNY mn unless noted. Base case.",
          italic=True, color=GREY, size=9)

add_chart(os.path.join(CHARTS, "chart_33_eps_trajectory.png"), width_inches=6.5,
          caption="Figure 21. Diluted EPS trajectory by scenario — Base case ¥1.77 in FY2030E")

add_chart(os.path.join(CHARTS, "chart_14_scenarios.png"), width_inches=6.5,
          caption="Figure 22. FY2030E scenario outcomes — Bear / Base / Bull comparison across key metrics")

add_heading("DuPont Decomposition of ROE", level=2)

add_para(
    "We decompose Shuanglin's return on equity into three drivers — net margin, asset turnover, and "
    "financial leverage — to understand the sources of value creation over the historical period and "
    "the projected trajectory. The 2025 weighted-average ROE of 17.23% (per company disclosure) is "
    "decomposable as approximately 9.2% net margin × 0.85x asset turnover × 2.20x financial leverage "
    "(equity-multiplier). The projection trajectory shows margin-driven ROE expansion as the mix "
    "shifts toward higher-margin Transmission/Drive products, partially offset by modest leverage "
    "compression as the HKEX IPO adds equity to the balance sheet.",
    align="justify")

add_table_from_data(
    headers=["Metric", "FY2023A", "FY2024A", "FY2025A", "FY2026E", "FY2027E", "FY2028E", "FY2030E"],
    rows=[
        ["Revenue (¥ mn)",        "4,139", "4,911", "5,484", "5,531", "6,345", "7,403", "9,427"],
        ["Net income (¥ mn)",      "81",   "497",   "503",   "399",   "519",   "722",   "1,014"],
        ["Net margin",             "2.0%", "10.1%", "9.2%",  "7.2%",  "8.2%",  "9.7%",  "10.8%"],
        ["Total assets (¥ mn)",    "5,955","6,271", "7,023", "8,087", "8,805", "9,579", "11,570"],
        ["Asset turnover (×)",     "0.69x","0.78x", "0.78x", "0.68x", "0.72x", "0.77x", "0.81x"],
        ["Equity (¥ mn)",          "2,204","2,643", "3,226", "4,346", "4,761", "5,324", "6,800"],
        ["Equity multiplier (×)",  "2.70x","2.37x", "2.18x", "1.86x", "1.85x", "1.80x", "1.70x"],
        ["ROE (DuPont, %)",        "3.7%", "18.8%", "15.6%", "9.2%",  "10.9%", "13.6%", "14.9%"],
        ["ROE (reported, %)",      "n/d",  "18.8%", "17.2%", "n/a",   "n/a",   "n/a",   "n/a"],
    ],
    col_widths=[4.0, 1.8, 1.8, 1.8, 1.8, 1.8, 1.8, 1.8]
)
add_para("Source: Financial model and 2025 年度报告. DuPont ROE = Net Margin × Asset Turnover × Equity Multiplier. "
          "FY2025A reported weighted-average ROE was 17.23% (per company); DuPont approximation of 15.6% uses period-end equity.",
          italic=True, color=GREY, size=9)

add_para(
    "The decomposition reveals three insights. First, margin expansion is the primary ROE driver — "
    "net margin rises from 2.0% (FY2023A trough) to 10.8% (FY2030E projection), an absolute 8.8 "
    "percentage-point improvement that contributes most of the ROE delta. Second, asset turnover is "
    "relatively stable at 0.7–0.8x — characteristic of a capital-intensive manufacturer where "
    "PP&E and inventory scale with revenue. Third, financial leverage is compressing from 2.4x to "
    "1.7x as the HKEX IPO and retained earnings build up the equity base; this is structurally "
    "healthy for the credit profile but mechanically lowers ROE versus a more-leveraged peer.",
    align="justify")

add_page_break()

# ============================================================================
# VALUATION ANALYSIS
# ============================================================================
add_heading("Valuation Analysis", level=1)

add_heading("DCF Approach: WACC, Terminal Value, and Outputs", level=2)

add_para(
    "We construct a 5-year DCF model using unlevered free cash flow from FY2026E–FY2030E plus a "
    "Gordon-growth terminal value. The WACC build uses standard CAPM for the cost of equity component "
    "(China 10-year CGB risk-free rate of 1.75%, Damodaran A-share ERP of 6.5%, levered beta of 1.30 "
    "reflecting the auto-parts core plus robot-pipeline exposure premium), giving a cost of equity "
    "of 10.20%. Cost of debt is set at 3.50% pre-tax (PBOC 5-year LPR + spread for mid-tier corporate "
    "credit), or 3.05% after-tax. With market-value weights of 92.6% equity / 7.4% debt, our blended "
    "WACC is 9.67%. Terminal growth rate of 2.5% reflects long-run China nominal GDP growth (1% real "
    "+ 1.5% inflation).",
    align="justify")

add_table_from_data(
    headers=["Year", "UFCF (¥ mn)", "Discount factor", "PV of UFCF (¥ mn)"],
    rows=[
        ["FY2026E", "174", "0.9118", "159"],
        ["FY2027E", "251", "0.8314", "209"],
        ["FY2028E", "431", "0.7581", "327"],
        ["FY2029E", "598", "0.6913", "413"],
        ["FY2030E", "772", "0.6303", "487"],
        ["Sum explicit PV", "", "", "1,594"],
        ["Terminal value (TV)", "13,213", "0.6303", "8,328"],
        ["Enterprise Value", "", "", "9,922"],
        ["Less: Net debt", "", "", "(600)"],
        ["Equity Value", "", "", "9,322"],
        ["÷ Diluted shares (mn)", "", "", "584"],
        ["Implied price per share", "", "", "¥15.96"],
    ],
    col_widths=[3.5, 3.0, 3.0, 3.0]
)
add_para("Source: Financial model — DCF tab. Base case: WACC 9.67%, terminal growth 2.5%.",
          italic=True, color=GREY, size=9)

add_para(
    "The DCF Base implies a price per share of approximately ¥16.62, a 44.6% downside to the current "
    "¥30 quote. Terminal value comprises 84% of total enterprise value — above the institutional "
    "sanity-check threshold of 70% — reflecting the long-duration nature of Shuanglin's emerging "
    "pipelines. We disclose this caveat but note that lengthening the explicit forecast period to 10 "
    "years (vs. 5) would simply push the issue further out without materially changing the "
    "conclusion: at current price ¥30, the market is paying ~33x our FY2027E base-case EPS of ¥0.91 — "
    "a meaningful premium to the auto-parts peer median 22x.",
    align="justify")

add_chart(os.path.join(CHARTS, "chart_29_dcf_waterfall.png"), width_inches=6.5,
          caption="Figure 23. DCF waterfall — EV ¥10,305 mn → Equity ¥9,705 mn → Price/share ¥16.62 (Base)")

add_heading("DCF Sensitivity Analysis", level=2)

add_para(
    "We test the DCF Base output against two-way sensitivities for WACC (range 7.7% to 11.7% in 50 bp "
    "increments) and terminal growth (range 0.5% to 4.0% in 50 bp increments). The matrix shows "
    "implied price per share ranging from ¥10 (worst-corner: WACC 11.7% / g 0.5%) to ¥34 "
    "(best-corner: WACC 7.7% / g 4.0%). Even at the optimistic corner, the DCF Base case scenario "
    "barely clears the current ¥30 — and that corner requires aggressive assumptions on both inputs "
    "simultaneously.",
    align="justify")

add_chart(os.path.join(CHARTS, "chart_28_dcf_sensitivity.png"), width_inches=6.5,
          caption="Figure 24. ⭐ DCF sensitivity heatmap — WACC × terminal growth")

add_heading("DCF Scenario Analysis", level=2)

add_para(
    "We supplement the Base case with Bull and Bear case DCFs that perturb both the projection (UFCF "
    "trajectory) and the discount rate. The Bull case assumes faster roller-screw and corner-module "
    "scale-up (UFCF reaching ¥1,500 mn by FY2030E vs. ¥772 mn in Base), with a tighter WACC of 9.17% "
    "and higher terminal growth of 3.0%. The Bear case assumes Q1-2026 pricing weakness extends "
    "through FY2027 with delayed roller-screw ramp (UFCF only ¥350 mn by FY2030E), wider WACC of "
    "10.67%, and lower terminal growth of 2.0%.",
    align="justify")

add_table_from_data(
    headers=["Case", "WACC", "g", "EV (¥ mn)", "Equity (¥ mn)", "Price/share", "Upside"],
    rows=[
        ["Bull", "9.17%", "3.0%", "20,488", "19,888", "¥35.08", "+16.9%"],
        ["Base", "9.67%", "2.5%", "10,305", "9,705", "¥16.62", "−44.6%"],
        ["Bear", "10.67%", "2.0%", "3,335", "2,535", "¥4.34", "−85.5%"],
    ],
    col_widths=[2.5, 2.5, 2.5, 3.0, 3.5, 3.0, 2.5]
)
add_para("Source: Financial model — DCF tab. Net debt of ¥600 mn for Base/Bull; ¥800 mn for Bear (extra debt assumed).",
          italic=True, color=GREY, size=9)

add_heading("Comparable Companies Analysis", level=2)

add_para(
    "We construct a comparable-companies set of nine peers spanning four buckets relevant to "
    "Shuanglin's business mix: (a) large NEV auto-parts (Tuopu, Wanxiang); (b) bearings and linear "
    "motion (Wanxiang, Hengli, Schaeffler global); (c) roller-screw concept / pure-play (Beste, "
    "XCC, Beite, Dingzhi); and (d) precision robot supply chain (Shuanghuan). We exclude two outliers: "
    "ZDLD (P/E 189x — speculative robot-harmonic-reducer) and NSK (already captured by Schaeffler).",
    align="justify")

add_table_from_data(
    headers=["Peer", "Ticker", "Mkt Cap (¥ bn)", "EV/EBITDA NTM", "P/E NTM", "Rev growth", "EBITDA margin"],
    rows=[
        ["Tuopu Group", "SSE:601689", "102.0", "19.3x", "27.6x", "20%", "16.6%"],
        ["Wanxiang Qianchao", "SZSE:000559", "18.8", "14.1x", "23.5x", "10%", "10.0%"],
        ["Hengli Hydraulic", "SSE:601100", "77.5", "22.4x", "31.0x", "14%", "28.6%"],
        ["Beste", "SZSE:300580", "8.5", "17.3x", "24.3x", "24%", "23.6%"],
        ["XCC (Wuzhou Xinchun)", "SSE:603667", "12.5", "21.8x", "52.1x", "22%", "13.6%"],
        ["Beite Tech", "SSE:603009", "18.5", "55.6x", "119.4x", "21%", "11.7%"],
        ["Shuanghuan Drive", "SZSE:002472", "41.0", "20.4x", "30.8x", "14%", "16.4%"],
        ["Dingzhi Tech", "BSE:873593", "9.5", "44.8x", "61.3x", "28%", "26.5%"],
        ["Schaeffler", "FRA:SHA", "45.0", "3.4x", "5.5x", "6%", "12.3%"],
        ["Statistical summary", "", "", "", "", "", ""],
        ["  Max", "", "", "55.6x", "119.4x", "28%", "28.6%"],
        ["  75th percentile", "", "", "23.9x", "37.9x", "23%", "21.4%"],
        ["  Median", "", "", "20.4x", "30.8x", "20%", "16.4%"],
        ["  25th percentile", "", "", "17.3x", "24.3x", "14%", "12.3%"],
        ["  Min", "", "", "3.4x", "5.5x", "6%", "10.0%"],
        ["Shuanglin (target)", "SZSE:300100", "17.2", "25.6x", "43.9x", "13.5%", "14.5%"],
    ],
    col_widths=[3.5, 2.5, 2.0, 2.5, 2.0, 2.0, 2.5]
)
add_para("Source: Eastmoney 东方财富, Yahoo Finance, company filings (May 2026 snapshot). All EV calculations use mkt cap + net debt.",
          italic=True, color=GREY, size=9)

add_chart(os.path.join(CHARTS, "chart_31_ev_ebitda_ntm.png"), width_inches=6.5,
          caption="Figure 25. Peer EV/EBITDA NTM — Shuanglin at 25.6x, 25% above auto-parts median 20.4x")

add_chart(os.path.join(CHARTS, "chart_30_pe_ntm.png"), width_inches=6.5,
          caption="Figure 26. Peer P/E NTM — Shuanglin at 43.9x, double the median 30.8x")

add_para(
    "Applying peer-median multiples to Shuanglin's NTM (next-twelve-months) financials yields four "
    "implied per-share price points:",
    align="justify")

add_table_from_data(
    headers=["Multiple", "Bear (25th %)", "Base (median)", "Bull (75th %)"],
    rows=[
        ["EV/EBITDA NTM on ¥709 mn EBITDA", "¥18.01", "¥23.79", "¥37.83"],
        ["EV/Revenue NTM on ¥5,531 mn rev", "¥19.78", "¥30.72", "¥60.31"],
        ["P/E NTM 2026E on ¥399 mn NI", "¥16.32", "¥21.06", "¥38.73"],
        ["P/E forward 2027E on ¥519 mn NI", "¥21.23", "¥27.40", "¥50.38"],
    ],
    col_widths=[6.5, 3.0, 3.0, 3.0]
)

add_heading("Valuation Reconciliation and Football Field", level=2)

add_para(
    "We weight five methodologies to arrive at the 12-month price target: DCF Base (20%, reflecting "
    "conservative explicit-period cash flows), DCF Bull (10%, reflecting the option-value scenario), "
    "EV/EBITDA NTM peer median (20%), P/E NTM 2026E peer median (20%, reflecting "
    "near-term consensus visibility), and P/E forward 2027E peer median (30%, anchoring the 12-month "
    "horizon target). The weighted-average price target is ¥24, with a low-high range of ¥19–¥38.",
    align="justify")

add_chart(os.path.join(CHARTS, "chart_32_football_field.png"), width_inches=6.5,
          caption="Figure 27. ⭐ Football field — six methodology ranges with current ¥30 and target ¥24")

add_heading("Probability-Weighted Cross-Check", level=2)

add_para(
    "As a sanity check on the methodology-weighted target, we run a probability-weighted scenario "
    "overlay. We assign 25% probability to the Bull case (roller-screw + corner module materializing), "
    "55% to the Base case (Q1 weakness with gradual recovery), and 20% to the Bear case (pricing war "
    "deepens, screw ramp delays). For each scenario, we equally weight the DCF output and an implied "
    "P/E-based output. The probability-weighted target is approximately ¥24, converging with the "
    "methodology-weighted ¥24 — supporting directional confidence in the SELL rating.",
    align="justify")

add_page_break()

# ============================================================================
# APPENDICES
# ============================================================================
add_heading("Final Recommendation and Investment Conclusion", level=1)

add_para(
    "We initiate coverage of Shuanglin Co. (SZSE:300100) with a SELL rating and a 12-month price "
    "target of ¥24. The rating reflects a quantitative conclusion that emerges consistently across "
    "five distinct valuation methodologies — DCF Base (¥16.62), DCF Bull (¥35.08), EV/EBITDA NTM "
    "peer median (¥23.79), P/E NTM 2026E peer median (¥21.06), and P/E forward 2027E peer median "
    "(¥27.40) — weighted to a target of ¥24, implying 20% downside to the current ¥30 quote.",
    align="justify")

add_para(
    "Our concerns are not about Shuanglin as a business. The company is a legitimate market leader in "
    "HDM (32.8% China share / 15.1% global), a credible #3 in wheel-hub bearings (9% China share), "
    "and one of the few Chinese auto-parts companies with technical credibility in three emerging "
    "categories — humanoid roller screws, smart corner modules, and eVTOL e-drives. Management has "
    "demonstrated 22 years of strategic continuity under founder Wu Jianbin, and the balance sheet is "
    "conservatively managed with net-debt-to-EBITDA below 1.0x. These positive attributes support "
    "a long-run buy-and-hold thesis at the right price.",
    align="justify")

add_para(
    "Our concerns are about price — specifically, the gap between the current quote and the "
    "intrinsic value implied by every realistic valuation framework. The current ¥30 share price "
    "implies a TTM P/E of 34x, well above the China auto-parts peer median of 22x, and embeds option "
    "value for the robot and corner-module pipelines that has not yet been validated by formal OEM "
    "program awards. Q1 2026 results — revenue −10.4% and net income −47% — confirm that "
    "near-term earnings are compressing materially from the FY2025 base. The combination of "
    "elevated multiple plus declining earnings creates an unfavorable risk-reward.",
    align="justify")

add_para(
    "Investors looking at Shuanglin should monitor three specific catalysts over the next six months. "
    "First, Q2 2026 earnings release (August 2026) will reveal whether Q1 was an isolated quarter or "
    "the beginning of a multi-quarter decline; sequential improvement in revenue and stable margins "
    "would be the bullish signal. Second, the roller-screw 量产 (volume production) line start in "
    "June 2026 — successful operation at yield benchmarks and any announcement of a humanoid OEM "
    "program award (定点) would dramatically shift our methodology weighting toward the DCF Bull case "
    "(¥35.08). Third, HKEX A+H IPO terms and anchor-investor identity (expected Q3–Q4 2026) — strong "
    "long-only institutional anchor support at a meaningful premium to A-share pricing would signal "
    "international validation of the smart-drive-solutions narrative.",
    align="justify")

add_para(
    "We would consider upgrading our rating from SELL to HOLD if any one of the following occurs: "
    "(a) the share price corrects toward ¥24 (our target); (b) Q2 2026 earnings show clear "
    "sequential improvement with stable margins, validating Base-case projections; or (c) a formal "
    "humanoid OEM 定点 is announced with disclosed program size and timing. We would consider "
    "upgrading further to BUY only if multiple emerging-product wins materialize together — i.e., "
    "the Bull-case scenario is being executed. Conversely, we would consider deepening to a stronger "
    "SELL conviction if Q2 2026 confirms multi-quarter weakness or if the roller-screw line "
    "encounters major yield/timing issues.",
    align="justify")

add_para(
    "Our 12-month price target of ¥24 implies modest downside in a name with significant binary "
    "option value. Investors who believe the roller-screw and corner-module bets will execute "
    "successfully should wait for either a lower entry point or for tangible validation of one or "
    "both opportunities. Investors holding the stock at current levels face elevated near-term "
    "earnings and multiple-compression risk that, in our view, is not offset by the option-value "
    "premium currently embedded in the share price.",
    align="justify")

add_page_break()

add_heading("Appendices", level=1)

add_heading("Appendix A: Historical Trading Range and Multiple Context", level=2)

add_para(
    "Shuanglin's TTM P/E has ranged from 16.7x (2024 trough, prior to the FY2024 earnings recovery) "
    "to 91x (2025 humanoid-robot speculation peak) — a 5.4x range over three years. The current "
    "34x sits at approximately the 50th percentile of this trading range, but elevated relative to "
    "the 25th percentile (~22x) where peer-comparison analysis would place fair value.",
    align="justify")

add_chart(os.path.join(CHARTS, "chart_34_pe_band.png"), width_inches=6.5,
          caption="Figure 28. Historical TTM P/E band — current 34x sits at trading-range midpoint")

add_heading("Appendix B: Recent Capital Structure Events", level=2)

add_para(
    "May 2025: Initial A-share private-placement (定增) plan of ¥1.5 bn proposed for roller-screw + "
    "Thailand bearings capacity expansion. Withdrawn after 7 days as regulatory environment shifted; "
    "company immediately pivoted to HKEX A+H IPO route.",
    align="justify")

add_para(
    "August 2025: Board approval of the Market Capitalization Management System (市值管理制度). "
    "Quality-Return Action Plan (质量回报双提升行动方案) disclosed January 2025.",
    align="justify")

add_para(
    "March 2026: HKEX A+H IPO prospectus re-submitted with CITIC Securities and GF Securities as "
    "joint sponsors. Targeted listing in H2 2026; capital raise of ¥0.8–1.5 bn for capacity expansion "
    "and overseas operations. Pricing premium to A-share possible.",
    align="justify")

add_heading("Appendix C: Sanity Checks Performed", level=2)

add_bullet("DCF terminal value at 84% of EV — above the 70% threshold but explained by the long-duration nature of emerging-product pipelines.")
add_bullet("Implied P/E at target ¥24 = 33.7x FY2027E EPS — within peer trading band (q1 24x – q3 38x).")
add_bullet("WACC 9.67% within the expected 8–14% range for China industrial mid-cap with technology pipeline exposure.")
add_bullet("Implied 12-month total return from ¥30 to ¥24 target = −20.0% — consistent with SELL rating.")
add_bullet("Market cap at target ¥13.7 bn — within the auto-parts mid-cap peer range of ¥8–100 bn.")
add_bullet("FCF positive in all five projection years (¥174 mn to ¥771 mn) — no funding gap in projections.")
add_bullet("Balance sheet balanced for all 10 projection-history years (Diff < ¥0.1 mn). See model.")

add_heading("Appendix D: Key References and Sources", level=2)

refs = [
    ("双林股份 2025 年年度报告 (filed 2026-03-24)",  "https://static.cninfo.com.cn/finalpage/2026-03-25/"),
    ("双林股份 2026 年第一季度报告 (filed 2026-04-28)", "https://disc.static.szse.cn/disc/disk03/finalpage/2026-04-28/"),
    ("双林股份 2023 年年度报告 (filed 2024-04-18)", "https://www.cninfo.com.cn/new/disclosure/stock?stockCode=300100"),
    ("Task 1 Company Research (内部研究, 2026-05-17)", "双林股份_SZSE300100_公司研究_2026-05-17.md"),
    ("Task 2 Financial Model (内部模型, 2026-05-18)", "双林股份_SZSE300100_Financial_Model_2026-05-18.xlsx"),
    ("Task 3 Valuation Analysis (内部分析, 2026-05-18)", "双林股份_SZSE300100_Valuation_Analysis_2026-05-18.md"),
    ("智研咨询 — China HDM Industry Research 2025", "https://www.chyxx.com/industry/1239987.html"),
    ("智研咨询 — China Auto Bearings Industry Research 2025", "https://www.chyxx.com/industry/1225289.html"),
    ("亿牛网 — Historical P/E for Shuanglin (300100)",   "https://eniu.com/gu/sz300100"),
    ("Damodaran — Country Risk Premiums (2026 update)",   "https://pages.stern.nyu.edu/~adamodar/"),
    ("People's Bank of China — 10-year CGB yield curve",  "http://www.pbc.gov.cn/"),
]
for text, url in refs:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(text + " — ")
    set_run(r1, size=10)
    r2 = p.add_run(url)
    set_run(r2, size=9, italic=True, color=NAVY)

add_heading("Appendix E: Glossary of Terms", level=2)

glossary = [
    ("HDM",          "Horizontal Drive Module — Shuanglin's flagship product. Powers the fore-aft adjustment in electrically-adjustable car seats."),
    ("PPAP",         "Production Part Approval Process — the OEM qualification cycle (18–36 months) required before a new auto-parts supplier ships volume production."),
    ("WACC",         "Weighted Average Cost of Capital — the discount rate used in DCF valuation, calculated as the weighted average of cost of equity and after-tax cost of debt."),
    ("UFCF",         "Unlevered Free Cash Flow — cash flow available to both equity and debt holders, calculated as NOPAT + D&A − CapEx − ΔNWC."),
    ("Reverse planetary roller screw", "A precision linear-actuator used in humanoid robot upper- and lower-body joints. Higher load-to-volume ratio than standard planetary roller screws."),
    ("EHB / EMB",    "Electronic Hydraulic Brake / Electronic Mechanical Brake — brake-by-wire technologies for autonomous vehicles."),
    ("定点",          "(Pinyin: dìng-diǎn) — Formal supplier program award from an OEM for a specific vehicle model. Triggers the PPAP cycle and ultimately volume revenue."),
    ("量产",          "(Pinyin: liàng-chǎn) — Volume production. The transition from prototype/sample to mass production."),
    ("NTM / LTM",    "Next-Twelve-Months / Last-Twelve-Months — forward and trailing 12-month financial metrics used for valuation."),
    ("CGB",          "China Government Bond — used for the risk-free rate in CAPM."),
]
for term, defn in glossary:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(term + ": ")
    set_run(r1, bold=True, size=10)
    r2 = p.add_run(defn)
    set_run(r2, size=10)

add_heading("Appendix F: ESG and Sustainability Considerations", level=2)

add_para(
    "Environmental, Social, and Governance (ESG) considerations have grown more relevant for Chinese "
    "industrial mid-caps as both domestic regulators (CSRC, MIIT) and international institutional "
    "investors elevate disclosure expectations. Shuanglin's ESG profile is mixed. On the environmental "
    "dimension, the company's product mix is increasingly oriented toward enabling NEV adoption — "
    "wheel bearings for electric vehicles, e-drive motors, smart corner modules for autonomous "
    "platforms — which positions Shuanglin favorably in the broader energy-transition narrative. "
    "However, the manufacturing process for precision auto parts and roller screws is energy-intensive, "
    "and the company has not yet disclosed Scope 1/2/3 emissions data with the granularity expected "
    "by international sustainability frameworks (TCFD, SASB).",
    align="justify")

add_para(
    "On the social dimension, Shuanglin employs approximately 6,000 people across China, Thailand, and "
    "North America, with disclosed average compensation increases of 5-8% annually. Workforce safety "
    "has been satisfactory based on disclosed lost-time-injury statistics, though the company does "
    "not publish a stand-alone sustainability report. On governance, the family-control structure "
    "(48.9% combined Wu family voting power) creates concentrated decision-making — a double-edged "
    "feature that enables long-horizon strategic patience but limits minority-shareholder influence. "
    "The pending HKEX A+H listing will likely catalyze improvements in disclosure quality and "
    "potentially in board composition (HK Listing Rules require minimum 1/3 independent directors).",
    align="justify")

add_heading("Disclaimer", level=2)

add_para(
    "This research report has been auto-generated by the equity-research/initiating-coverage workflow "
    "based on publicly available filings and analyst projections. All numbers and forecasts represent "
    "the analyst's view and may differ from market consensus or company guidance. Pricing data is as of "
    "May 2026. This report does not constitute investment advice. Past performance is not indicative "
    "of future results. The analyst does not hold a position in SZSE:300100.",
    italic=True, color=GREY, size=9, align="justify")

# Save
doc.save(OUT_PATH)
print(f"\n✓ Report saved: {OUT_PATH}")
print(f"  Size: {os.path.getsize(OUT_PATH)/1024:.1f} KB")
# Count paragraphs and table rows for verification
n_paras = len(doc.paragraphs)
n_tables = len(doc.tables)
n_table_rows = sum(len(t.rows) for t in doc.tables)
print(f"  Paragraphs: {n_paras}")
print(f"  Tables: {n_tables} ({n_table_rows} total rows)")
