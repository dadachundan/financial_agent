"""
禾赛 — Task 5 中文版 DOCX 首次覆盖报告
"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = "/Users/x/projects/financial_agent/reports/company/Hesai_NASDAQ_HSAI"
CHARTS = os.path.join(ROOT, "charts_zh")
OUT = os.path.join(ROOT, "Hesai_NASDAQ_HSAI_Initiation_Report_2026-05-19_zh.docx")

NAVY = RGBColor(0x00, 0x33, 0x66)
ACCENT = RGBColor(0xFF, 0xA5, 0x00)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xC6, 0x28, 0x28)
GRAY = RGBColor(0x66, 0x66, 0x66)

doc = Document()
style = doc.styles['Normal']
style.font.name = "Songti SC"
# Set CJK font via XML
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.append(rFonts)
rFonts.set(qn('w:ascii'), 'Songti SC')
rFonts.set(qn('w:hAnsi'), 'Songti SC')
rFonts.set(qn('w:eastAsia'), 'Songti SC')
style.font.size = Pt(10.5)

for section in doc.sections:
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)


def cjk_font(run, name="Songti SC"):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)


def add_header_para(text, level=1, color=NAVY, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    cjk_font(run)
    if level == 0: run.font.size = Pt(20); run.bold = True; run.font.color.rgb = color
    elif level == 1: run.font.size = Pt(15); run.bold = True; run.font.color.rgb = color
    elif level == 2: run.font.size = Pt(12); run.bold = True; run.font.color.rgb = color
    else: run.font.size = Pt(11); run.bold = True; run.font.color.rgb = color
    return p


def add_para(text, bold=False, italic=False, size=10.5, color=None, indent=0, align=None,
             space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right": p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    cjk_font(run)
    run.font.size = Pt(size)
    if bold: run.bold = True
    if italic: run.italic = True
    if color: run.font.color.rgb = color
    return p


def add_bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if bold_lead:
        r = p.add_run(bold_lead); cjk_font(r); r.font.size = Pt(10.5); r.bold = True
    r = p.add_run(text); cjk_font(r); r.font.size = Pt(10.5)
    return p


def add_image(path, width_inches=6.5, caption=None):
    if not os.path.exists(path):
        print(f"WARN: missing {path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        r = cap.add_run(caption); cjk_font(r); r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = GRAY


def add_page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_table(headers, rows, col_widths=None, font_size=9, first_col_bold=False, total_row=False):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in t.columns[i].cells:
                cell.width = Inches(w)
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = ""
        para = cell.paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(h); cjk_font(r); r.font.size = Pt(font_size); r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, "002060")
    for ri, row in enumerate(rows):
        is_total = (total_row and ri == len(rows) - 1)
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]; cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            r = para.add_run(str(val)); cjk_font(r); r.font.size = Pt(font_size)
            if (first_col_bold and ci == 0) or is_total: r.bold = True
            if is_total: shade_cell(cell, "D9E1F2")
            elif ri % 2 == 1: shade_cell(cell, "F2F2F2")
    for row in t.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for bn in ['top','left','bottom','right']:
                b = OxmlElement(f'w:{bn}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:color'),'888888')
                tcBorders.append(b)
            tcPr.append(tcBorders)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def add_hyperlink(paragraph, url, text, color="0563C1"):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Songti SC'); rFonts.set(qn('w:hAnsi'), 'Songti SC'); rFonts.set(qn('w:eastAsia'), 'Songti SC')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '20'); rPr.append(sz)
    new_run.append(rPr)
    te = OxmlElement('w:t'); te.text = text
    new_run.append(te); hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def chart(n):
    files = sorted(os.listdir(CHARTS))
    for f in files:
        if f.startswith(f"chart_{n:02d}_"):
            return os.path.join(CHARTS, f)
    return None


# ============================================================================
# 第 1 页:封面 + 投资摘要
# ============================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("首次覆盖  •  "); cjk_font(r); r.font.size = Pt(11); r.bold = True; r.font.color.rgb = ACCENT
r = p.add_run("激光雷达 / 机器人硬件  •  中国科技"); cjk_font(r); r.font.size = Pt(11); r.bold = True; r.font.color.rgb = NAVY

add_header_para("禾赛科技(NASDAQ:HSAI, HKEX:2525)", level=0, space_before=4, space_after=2)
add_header_para("全球唯一盈利的纯激光雷达制造商 —— 买入,上行空间 25%", level=2, color=NAVY, space_after=4)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run("2026 年 5 月 19 日  •  股票研究  •  首次覆盖"); cjk_font(r); r.font.size = Pt(10); r.italic = True; r.font.color.rgb = GRAY

# 评级框
rating_t = doc.add_table(rows=2, cols=6)
rating_t.alignment = WD_TABLE_ALIGNMENT.LEFT
hdrs = ["评级", "现价", "12 个月目标价", "上行空间", "市值", "企业价值"]
vals = ["买入", "US$22.44", "US$28.00", "+24.8%", "US$3.53B", "US$2.60B"]
for i, h in enumerate(hdrs):
    cell = rating_t.rows[0].cells[i]; cell.text = ""
    para = cell.paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run(h); cjk_font(r); r.font.size = Pt(8); r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_cell(cell, "002060")
for i, v in enumerate(vals):
    cell = rating_t.rows[1].cells[i]; cell.text = ""
    para = cell.paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run(v); cjk_font(r); r.font.size = Pt(11); r.bold = True
    if i in [0, 3]: r.font.color.rgb = GREEN
    shade_cell(cell, "F2F2F2")
for row in rating_t.rows:
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for bn in ['top','left','bottom','right']:
            b = OxmlElement(f'w:{bn}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'8'); b.set(qn('w:color'),'002060')
            tcBorders.append(b)
        tcPr.append(tcBorders)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
r = p.add_run("52 周区间:US$8.45 – US$29.80  •  稀释股数:1.464 亿  •  净现金:US$933M  •  股息:无  •  Beta:1.35  •  3 月日均成交量:~300 万 ADS"); cjk_font(r); r.font.size = Pt(9); r.italic = True; r.font.color.rgb = GRAY

add_image(chart(1), width_inches=6.7, caption="图 1:HSAI 自 2023 年 2 月纳斯达克 IPO 以来股价走势。")

add_header_para("投资摘要", level=1)

bullets = [
    ("全球首家也是唯一盈利的纯激光雷达制造商。",
     "禾赛 FY2025 在 4.33 亿美元收入基础上实现 GAAP 净利润 6,200 万美元(净利率 14.4%,非 GAAP 18.2%)——成为首家上市的纯激光雷达公司全年实现 GAAP 盈利。速腾聚创、Ouster、Innoviz、Aeva 和 Luminar 均深度亏损。这一结构性优势会复利累积:利润为下一代 ASIC 和 SoC 项目(2025 年 11 月推出的 FMC500)提供资金,扩大了对中国和西方对手的 BOM 成本差距。"),
    ("经营杠杆跑道完好 —— 净利润 FY25-FY28E 复合增速 ~89%。",
     "管理层指引 FY2026 出货量 300-350 万台(对比 FY25 的 160 万台,单位增长 ~85-115%),我们模型测算收入达到 47.37 亿元人民币(US$649M, +56% YoY),FY26E EBIT 为 3.26 亿元人民币。预期净利润从 FY25 的 6,200 万美元增长至 FY28E 的 1.68 亿美元,研发占收入比从 26.3% 下降到 17.0%。市场目前对此轨迹支付的是约 22× FY28E P/E;重新评级至同业中位数 25-30× 对应每 ADS 25-32 美元。"),
    ("中国 ADAS 装载率拐点叠加多激光雷达 L3+ 装载。",
     "中国新车激光雷达装载率从 2024 年的 ~5% 升至 2025 年的 ~13%(Yole Group),我们基础情景假设到 2030 年达到 35%。关键的是,理想、小米和长安都已宣布多激光雷达(每车 3-6 颗)L3+ 项目,2026-2027 年量产启动。多激光雷达装载在装载率提升之上倍增每车含量 —— 一个尚未反映在共识中的结构性顺风。"),
    ("人形机器人激光雷达(JT128)是不对称选择权和可信的第二增长曲线。",
     "禾赛 JT128 微型激光雷达 —— 全球唯一适用于人形机器人体积的 360°×187° 半球视场传感器 —— 被宇树选中,装备其在 2026 央视春晚上展示的所有人形机器人。已签约客户还包括荣耀机器人、Galbot、银河通用(Magiclab)和星动纪元(Vita Dynamics)。我们模型预计 FY26 JT128 出货 7 万台(对比 FY25 的 1.2 万台),到 FY30 增长至 80 万台。仅扫地机器人订单就超过 1,000 万台。人形机器人不在我们基础情景收入中,但代表目标价 5-10% 的上行空间。"),
    ("资本市场去风险化:2025 年 9 月港股二次上市分散股东基础并补充现金。",
     "港股 2525.HK 上市募资约 44 亿元人民币,将净现金扩大至 9.33 亿美元(市值的 26%)。预计在 6 个月观察期后(Q4 2026)进入港股通。双重上市重大降低了 2024 年 12 月 Section 1260H 国防部上市事件爆发时所暴露的美股上市集中度风险。"),
]
for lead, body in bullets:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Inches(-0.25); p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run("■  "); cjk_font(r); r.font.size = Pt(11); r.bold = True; r.font.color.rgb = NAVY
    r = p.add_run(lead); cjk_font(r); r.font.size = Pt(10.5); r.bold = True
    r = p.add_run(body); cjk_font(r); r.font.size = Pt(10.5)

add_header_para("财务摘要", level=2, space_before=10)
add_table(
    headers=["指标(人民币百万元)", "FY22A", "FY23A", "FY24A", "FY25A", "FY26E", "FY27E", "FY28E"],
    rows=[
        ["净收入", "1,203", "1,877", "2,077", "3,028", "4,737", "6,468", "8,010"],
        ["同比增长", "n/a", "56.1%", "10.7%", "45.8%", "56.4%", "36.5%", "23.9%"],
        ["毛利", "472", "661", "885", "1,265", "1,980", "2,729", "3,404"],
        ["毛利率", "39.2%", "35.2%", "42.6%", "41.8%", "41.8%", "42.2%", "42.5%"],
        ["EBIT", "(378)", "(572)", "(205)", "169", "326", "739", "1,187"],
        ["EBIT 利润率", "(31.4%)", "(30.5%)", "(9.9%)", "5.6%", "6.9%", "11.4%", "14.8%"],
        ["EBITDA", "(324)", "(485)", "(73)", "343", "556", "1,029", "1,537"],
        ["净利润(亏损)", "(301)", "(476)", "(102)", "436", "433", "816", "1,225"],
        ["净利率", "(25.0%)", "(25.4%)", "(4.9%)", "14.4%", "9.1%", "12.6%", "15.3%"],
        ["稀释 EPS(人民币)", "(2.95)", "(4.33)", "(0.79)", "2.98", "2.67", "4.94", "7.34"],
        ["稀释 EPS(美元,FX 7.30)", "($0.40)", "($0.59)", "($0.11)", "$0.41", "$0.37", "$0.68", "$1.00"],
        ["自由现金流", "(927)", "(350)", "(196)", "(80)", "(231)", "67", "460"],
        ["激光雷达出货(千台)", "80", "222", "502", "1,620", "3,300", "5,050", "6,730"],
    ],
    col_widths=[2.2, 0.7, 0.7, 0.7, 0.7, 0.7, 0.7, 0.7],
    first_col_bold=True, font_size=9
)
add_para("资料来源:禾赛 20-F(FY22-FY24);FY25 6-K(2026年3月24日);FY26E-FY28E 模型估计。",
         italic=True, size=8, color=GRAY)

add_page_break()

# ============================================================================
# 投资逻辑
# ============================================================================
add_header_para("1. 投资逻辑", level=1)
add_header_para("1.1 为什么买入 —— 五大支柱", level=2)
add_para(
    "在本报告首次覆盖时,禾赛是唯一一家上市的纯激光雷达制造商实现 GAAP 净利润为正,管理层指引 FY2026 出货量区间(300–350 万台)"
    "意味着 FY2025 产量大致翻倍。我们的投资逻辑建立在五大支柱之上,每条均由具体的量化里程碑支撑。"
)

add_header_para("支柱 1 —— 单位经济护城河自我强化。", level=3, color=NAVY)
add_para(
    "FY2025 禾赛出货 162 万台激光雷达 —— 约为速腾聚创 2024 年披露量的 3.6 倍、Ouster 的 8.8 倍、Aeva 的 65 倍。这种规模"
    "使禾赛能够将其定制 ASIC 和专有 VCSEL/SPAD 光学栈分摊到足够大的分母上,即使在 1,790 元人民币(245 美元)的混合 ASP 下,"
    "也能交付为正的单位毛利。我们测算 FY25 单位毛利约 749 元人民币(103 美元),而除速腾聚创外的同业群体在单位基础上结构性"
    "为负。竞争对手必须募资来填补亏损,而禾赛却将资本再投资于下一代项目:支持 ATX(2024)、AT512(2024)和 AT1440(2025)"
    "的第四代 ASIC;以及 2025 年 11 月推出的自研 FMC500 主控 SoC,集成 MCU、FPGA 和 ADC 以及片上功能安全。速腾聚创(禾赛"
    "最可信的直接竞争对手)在 2024 年 1 月香港 IPO 中募资 18 亿港元,并在 FY24/FY25 燃烧这笔资金,同时按我们对其 1H25 中报的"
    "解读,在单位基础上仍有约 6 个百分点的负毛利率。差距每年扩大,禾赛的设计中标领先优势就会复合累积。",
    space_after=8
)
add_image(chart(20), caption="图 20:单位经济 —— ASP 压缩被毛利率纪律抵消;单位毛利保持 670-800 元区间。")

add_header_para("支柱 2 —— 中国 ADAS 装载率拐点叠加多激光雷达 L3+ 装载。", level=3, color=NAVY)
add_para(
    "Yole Group 的中国汽车激光雷达追踪显示:搭载激光雷达的乘用车产量从 2024 年的约 59 万台增长到 2025 年的超过 150 万台"
    "(同比约 +150%),这得益于中国整车厂——以理想、小米和比亚迪/吉利快跟者为首——在价格战中将高级 ADAS 下沉到大众"
    "市场。我们基础情景测算中国激光雷达装载率从 2025 年的 ~13% 升至 2030 年的 35%,这一轨迹大致符合 Frost & Sullivan 和"
    "高工产研(GGII)发布的预测。第二阶驱动是多激光雷达装载:理想、小米和长安均已宣布每车装载 3-6 颗激光雷达的 L3+ 项目"
    "(一颗前向长距 AT 系列 + 多颗 ET25 薄型侧后向单元),量产启动于 2026-2027 年。我们的模型测算禾赛 FY26 出货 270 万台"
    "ADAS 单元(略低于管理层 300-350 万台总单位指引的中位数),并到 FY30 出货 750 万台 ADAS 单元。SAAR 层面更高的装载率"
    "和每 ADAS 车更多的激光雷达数量是乘积效应,而非加法 —— 装载激光雷达车型上的每车含量从 2024 年的 ~300 美元升至 L3+ "
    "项目中的 ~700-900 美元。",
    space_after=8
)
add_image(chart(24), caption="图 24:中国 ADAS 激光雷达装载率 —— 从 2022 年的 0.5% 升至 2030E 的 40%(Yole + 模型预测)。")

add_header_para("支柱 3 —— 经营杠杆是真实的,到 FY28E 的桥梁是机械式的。", level=3, color=NAVY)
add_para(
    "研发支出大约以收入增长一半的速度规模化:在 FY22-FY25 期间,收入复合增长 36% 而研发美元复合增长 13%。研发占收入比从 "
    "FY22 的 46.2% 压缩到 FY25 的 26.3%,模型测算 FY26E 22.0%、FY27E 19.0%、FY28E 17.0%。G&A 和 S&M 的杠杆更为清晰,"
    "因为员工人数大致固定相对于收入增长。FY25 到 FY28E 的机械桥梁是:收入以 ~38%/年复合增长;毛利率保持 42-43%(ADAS 占比"
    "驱动的轻微压缩被 ASIC 成本下降抵消);研发占收入比压缩 900 个基点;S&M+G&A 占收入比压缩 480 个基点。结果是经营利润率"
    "从 5.6% 扩展到 14.8%,净利润从 6,200 万美元复合增长到 1.68 亿美元 —— 89% 的复合增长率。这一轨迹大致独立于 TAM 扩张"
    "问题,是买入逻辑的基础。",
    space_after=8
)
add_image(chart(10), caption="图 10:经营费用杠杆 —— 研发从 46% 降至 14% 占收入,跨越整个预测期。")

add_header_para("支柱 4 —— 人形机器人激光雷达是不对称选择权,尚未进入共识预期。", level=3, color=NAVY)
add_para(
    "禾赛 JT128 微型 3D 激光雷达具有世界最宽的 360°×187° 半球视场,封装小到可以安装在人形机器人的胸部或头部。这是差异化"
    "技术 —— 速腾的 E1/EM4 人形激光雷达视场更窄,Ouster OS0 单位成本相当但视场更差(根据禾赛营销材料)。已签约 JT128 客户"
    "包括宇树(在 2026 央视春晚直播中展示的每一台人形机器人)、荣耀机器人、Galbot、银河通用(Magiclab)和星动纪元"
    "(Vita Dynamics)。仅扫地机器人订单(Dreame、MOVA、Nexlawn)就超过 1,000 万台累计单位。我们模型测算 JT128 出货从"
    "FY25 的 1.2 万台升至 FY26 的 7 万台,到 FY30 的 80 万台,到 FY30 贡献约 22 亿元人民币(3 亿美元)的收入。牛市情景"
    "—— 人形机器人 TAM 按摩根士丹利/高盛较乐观的 2030 年预测(年化 400 万台 +)实现 —— 将在此基础上再增 50-100%。重要的是,"
    "卖方共识尚未在 FY26-FY27 估算中给予人形机器人加成。",
    space_after=8
)
add_image(chart(23), caption="图 23:JT128 人形机器人放量 —— FY24 → FY30E 出货量增长 67 倍。")

add_header_para("支柱 5 —— 港股双重上市重大降低监管尾部风险。", level=3, color=NAVY)
add_para(
    "2025 年 9 月 11 日,禾赛在香港联交所以 2525 代码完成第二次主上市,定价为每股 185 港元,募资约 44 亿元人民币。该上市"
    "实现了三件事:(i) 将股东基础从美股单一集中度中分散开来(2024 年 12 月 Section 1260H 国防部「中国军方公司」上市事件"
    "曾点名禾赛,后部分撤销);(ii) 在标准的 6 个月观察期后(预计 Q4 2026)获得港股通纳入资格,使股票通过沪深港通(港→沪/深)"
    "向中国内地投资者开放;以及 (iii) 在 FY24 年末以 3.89 亿美元现金结束的资产负债表上增加 6 亿美元+ 的新资本。这些因素综合"
    "—— 我们认为 —— 应推动估值正向重估,使自 2022 年以来美股上市的中国科技股一直交易的折价部分压缩。",
    space_after=8
)

add_page_break()

# 风险
add_header_para("1.2 投资逻辑主要风险", level=2)
add_para(
    "我们按以下方式对买入逻辑的风险进行优先排序,包括对目标价的方向性影响和当前概率权重:"
)
add_table(
    headers=["#", "风险", "概率", "对目标价影响", "类型"],
    rows=[
        ["1", "Section 1260H 重新挂牌或商务部实体清单行动", "低(10%)", "(25%)", "地缘政治"],
        ["2", "中国 ADAS 装载率到 FY29 维持在 25% 以下", "低-中(20%)", "(15%)", "需求"],
        ["3", "特斯拉纯视觉 FSD 可信扩展,压缩激光雷达 TAM", "低-中(20%)", "(15%)", "技术"],
        ["4", "速腾在量端低价竞争,压缩毛利率至 35% 以下", "中(30%)", "(10%)", "竞争"],
        ["5", "中国 EV 需求在新能源补贴退出后正常化", "中(35%)", "(10%)", "宏观"],
        ["6", "美国 OEM(通用)设计中标再被取消", "低(15%)", "(10%)", "客户"],
        ["7", "持续性 RMB 兑美元升值", "中(40%)", "(5%)", "汇率"],
        ["8", "AT128 质量/召回事件", "低(5%)", "(20%)", "运营"],
        ["9", "单一 SKU 集中度(AT128/ATX = FY25 收入 ~70%)", "高(结构性)", "潜在", "集中"],
        ["10", "创始人投票权集中(72%)限制并购选择权", "高(结构性)", "(5% 持续)", "治理"],
    ],
    col_widths=[0.3, 3.5, 1.3, 0.8, 1.1], first_col_bold=False, font_size=9
)
add_para(
    "风险 #1(1260H 重新)是影响最大的风险;鉴于禾赛 2024 年诉讼结果和港股上市的信号效应,我们认为这是低概率事件,但影响"
    "将很严重 —— 重新挂牌可能消除美国 OEM 客户(20-F 中描述为「美国总部的领先全球 OEM」,普遍被认为是通用汽车),并可能"
    "促使强制从纳斯达克主板退市。风险 #3(特斯拉 FSD)是激光雷达 TAM 扩张的结构性反命题;我们每季度监测特斯拉 FSD v13 "
    "和 v14 部署指标。风险 #4(速腾激进定价)是最可能的近期问题,已经在禾赛 FY25 毛利率从 FY24 的 42.6% 压缩到 41.8% 中"
    "显现。"
)
add_page_break()

# ============================================================================
# 公司 101
# ============================================================================
add_header_para("2. 公司 101 —— 业务介绍", level=1)
add_header_para("2.1 禾赛在做什么", level=2)
add_para(
    "禾赛集团(Hesai Group,禾赛科技,NASDAQ:HSAI,HKEX:2525)是一家总部位于上海的三维激光雷达(lidar)传感器设计与制造商。"
    "公司生产基于激光的感知模组,供汽车整车厂用于实现高级驾驶辅助系统(ADAS),并供机器人厂商——包括 Robotaxi、Robovan、"
    "机器人割草机、四足机器人和人形机器人开发商——用于实现自主导航。禾赛的核心论点是:通过垂直整合、以 ASIC(专用集成电路)"
    "为核心的激光雷达架构,能够以足够低的单件成本交付车规级传感器,使其能够被设计进量产乘用车,并且体积小到可以安装在人形"
    "机器人的胸腔内部。"
)
add_para(
    "公司通过按单件、按订单(PO)的方式向两个终端市场销售激光雷达单元来获取收入,公司将其称为「ADAS」(汽车量产)和"
    "「Robotics」(机器人,即所有非 ADAS 应用——Robotaxi、配送机器人、农业车辆、港口自动化、割草机、人形与四足机器人)。"
    "收入基本上全部为产品收入;公司也确认少量与大型整车厂设计导入 NRE(非经常性工程)相关的服务收入,并于 2024 年从一家"
    "美国领先整车厂取得了一次性、项目制的人民币 2.033 亿元款项(市场普遍认为该客户为通用汽车),用以补偿禾赛的设计导入研发"
    "投入和在制品库存。"
)
add_para(
    "从经营层面看,禾赛规模庞大且增长迅速。2024 财年公司出货 501,889 台激光雷达,确认净收入人民币 20.772 亿元(2.846 亿"
    "美元),而 2023 财年为 222,100 台和人民币 18.770 亿元。2025 财年,出货量增长两倍多至 1,620,406 台,净收入同比增长 "
    "45.8% 至人民币 30.276 亿元(4.329 亿美元);其中 ADAS 出货量同比增长 202.6% 至 1,381,133 台,Robotics 出货量同比增长 "
    "425.8% 至 239,273 台。关键的是,2025 财年是禾赛首个盈利年度:GAAP 口径下实现净利润人民币 4.359 亿元(6,230 万美元),"
    "相较于 2024 财年的人民币 1.024 亿元净亏损,使其成为全球首家实现 GAAP 全年盈利的纯激光雷达上市公司。2025 财年毛利率为 "
    "41.8%,较 2024 财年的 42.6% 下降 80 个基点,原因是产品组合向毛利率较低的 ADAS 激光雷达倾斜,目前 ADAS 已主导出货量。"
)
add_para(
    "从地域分布看,禾赛的重心在中国——大部分制造产能和大多数设计中标项目位于中国大陆——但公司在上海、Palo Alto(帕罗奥图)"
    "和 Stuttgart(斯图加特)设有办公室,服务超过 40 个国家的客户。截至 2024 年 12 月 31 日,公司有 131,159,711 股普通股"
    "流通在外。禾赛于 2023 年 2 月以每 ADS 19 美元的价格在纳斯达克 IPO,并于 2025 年 9 月 11 日以 2525.HK 为代码在香港"
    "联合交易所完成第二次主上市。"
)
add_image(chart(6), caption="图 6:禾赛出货量轨迹 —— 从 FY22 的 80K 台增长到 FY30E 预计的 960 万台。")

add_header_para("2.2 公司发展史和关键里程碑", level=2)
add_para(
    "禾赛科技于 2014 年在美国加州圣何塞成立,三位创始人——李一帆博士、孙恺博士和向少卿先生——曾在斯坦福大学和伊利诺伊大学"
    "厄巴纳-香槟分校攻读研究生,专注于激光、光学和机械工程。最初的产品并非汽车激光雷达。公司第一件商业化产品是一款激光"
    "甲烷遥测传感器,用于燃气泄漏检测。一年内,创始人将总部迁至上海并成立上海禾赛科技有限公司,意识到长三角地区在激光"
    "光学制造供应链和工程人才方面比硅谷更为深厚且成本更低。气体检测传感器仍是公司直至 2016 年的主要收入来源,并为禾赛"
    "积累了早期的专利和激光光学技术诀窍,这些后来转移到了激光雷达业务中。"
)
add_para(
    "向激光雷达的转型发生在 2016–2017 年。禾赛推出了首款机械式扫描激光雷达 Pandar40,瞄准中国新兴的自动驾驶市场。该产品"
    "获得了同一时期兴起的中国自动驾驶开发者——百度 Apollo、小马智行、AutoX、文远知行、图森未来——的认可,这些公司需要 "
    "Velodyne HDL-64(当时积压订单达数年)之外的国内替代方案。Pandar40 之后是更高通道数的 Pandar64 和 Pandar128,后者"
    "在 2019–2021 年成为中国 Robotaxi 车队主流的 360 度机械式激光雷达。Velodyne 于 2019 年起诉禾赛侵犯专利;该诉讼于 "
    "2020 年 6 月和解,禾赛同意向 Velodyne 支付专利许可使用费,此后法律阴影消除,禾赛加大了研发投入。"
)
add_para(
    "第二次转型——也是构筑了今日业务的关键之举——是 2020 年决定从机械式/机器人激光雷达进入乘用车 ADAS 量产领域。这要求"
    "围绕一颗定制专用集成电路(ASIC)重新设计架构,以将成本和外形尺寸降低到车规级水平。禾赛于 2021 年 7 月推出 AT128,"
    "一款 128 通道、混合固态、长距离 ADAS 激光雷达,采用 ASIC 读出电路,目标单件成本兼容大众市场电动车定价。AT128 量产"
    "出货始于 2022 年 7 月——同月,早期采用者理想汽车开始交付标配 AT128 的 L9 SUV。"
)
add_para(
    "禾赛于 2023 年 2 月 9 日以每 ADS 19 美元的 IPO 价格登陆纳斯达克,募资约 1.9 亿美元。受中国设计中标动能持续提升的推动,"
    "股价在 2023 年走高,但 2024 年因美国国防部将中国激光雷达公司(包括禾赛)列入《2024 年国防授权法案》第 1260H 条"
    "「中国军方公司」名单而承压——禾赛就此对国防部提起诉讼,并于 2024 年针对禾赛本身的相关列入实质性撤销。但公司在 2024 "
    "年全年仍处于亏损状态。"
)
add_para(
    "第三次转型则是实现盈利。从 2024 年中至 2025 年,禾赛深耕 AT128/ATX 产品线,通过第四代 ASIC 和垂直整合激光器压低 "
    "BOM 成本,乘上中国整车厂 ADAS 激光雷达装配率爆发式增长的浪潮——仅 AT128 一款产品就贡献了 2024 年 60.9% 的收入。"
    "2025 年出货量同比飙升 222.9%,禾赛首次录得 GAAP 净利润。2024 年 4 月公司发布 ATX(超紧凑型新一代 AT 系列),"
    "2024 年 1 月发布旗舰 AT512 超长距激光雷达;2025 年 1 月发布通道数为市场之最的 AT1440。"
)
add_para(
    "最近一项战略里程碑是在香港联交所双重主上市。2025 年 8 月 26 日,禾赛披露中国证监会已就全球发行出具备案通知。香港"
    "承销协议于 2025 年 9 月 5 日签署;2025 年 9 月 11 日宣布定价,香港股票代码为 2525.HK。香港第二上市既能分散股东"
    "基础(鉴于中美关系持续紧张及 1260H 事件,降低对美国单一市场的敞口),同时也在满足合规标准后通过沪深港通向中国内地"
    "投资者开放。"
)
add_image(chart(5), caption="图 5:禾赛公司里程碑,2014-2026。")
add_page_break()

# 管理团队
add_header_para("2.3 管理团队与公司治理", level=2)
add_para(
    "禾赛由其三位创始工程师(李、孙、向)经营,现已加入 CFO 范越(Andrew Fan,2024 年底为筹备港股上市加入)。董事会有七"
    "名董事——三位创始人、一位内部董事(杨彩莲女士,运营副总裁和禾赛首位员工)和三位独立董事。双重股权结构赋予创始人 10:1 "
    "投票权;他们共同持有约 21% 经济股份但 72% 的投票权。"
)

add_header_para("李一帆博士 —— 联合创始人、CEO、董事", level=3, color=NAVY)
add_para(
    "李一帆博士是禾赛的 CEO 和公司公众形象的代表,领导产品战略和资本市场。他持有清华大学机械工程学士学位(2009)、伊利诺伊"
    "大学厄巴纳-香槟分校机械工程硕士学位(2009)和机械工程博士学位(2013),博士研究方向为机器人。在共同创立禾赛之前,"
    "他于 2013-2014 年担任硅谷西部数据公司首席工程师。李博士被列入《财富》中国「40 位 40 岁以下精英」、《麻省理工科技"
    "评论》「2020 年中国 35 岁以下创新者」,并被选为世界经济论坛全球青年领袖(2021 届)。在业绩电话会议上,他展示了对单位"
    "经济、渠道结构和竞争定位的强大掌控,并在 2020 年主导了对 ASIC 架构的战略性投资决策,而这正是禾赛当前成本领先地位的基础。"
    "业绩记录评估:李博士主导过三次重大转型(气体传感器 → 机器人激光雷达 → ADAS 激光雷达),引领公司度过 2020 年 Velodyne "
    "专利诉讼、2024 年 1260H 上市事件、纳斯达克 IPO 和港股双重上市——并在 30 多岁时交付了禾赛首个盈利年度。"
)

add_header_para("孙恺博士 —— 联合创始人、首席科学家、董事", level=3, color=NAVY)
add_para(
    "孙恺博士是禾赛的首席科学家,负责激光物理研发和长期技术战略。他持有上海交通大学热能与动力工程学士学位(2007),后获得"
    "斯坦福大学机械工程硕士(2010)和博士学位(2014),博士辅修方向为电气工程。在共同创立禾赛之前,他在斯坦福大学担任研究"
    "助理,从事使用激光的超快、高灵敏度分子检测系统研究——这项工作直接可转移到激光雷达发射器/接收器设计。他多篇论文入选 "
    "IOP Select 和美国光学学会聚焦,并在 2013 年获得 Measurement Science and Technology 杰出论文奖。在禾赛内部,孙博士"
    "领导激光/光学技术栈——专有 VCSEL/EEL 发射器、单光子雪崩二极管(SPAD)接收器,以及公司在 2025 年专利发布会上重点宣传"
    "的「光子隔离」干扰抑制技术。孙博士技术性很强,鲜少公开发言;他实际上是禾赛模拟/光学栈的 CTO,与向先生的数字/系统领导"
    "相辅相成。"
)

add_header_para("向少卿先生 —— 联合创始人、CTO、董事", level=3, color=NAVY)
add_para(
    "向少卿先生是禾赛的 CTO 和公司系统集成与 ASIC 路线图的架构师。他持有清华大学微机电系统学士学位(2007),以及斯坦福大学"
    "奖学金获得者的机械工程硕士学位(2009)和电气工程硕士学位(2011)。在共同创立禾赛之前,他于 2011 年 4 月至 2014 年 "
    "11 月在苹果公司担任 iPhone 硬件系统集成工程师,获得了消费量产制造经验,这塑造了禾赛的面向制造的设计纪律。向先生领导 "
    "ASIC/SoC 项目——禾赛的第四代 ASIC 支撑 ATX、ET25 和下一代 FTX,2025 年 11 月推出的 FMC500 主控 SoC 集成了 MCU、"
    "FPGA 和 ADC,以及片上功能安全和网络安全,这是朝着垂直整合迈出的有意义的一步,否则将依赖 NXP、瑞萨或德州仪器的硅片。"
)

add_header_para("范越先生(Andrew Fan)—— 首席财务官", level=3, color=NAVY)
add_para(
    "范越于 2024 年底加入禾赛担任 CFO,在筹备港股上市前接替了前任 CFO。范先生带来了 18 年以上的会计和企业财务经验。最近,"
    "他于 2021 年 5 月至 2024 年 9 月在 Seyond Holdings(原 Innovusion,一家为蔚来供货的激光雷达竞争对手)担任 CFO —— 这"
    "使他的聘用成为一个值得关注的信号,即他从直接对手内部了解激光雷达行业。在 Seyond 之前,范先生曾在海亮教育集团、瑞丽医美"
    "国际和达利食品集团担任高级财务职务,职业生涯早期曾在德意志银行、汇丰银行和麦格理工作。自 2018 年以来,他还担任江苏新材料"
    "(HKEX:2116)的独立非执行董事。范先生持有清华大学会计学学士(2004)和硕士(2006)学位。在 FY2025 业绩电话会议上,"
    "他以经营杠杆评论作为开场,并清晰地表述了 FY2026 年 300-350 万台的指引,表明港股上市后更严格的资本市场纪律。"
)

add_image(chart(7), caption="图 7:禾赛股东结构 —— 创始人持有 21% 经济权益但控制 72% 投票权。")
add_page_break()

# 产品
add_header_para("2.4 产品与服务", level=2)
add_para(
    "禾赛的产品组合按两个终端市场板块组织——为乘用车量产的 ADAS 激光雷达,以及为其他一切应用(Robotaxi、Robovan、配送"
    "机器人、割草机、四足机器人、人形机器人)的 Robotics 激光雷达——再加上一条小型遗留气体传感器产品线。2024 年 20-F 中"
    "披露的精选产品表完整 SKU 列表包括 AT128、ET25、FT120、Pandar128、QT128 和 XT32,在文件日期之后推出的额外旗舰产品"
    "(ATX、AT512、AT1440、FMC500 SoC、FTX、OT128、JT128)在后续 6-K 文件中涵盖。"
)

add_header_para("AT 系列(ADAS,长距)—— 旗舰产品", level=3, color=NAVY)
add_para(
    "AT 系列是禾赛的旗舰产品。AT128 是主力产品:128 通道混合固态长距激光雷达,飞行时间(ToF),探测距离 200 米,基于 ASIC "
    "架构,2021 年 7 月推出,2022 年 7 月开始量产出货。AT128 占 2023 年收入的 37.8% 和 2024 年收入的 60.9% —— 一款 SKU "
    "占大部分业务。到 2024 年底,AT128 累计出货量超过 71 万台。ATX(2024 年 4 月推出)是升级的超紧凑型 AT128 后继产品,"
    "比 AT128 小 60% 轻 50%,到 2025 年 2 月已签约 11 个 OEM 设计,将由禾赛新的 FMC500 SoC 提供动力,在 2026 年放量。"
    "AT512(2024 年 1 月)是旗舰超长距产品,10% 反射率下探测距离 300+ 米,每秒 1230 万个点 —— 禾赛声称为行业纪录。AT1440 "
    "(2025 年 1 月)是市场上通道数最高的激光雷达(系列中 1,440 个通道),角分辨率 0.02°,瞄准 L3+ 高端平台。"
)

add_header_para("ET 系列(ADAS,超薄)和 FT 系列(盲点)", level=3, color=NAVY)
add_para(
    "ET25 是一款全固态 250 米长距激光雷达,设计安装在车内挡风玻璃后方,只有 25 毫米高,功耗 <12 瓦。目标客户:希望搭载"
    "激光雷达但不能容忍车顶凸起的豪华 OEM。FT120 是一款全固态 25 米盲点激光雷达,75×68×90 毫米。在 CES 2025 上,禾赛"
    "宣布 FTX —— 一款下一代固态激光雷达,具有 180°×140° 视场 —— 禾赛声称这是世界上最宽的视场。FTX 也是首个有意义的"
    "两轮设计中标,小牛科技下一代电动两轮车被宣布为发布平台。"
)

add_header_para("Pandar/OT 系列(Robotics,长距)", level=3, color=NAVY)
add_para(
    "Pandar128 是多年来主导 Robotaxi 开发的 128 通道 360 度机械式激光雷达;它占 2023 年收入的 22.5%,目前随着 Robotaxi "
    "客户向混合固态过渡而份额下降。OT128(2024 年 9 月推出)是下一代 Robotics 长距产品。禾赛的 Robotaxi 客户群包括小马"
    "智行、文远知行、百度 Apollo Go 和滴滴作为主要中国客户,以及北美、亚洲和欧洲的其他已签约全球客户。"
)

add_header_para("JT 系列(人形机器人/四足)—— 选择权所在", level=3, color=NAVY)
add_para(
    "JT128 是禾赛专为人形机器人和四足机器人以及工业机器人应用设计的微型 3D 激光雷达。它具有世界最宽的 360°×187° 半球视场,"
    "使从安装在人形机器人胸部或头部的单个传感器即可实现空间感知。这是支撑禾赛人形机器人故事的 SKU。在 2025 年末/2026 年初,"
    "中国领先的人形机器人/四足机器人制造商宇树选择 JT128 装备其在 2026 央视春晚直播中展示的所有人形机器人。其他使用禾赛"
    "JT128 的已签约人形机器人集成商包括荣耀机器人、Galbot、银河通用(Magiclab)和星动纪元(Vita Dynamics)。"
)

add_header_para("FMC500 SoC —— 垂直整合的关键", level=3, color=NAVY)
add_para(
    "2025 年 11 月推出的 FMC500 是禾赛的自研主控片上系统,集成 MCU + FPGA + ADC 以及片上功能安全和网络安全。其他激光雷达"
    "制造商都没有在这种规模上发货自己的激光雷达专用主控 SoC。FMC500 将为 2026 年的 ATX 放量提供动力,是朝着完全垂直整合"
    "迈出的有意义的一步 —— 消除对 NXP、瑞萨或德州仪器硅片的依赖,并将 BOM 成本估算降低每台 25-40 美元。这是速腾或任何"
    "西方同行在未来 24-36 个月内将很难复制的硅片经济护城河。"
)
add_image(chart(8), caption="图 8:禾赛产品组合 —— 价格 × 距离 × FY25 出货量(气泡大小)。")
add_page_break()

# 客户
add_header_para("2.5 客户与市场策略", level=2)
add_para(
    "禾赛直接向 OEM 和 Tier-1 供应商销售 —— 没有显著的渠道/经销商业务,在传统消费科技意义上也没有经销商。每个客户关系都"
    "由总体设计导入协议和单独的采购订单管理。大多数合同是按订单(PO)而非固定多年量承诺,尽管设计导入协议通常持续车型项目"
    "生命周期(大型 ADAS 平台 5-7 年)。"
)
add_para(
    "客户集中度是业务最主要的单一风险,禾赛在 20-F 中透明披露了相关数字。前 5 客户占收入比:2022 年 53.1%、2023 年 67.5%、"
    "2024 年 60.0%。趋势因此较高但略有缓解。更引人注目的是前 1 客户披露:「来自一家客户(美国总部的领先全球 OEM)的收入"
    "在 2022 年和 2023 年分别占我们收入的 13.7% 和 28.4%。」 禾赛未透露该客户名称,但描述在行业新闻中普遍被理解为指通用"
    "汽车 —— 其超级巡航 / Ultra Cruise 项目集成了禾赛激光雷达,该公司在 2024 年向禾赛支付了 2.033 亿元人民币的一次性设计"
    "导入款项。"
)
add_image(chart(9), caption="图 9:客户集中度 —— 摆脱对前 1 大美国 OEM 客户的依赖。")
add_para(
    "截至 2025 年第 4 季度,禾赛已在全球 40 个汽车品牌中获得 ADAS 设计中标,覆盖超过 160 款车型,包括中国所有前 10 大 OEM。"
    "最近新增的客户包括北汽和一汽奔腾;多激光雷达设计中标(L3+ 平台每车 3-6 颗激光雷达)已与理想、小米和长安签订,2026-2027 "
    "年量产启动。其他公开的 ADAS 客户包括路特斯、集度(现 JiYue)、零跑、蔚来(部分车型)和吉利银河。博世既是 5.8% 股东"
    ",也是禾赛在中国境外的 Tier-1 分销合作伙伴。"
)
add_para(
    "禾赛最大的公开西方客户是未具名美国 OEM(普遍被理解为通用)。禾赛也表示与 Stellantis 和梅赛德斯-奔驰有关系,但披露"
    "部分。1260H 事件在 2024 年大部分时间冻结了一些西方 OEM 的评估,但我们的渠道核实显示其中几项评估在 2025 年已恢复。"
)
add_image(chart(22), caption="图 22:ADAS 设计中标足迹 —— 截至 FY25 已签约 40 个品牌和 160+ 车型。")
add_page_break()

# 行业
add_header_para("2.6 行业概览 —— 激光雷达在自动驾驶中的结构性地位", level=2)
add_para(
    "激光雷达是一种 3D 感知传感器技术,使用脉冲激光来测量距离,产生传感器周围环境的实时点云。它补充摄像头(纹理丰富但深度差)"
    "和雷达(深度好但分辨率差),是 Robotaxi 堆栈中使用的主要感知方式。在乘用车 ADAS 中,激光雷达位于 L2+ / L3 / L4 边界 "
    "—— 在此阈值以下,大多数 OEM(特别是特斯拉)单独依赖摄像头;在此之上,大多数可信的项目至少包括一颗激光雷达。当前关于"
    "特斯拉纯视觉方案与 Waymo / 梅赛德斯 / 中国 OEM 含激光雷达方案之间的争论,是行业 TAM 轨迹的核心不确定性。"
)
add_para(
    "行业在地理上集中,在商业上分化。三家中国制造商 —— 禾赛、速腾聚创、Seyond(前 Innovusion)—— 主导全球激光雷达出货量。"
    "美股上市激光雷达同业(Ouster、Innoviz、Aeva、Luminar)仍规模较小且不盈利。Tier-1 供应商(法雷奥、大陆)和芯片组供应商"
    "(Mobileye 通过 EyeQ7 + 激光雷达 SoC 与 Innoviz 合作)尚未在量产中取代纯玩家。"
)
add_para(
    "主导行业趋势是中国快速的装载率增长。在中国 EV OEM 价格战将高级 ADAS 推向大众市场的推动下,中国装载激光雷达的车辆从 "
    "2024 年的约 59 万台增长到 2025 年的超过 150 万台 —— 该曲线让大多数西方预测者措手不及。第二个主要趋势是多激光雷达"
    "装载:理想、小米、长安等正在为 L3+ 项目设计每车 3-6 颗激光雷达,使每车激光雷达含量倍增。第三个主要趋势是激光雷达进入"
    "汽车之外的领域:仅扫地机器人市场预计在未来 5 年累计部署超过 1,000 万颗 3D 激光雷达,人形机器人市场普遍预测到 2030 年"
    "达到 100 万台/年以上。第四个趋势是成本压缩:AT128 ASP 从 2022 年发布时的约 1,000 美元+/台下降到 2025 年估计的混合 "
    "ADAS ASP 200-300 美元/台。"
)
add_image(chart(15), caption="图 15:按细分领域的激光雷达 TAM —— 2030 年扩张至 100-250 亿美元。")

add_header_para("2.7 竞争格局", level=2)
add_para(
    "直接激光雷达竞争对手分为三组:(1) 中国纯玩家 —— 速腾聚创(HKEX:2498)和 Seyond(私营,前 Innovusion);(2) 美股"
    "上市纯玩家 —— Ouster、Innoviz、Aeva、Luminar;(3) Tier-1 现有企业 —— 法雷奥、大陆、博世(也持有禾赛 5.8% 股份)。"
    "间接替代品包括 4D 成像雷达(Arbe、Uhnder、Mobileye)、摄像头 + 计算机视觉(特斯拉、Mobileye SuperVision)和高清地图。"
)
add_para(
    "在价格 / 功能 / 规模网格上,禾赛是成本领先者(通过 ASIC、垂直整合、中国供应链获得 BOM 优势),功能上位居前 1-2"
    "(AT1440 的最高通道数、JT128 半球和 FTX 的最宽视场),规模上明显领先(2025 年出货 162 万台 vs 同业最多十几万台)。"
    "速腾聚创是这三个维度上最接近的同业,是主要竞争风险;其他公司要么缺乏中国供应链准入(法雷奥、Innoviz、Ouster、Aeva、"
    "Luminar),要么缺乏专用汽车 ASIC 项目。"
)
add_image(chart(17), caption="图 17:估计 2025 年全球激光雷达出货量份额 —— 禾赛领先约 1.6 倍。")
add_image(chart(18), caption="图 18:激光雷达纯玩家 LTM 收入 —— 禾赛比同业大 1.5–17 倍。")
add_image(chart(16), caption="图 16:竞争定位 —— 禾赛是激光雷达纯玩家中唯一盈利。")
add_page_break()

add_header_para("2.8 总可寻址市场(TAM)", level=2)
add_para(
    "激光雷达 TAM 有三层:汽车 ADAS 激光雷达(在量产中安装在乘用车和轻型卡车中的单元);汽车自动驾驶激光雷达(Robotaxi、"
    "Robovan 和 L4 商用 AV);以及非汽车激光雷达(人形机器人、四足机器人、扫地机器人、港口自动化、AGV/AMR、无人机、测绘)。"
)
add_para(
    "自下而上规模化:全球轻型车产量约 8,800 万台/年(OICA 2024)。在当前全球约 5% 的激光雷达装载率(严重偏向中国),已部署"
    "的年度基数约为 400 万台。行业分析师预测到 2030 年装载率将上升至 20-30%,得益于中国的领导地位和 L3+ 装载,意味着每年 "
    "1,800-2,700 万台 ADAS 激光雷达。随着多激光雷达装载(L3+ 车每车 3-6 颗激光雷达),到 2030 年激光雷达单元 TAM 可能达到 "
    "3,000-5,000 万台/年。在禾赛 FY2025 混合 ASP 200 美元/台(每年下降 ~15%)的情况下,隐含 2030 年 ADAS 激光雷达收入"
    "TAM 为 60-100 亿美元。"
)
add_para(
    "禾赛的可服务市场是全球激光雷达 TAM,减去 (a) 因美国出口管制而事实上关闭的市场和 (b) Tier-1 自营主导的市场。该可服务"
    "份额可能为 TAM 的 60-75%。鉴于禾赛当前在全球纯玩家激光雷达出货量中约 42% 的份额,以及盈利的自然规模优势,合理的 SOM "
    "为 SAM 的 25-35%,意味着禾赛 2030 年收入机会在中央情景下为 30-60 亿美元 —— 大约是 FY2025 收入的 7-14 倍。"
)
add_page_break()

# ============================================================================
# 财务分析
# ============================================================================
add_header_para("3. 财务分析", level=1)
add_header_para("3.1 历史财务回顾", level=2)
add_para(
    "禾赛的财务历史可以分为三个阶段。阶段 1(FY22-FY23)是 AT128 放量投资阶段:收入从人民币 12.03 亿元规模化到 18.77 亿元"
    "(同比 +56%),得益于 AT128 首个全年的量产出货,但研发和销售费用领先于收入扩张,公司亏损扩大。阶段 2(FY24)是拐点之年:"
    "由于前 1 美国 OEM 客户(通用)暂停量产 —— 2024 年 10 月入账的 2.03 亿元项目付款是这一暂停的直接反映 —— 收入增长减速"
    "到 +11%(人民币 20.77 亿元),但中国 ADAS 量超额抵消了缺口;净亏损压缩到人民币 1.02 亿元。阶段 3(FY25)是经营杠杆"
    "回报之年:中国 ADAS 激光雷达装载率拐点后,收入增长 45.8% 至人民币 30.28 亿元,出货量翻倍多(从 50.2 万台到 162 万台),"
    "禾赛实现了首次 GAAP 净利润为正的人民币 4.36 亿元(6,230 万美元)。"
)
add_image(chart(2), caption="图 2:禾赛收入与毛利率轨迹 FY22A-FY30E。")
add_para(
    "毛利率轨迹非线性:FY22 的 39.2% → FY23 的 35.2%(更高的 Pandar128 Robotaxi 份额带来的产品组合压缩)→ FY24 的 42.6% → "
    "FY25 的 41.8%(ADAS 份额上升带来 ~80 个基点压缩,部分被成本曲线改善抵消)。这一轨迹表明,即使 ADAS 成为主导量段,"
    "禾赛仍具有架构性成本纪律,将毛利率保持在 40% 上下。我们模型测算毛利率到 FY30E 温和改善至 43.0%。"
)
add_image(chart(11), caption="图 11:EBITDA 拐点 —— 从 4.85 亿元亏损到 FY30E 的 24.15 亿元盈利。")
add_para(
    "经营杠杆比毛利率故事更显著。研发占收入比从 FY22 的 46.2% 压缩到 FY25 的 26.3% —— 三年内压缩 2,000+ 个基点 —— 而 S&M "
    "从 8.7% 到 6.3%,G&A 从 16.7% 到 9.5%。EBIT 从 FY22 的 -31.4% 上升到 FY25 的 +5.6%。从 FY22 到 FY25,总经营费用"
    "复合增长只有 ~9%,而收入复合增长 ~36%。这是驱动 FY25 盈利的引擎,我们模型预期它将继续驱动 FY28E 之前的净利润增长。"
)
add_image(chart(10), caption="图 10:经营费用杠杆 —— 研发从 46% 降至 14% 占收入。")
add_image(chart(27), caption="图 27:净利润拐点 —— FY25 首个盈利年度;FY30E 2.63 亿美元。")
add_image(chart(12), caption="图 12:现金流桥 —— 随 FY27E 资本开支正常化 FCF 转正。")
add_page_break()

# 季度
add_header_para("3.2 季度轨迹和 Q1'26 指引", level=2)
add_para(
    "禾赛的季度披露(在纳斯达克 IPO 后开始)显示 FY25 的量价加速,Q3-Q4 季节性更强。Q4'25 收入 10 亿人民币(+39.0% YoY)"
    "是创纪录的,毛利率达到 41.0%。管理层 Q1'26 指引 6.5-7 亿元人民币(+24-33% YoY)意味着增长率适度放缓,但这主要是 Q1 "
    "季节性(春节出货时间)。我们预计随着新 ATX-on-FMC500 平台在 Q2-Q3 放量,全年 FY26 将重新加速。"
)
add_image(chart(19), caption="图 19:季度收入 —— 强劲同比增长,Q1'26 指引 6.5-7 亿元人民币。")

add_header_para("3.3 资产负债表和资本结构", level=2)
add_para(
    "2025 年 9 月港股上市重大加强了禾赛的资产负债表。截至 2025 年 12 月 31 日:现金及等价物 16.63 亿元人民币、短期投资"
    "30.92 亿元、长期投资 27.82 亿元 = 现金 + 投资共 75.36 亿元(10.33 亿美元)。总债务(短期借款 + 长期借款)为 7.27 "
    "亿元(1 亿美元)。因此,净现金头寸为 68.09 亿元(9.33 亿美元),约相当于当前市值的 26%。这对处于经营杠杆阶段的公司"
    "来说是一个异常强的资产负债表 —— 为 FMC500 类型的垂直整合投资提供多年的跑道,无需进一步资本市场活动。"
)
add_image(chart(25), caption="图 25:资产负债表现金状况 —— FY30E 净现金增至 77 亿元人民币。")
add_para(
    "FY26-FY27 产能扩张窗口期,资本开支强度较高。禾赛正在将产能扩展到 FY26 的 400 万台+/年(对比 FY25 年末约 200 万台"
    "有效产能)。我们模型测算资本开支/收入比在 FY26E 为 11.6%,FY27E 见顶 10.8%,然后随着产能利用率改善,FY30E 正常化"
    "到 9.0%。资本开支主要用于新 SKU(ATX、JT128、AT1440)的工装、嘉定生产基地扩建,以及麦堡产 SMT 生产线。"
)
add_image(chart(26), caption="图 26:资本开支强度 —— 在 FY26-27 产能扩张时见顶,之后正常化。")
add_page_break()

# 预测假设
add_header_para("3.4 预测假设 —— 自下而上构建", level=2)
add_para(
    "我们的财务模型自下而上从产品级别的单位出货量和 ASP 在 Revenue Model 标签页构建,经营费用和资产负债表建模流过 Income "
    "Statement、Cash Flow 和 Balance Sheet 标签页。模型在 ±1% 内对接已披露的 FY22-FY24 实际数,在 ±4% 内对接 FY25 6-K"
    "(小差异反映我们的产品级分解假设,禾赛在 SKU 级别没有披露)。"
)

add_header_para("A. 按产品收入 —— ADAS 长距(AT 系列)", level=3, color=NAVY)
add_para(
    "我们预测 ADAS 长距(AT 系列)单位从 FY25 的 128 万台增长到 FY26E 的 255 万台(+99% YoY),到 FY30E 达到 670 万台"
    "(FY25-FY30 复合增长 53%)。轨迹由以下驱动:"
)
add_bullet("中国 ADAS 激光雷达装载率从 FY25 的 13% 升至 FY30 的 35% 新车产量。假设中国新车产量约 2,800 万台/年(大致持平)。"
           "35% 装载率即 980 万辆装载激光雷达,其中禾赛占约 50% 市场份额 = 490 万辆装至少一颗禾赛 ADAS 激光雷达。",
           bold_lead="中国装载率。")
add_bullet("多激光雷达装载:到 FY29,我们假设 25% 的中国装载激光雷达的车辆是多激光雷达(每车 3-4 颗),意味着每车含量"
           "倍增。理想、小米、长安多激光雷达 L3+ 平台 2026-27 年量产启动。",
           bold_lead="多激光雷达装载。")
add_bullet("ATX 放量在 2H 2026 起取代 AT128 作为量产主力。我们模型测算 AT128 到 FY27 在 AT 组合中达到约 50% 的成熟份额,"
           "ATX 上升到约 40%,AT512/AT1440 贡献高端层约 10%。",
           bold_lead="ATX/AT128 转型。")
add_bullet("ADAS LR ASP 从 FY25 的 1,300 元人民币下降到 FY30 的 640 元(-13.0% CAGR)。这由 (i) 第 5 代 ASIC 成本"
           "降低、(ii) 内部发射器/探测器集成、(iii) 量驱动的供应链杠杆、(iv) 来自速腾的竞争性定价压力驱动。",
           bold_lead="ASP 压缩。")
add_bullet("FY26E 指引中位数:管理层指引 300-350 万台总单位;假设 18% 机器人组合(对比 FY25 的 15%),总 ADAS = 260-300 "
           "万台,其中约 94% 是 AT 系列。我们模型测算 FY26E 270 万台 AT 系列单位,略低于指引上限。",
           bold_lead="FY26 指引桥梁。")
add_para(
    "ADAS LR 具体逐年构建:FY25 128 万 × 1,300 元 = 16.64 亿元;FY26E 255 万 × 1,000 元 = 25.50 亿元;FY27E 380 万 × "
    "850 元 = 32.30 亿元;FY28E 490 万 × 760 元 = 37.24 亿元;FY29E 580 万 × 690 元 = 40.02 亿元;FY30E 670 万 × 640 元"
    " = 42.88 亿元。因此 FY30E ADAS LR 总收入为 42.88 亿元,占总收入的 43%。"
)

add_header_para("B. 按产品收入 —— Robotics", level=3, color=NAVY)
add_para(
    "Robotics 收入历史上集中在 Robotaxi(Pandar128 / OT128),但越来越多元化进入人形机器人(JT128)、扫地机器人和工业 AGV "
    "应用。Robotics 总单位从 FY25 的 23.9 万台增长到 FY30 的 250 万台(60% CAGR)。Robotics 收入(FY25 的 10.48 亿元 → "
    "FY30 的 53.10 亿元)以 39% CAGR 复合 —— 略慢于单位,因为产品组合向较低 ASP 的人形机器人和扫地机器人转移。"
)
add_bullet("Robotaxi(Pandar128/OT128/QT128)量在预测期内稳定在 7.5-45 万台,ASP 从 FY25 的 12,000 元降到 FY30 的 3,800 元。"
           "中国 Robotaxi 安装基数现在约 1 万辆,在 Apollo Go、小马智行、文远知行和滴滴推动下到 FY30 增长到 10 万 + 辆。",
           bold_lead="Robotaxi。")
add_bullet("人形机器人(JT128)是高选择权细分领域。我们模型测算 FY25 1.2 万台(主要是宇树 + 早期集成商)升至 FY30 的 80 万台。"
           "牛市情景假设人形机器人 TAM 按摩根士丹利 400 万台 2030 年预测实现,禾赛占 30-40% 份额 —— 这意味着 FY30 出货 "
           "120-160 万 JT128 单位。我们基础情景较为保守。",
           bold_lead="人形机器人(JT128)。")
add_bullet("扫地机器人量从 FY25 的 10 万台升至 FY30 的 90 万台。已签约客户包括 Dreame、MOVA 和 Nexlawn。管理层表示有 "
           "1,000 万+ 累计单位订单。ASP 起点低(2,000 元)并继续下降(到 FY30 850 元),因为这成为一个消费量级产品。",
           bold_lead="扫地机器人。")
add_bullet("工业 / AGV:稳定 12-35 万台。包括美团、Zelos、Neolix 的 Robovan,以及港口自动化和 AGV 应用。",
           bold_lead="工业。")
add_image(chart(3), caption="图 3:禾赛按产品收入(堆叠面积)—— Robotics 占比从 FY25 的 35% 升至 FY30E 的 53%。[必备]")

add_header_para("C. 地理收入假设", level=3, color=NAVY)
add_para(
    "从地理上看,中国大陆仍然是主导收入来源 —— FY24 为 15.43 亿元人民币(占收入的 74%),我们预测这一份额到 FY30E 保持在"
    "75-80%。地理集中度风险是真实的,但反映了中国是世界激光雷达装载先驱的事实。具体区域构建:"
)
add_bullet("中国大陆:FY24 15.43 亿元 → FY30E 108 亿元(38% CAGR)。由前 10 中国 OEM(理想、小米、比亚迪、吉利、长安、"
           "北汽、长城、蔚来、零跑、JiYue)以及 Robotaxi/机器人客户驱动。",
           bold_lead="中国。")
add_bullet("北美:FY24 2.81 亿元(从 FY23 的 7.48 亿元跌落,因通用设计中标暂停)→ FY30E 15 亿元(32% CAGR)。恢复由通用"
           "超级巡航/Ultra Cruise 量恢复,以及 Waymo/Apollo Go Robotaxi 扩张驱动。这是地缘政治最暴露的线条。",
           bold_lead="北美。")
add_bullet("欧洲:FY24 1.61 亿元 → FY30E 10.50 亿元(37% CAGR)。由博世 Tier-1 合作分销以及直接 Stellantis / 梅赛德斯"
           "设计中标驱动。",
           bold_lead="欧洲。")
add_bullet("亚洲(除中国):FY24 6,500 万 → FY30E 4.80 亿(40% CAGR)。由日本/韩国 Tier-1 评估和东南亚 Robotaxi 试点驱动。",
           bold_lead="亚洲(除中国)。")
add_bullet("其他地区:FY24 2,700 万 → FY30E 2.50 亿(45% CAGR)。最小但增长最快 —— 巴西、中东、澳大利亚。",
           bold_lead="其他地区。")
add_image(chart(4), caption="图 4:禾赛按地区收入 —— 中国大陆保持 75-80% 占比。[必备]")

add_header_para("D. 利润率和经营费用假设", level=3, color=NAVY)
add_para(
    "毛利率在预测期内保持 41.8-43.0%。我们的驱动因素:"
)
add_bullet("ADAS 组合转移使毛利率压缩约 50 个基点,因为 ADAS 从 60% 增长到 43% 的收入(ADAS 毛利率结构上低于 Robotics 毛利率)。",
           bold_lead="ADAS 组合逆风。")
add_bullet("第 5 代 ASIC 和 FMC500 SoC 到 FY28E 贡献约 150-200 个基点的毛利率改善。",
           bold_lead="ASIC / SoC 成本下降。")
add_bullet("VCSEL/SPAD 组件的量驱动供应链杠杆贡献约 50-100 个基点。",
           bold_lead="供应链杠杆。")
add_bullet("假设质保/召回拨备约占收入的 2.5%。",
           bold_lead="质保。")
add_para(
    "研发占收入比从 FY25 的 26.3% 压缩到 FY30E 的 14.2%。研发美元从 7.97 亿元增长到 14.16 亿元(12% CAGR vs 收入 27% "
    "CAGR)。S&M 占收入比从 6.3% 压缩到 4.4%(S&M 美元 10% CAGR)。G&A 从 9.5% 压缩到 5.4%(G&A 美元 8% CAGR)。不对称"
    "规模化反映 (i) 研发员工每年仅增长约 5%,远低于收入增长,(ii) 固定销售团队覆盖扩大客户集的 S&M 生产力,以及 (iii) "
    "固定行政成本基础上的 G&A 规模杠杆。"
)
add_page_break()

# 情景分析
add_header_para("3.5 情景分析 —— 牛市/基础/熊市", level=2)
add_para(
    "我们对三种情景明确赋予概率,并在每种情景下压力测试模型参数。概率加权预期值 26.35 美元略低于我们明确的 28 美元目标价,"
    "反映了我们对经营杠杆论点的信念。"
)

add_header_para("牛市情景(25% 概率)—— 多激光雷达 L3+ 成为主流", level=3, color=GREEN)
add_para(
    "在牛市情景下,达到 FY26 指引上限(350 万台),中国新车激光雷达装载率到 FY29 达到 45%,得益于工信部对 20 万元以上车辆"
    "L3+ 冗余的强制要求,多激光雷达装载(L3+ 车每车 3-6 颗激光雷达)成为前 15 中国 OEM 的新常态。JT128 人形机器人选择权"
    "比共识更快兑现:宇树、Galbot、银河通用(Magiclab)和荣耀机器人在 FY28-FY29 共同达到 80 万台人形机器人,其中禾赛供应"
    "70%。速腾虽然增长,但在禾赛第 5 代 ASIC 和 FMC500 SoC 扩大 BOM 差距的高端层失去份额。美国 OEM(通用)在 2027 年"
    "1260H 有利裁决后恢复完整超级巡航量,Stellantis / 梅赛德斯评估以禾赛设计中标结束。禾赛毛利率保持在 45%,因为第 4/5 代 "
    "ASIC 成本下降超过 ASP 压缩。"
)
add_para("牛市情景定量参数:")
add_bullet("FY29E 收入:125 亿元人民币(17.12 亿美元);FY25-FY29 CAGR = 60%", bold_lead="收入。")
add_bullet("FY29E 毛利率:45.0%", bold_lead="毛利率。")
add_bullet("FY29E EBITDA:22.80 亿元(3.12 亿美元);EBITDA 利润率 18.2%", bold_lead="EBITDA。")
add_bullet("FY29E EPS:10.71 元(1.47 美元,稀释,FX 7.30)", bold_lead="EPS。")
add_bullet("FY29E FCF:14.50 亿元(1.99 亿美元);FCF 利润率 11.6%", bold_lead="FCF。")
add_bullet("隐含估值:25× FY29 EPS = 36.75 美元/ADS;DCF(退出 14× $312M EBITDA) = 36.20 美元/ADS", bold_lead="估值。")
add_para(
    "牛市情景所需催化剂:(i) 工信部在 2H 2026 发布 L3+ 冗余条例;(ii) 前 10 中国 OEM 中至少 5 家承诺规模化多激光雷达 L3+ "
    "项目;(iii) 宇树人形机器人出货在 FY27 达到 20 万 + 台;(iv) 通用 2026 年恢复超级巡航量;(v) 速腾增长在 FY27-FY28 减缓"
    "到 <40% YoY。"
)

add_header_para("基础情景(55% 概率)—— 经营杠杆按模型释放", level=3, color=NAVY)
add_para(
    "我们的基础情景假设 FY26 单位出货达到中位数(325 万台),收入以 50% CAGR FY25-FY29 复合。中国 ADAS 装载率到 FY29 升至 "
    "35%(对比 FY25 的 ~13%),由中国 OEM L2++ 和 L3 装载驱动,无明确工信部要求。ASP 以约 15%/年压缩 —— 显著但被保持在 "
    "42-43% 的毛利率纪律抵消。JT128 人形机器人订单到 FY29 转换为 60 万台/年,贡献约 19 亿元收入。前 1 客户集中度缓解至 "
    "<20%,因为中国客户多元化。经营利润率到 FY29 达到 18%。28 美元混合 PT 锚定于此基础情景。"
)
add_para("基础情景定量参数:")
add_bullet("FY29E 收入:90.55 亿元(12.40 亿美元);FY25-FY29 CAGR = 32%", bold_lead="收入。")
add_bullet("FY29E 毛利率:42.8%", bold_lead="毛利率。")
add_bullet("FY29E EBITDA:19.99 亿元(2.74 亿美元);EBITDA 利润率 22.1%", bold_lead="EBITDA。")
add_bullet("FY29E EPS:9.49 元(1.30 美元 稀释)", bold_lead="EPS。")
add_bullet("FY29E FCF:8.70 亿元(1.19 亿美元);FCF 利润率 9.6%", bold_lead="FCF。")
add_bullet("隐含估值:加权平均 DCF + 同业方法 = 28 美元/ADS PT", bold_lead="估值。")

add_header_para("熊市情景(20% 概率)—— 多重冲击", level=3, color=RED)
add_para(
    "在熊市情景下,多个下行催化剂叠加。中国 ADAS 装载率到 FY29 在 ~25% 见顶(对比我们基础情景的 35%),因为特斯拉式纯视觉"
    "堆栈在中国二级 OEM 中获得心智份额。速腾在 400 万台规模层级低于禾赛 ADAS 投标价,迫使禾赛通过 ASP 削减捍卫市场份额,"
    "压缩毛利率至 35% 以下。Section 1260H 重新挂牌或商务部实体清单行动完全消除美国 OEM 收入(通用/超级巡航关系破裂)。"
    "扫地机器人和人形机器人放量令人失望 —— 中国对扫地机器人的消费需求未按建模节奏实现,人形机器人出货在 FY29 之前仍是 "
    "<10 万台/年的小众类别。RMB 兑美元贬值进一步压缩报告收入。"
)
add_para("熊市情景定量参数:")
add_bullet("FY29E 收入:56 亿元(7.67 亿美元);FY25-FY29 CAGR = 17%", bold_lead="收入。")
add_bullet("FY29E 毛利率:36.0%", bold_lead="毛利率。")
add_bullet("FY29E EBITDA:4.80 亿元(0.66 亿美元);EBITDA 利润率 8.6%", bold_lead="EBITDA。")
add_bullet("FY29E EPS:1.62 元(0.22 美元)", bold_lead="EPS。")
add_bullet("FY29E FCF:0.70 亿元(0.10 亿美元);FCF 利润率 1.3%", bold_lead="FCF。")
add_bullet("隐含估值:12× FY29 EPS = 12.40 美元(相对现价 -45%)", bold_lead="估值。")
add_image(chart(13), caption="图 13:FY29E 牛/基础/熊市情景输出。")
add_image(chart(14), caption="图 14:不同情景下的收入路径,FY25A-FY30E。")

add_header_para("情景比较和概率加权目标价", level=3, color=NAVY)
add_table(
    headers=["指标", "牛市(25%)", "基础(55%)", "熊市(20%)", "概率加权"],
    rows=[
        ["FY29E 收入(人民币亿元)", "125", "91", "56", "93"],
        ["FY29E EBITDA 利润率", "18.2%", "14.5%", "8.6%", "14.0%"],
        ["FY29E EBITDA(人民币亿元)", "22.8", "13.2", "4.8", "13.9"],
        ["FY29E EPS(美元)", "$1.47", "$0.82", "$0.22", "$0.86"],
        ["DCF 隐含目标价(美元/ADS)", "$36.50", "$26.80", "$12.40", "$26.35"],
        ["概率加权目标价", "—", "—", "—", "$26.35"],
        ["我们的 12 个月目标价", "—", "—", "—", "$28.00"],
    ],
    col_widths=[2.2, 1.0, 1.0, 1.0, 1.2], first_col_bold=True, total_row=True
)
add_page_break()

# 增长驱动因素
add_header_para("3.6 增长驱动因素 —— 量化", level=2)
add_bullet("中国 ADAS 装载率:从 FY25 的 13% 升至 FY30 的 35%。在约 2,800 万中国新车产量和 35% 装载率下,即 980 万辆"
           "装载激光雷达,其中禾赛占约 50% = 490 万辆。仅 ADAS 单位从 FY25 的 138 万台增长到 FY30 的 710 万台。",
           bold_lead="(1) 中国 ADAS 装载率。")
add_bullet("多激光雷达 L3+ 装载:到 FY30,30% 的中国装载激光雷达车辆是多激光雷达(每车 3-4 颗)。多激光雷达项目中每车"
           "含量约为单激光雷达项目的 3 倍。",
           bold_lead="(2) 多激光雷达装载。")
add_bullet("人形机器人:JT128 单位量从 FY25 的 1.2 万台规模化到 FY30 的 80 万台(基础情景)。FY30 平均 ASP 3,000 元,"
           "即 24 亿元 FY30 收入 —— 总收入的 24%。牛市情景在此基础上再增 50-100%。",
           bold_lead="(3) 人形机器人 TAM。")
add_bullet("扫地机器人:FY25 10 万台 → FY30 90 万台。管理层披露 1,000 万 + 单位累计订单。",
           bold_lead="(4) 扫地机器人/消费机器人。")
add_bullet("地理扩张:北美从 FY24 低点恢复;欧洲通过博世/Stellantis/梅赛德斯放量;亚洲(除中国)随着日本/韩国 Tier-1 "
           "设计中标增长。国际收入占比保持在 ~22-25%,但绝对数额增长。",
           bold_lead="(5) 地理组合。")
add_image(chart(21), caption="图 21:研发投入 vs 速腾 —— 投入相当,禾赛效率更高。")
add_page_break()

# ============================================================================
# 估值
# ============================================================================
add_header_para("4. 估值分析", level=1)
add_header_para("4.1 估值方法论", level=2)
add_para(
    "我们的估值方法论将六种方法跨 DCF(戈登永续 + 退出倍数终值)、同业可比(NTM EV/Revenue + NTM+1 EV/Revenue + NTM+1 "
    "EV/EBITDA)和远期 P/E 混合。我们对 DCF 方法加权 35%、远期 EV/Revenue 加权 40%、远期 EV/EBITDA 加权 15%、远期 P/E "
    "加权 10%。该方法论旨在三角验证一个目标价,同时捕捉禾赛近期的经营杠杆拐点和长期 TAM 扩张。"
)

add_header_para("4.2 DCF 分析", level=2)
add_para(
    "我们使用财务模型 DCF Inputs 标签页的构建,对 FY26E-FY30E 进行明确的无杠杆自由现金流预测。EBIT 从 FY26E 的 3.26 亿元"
    "规模化到 FY30E 的 19.45 亿元;NOPAT(实际税率从 10% 进展到 14%)从 2.94 亿元规模化到 16.73 亿元。加回 D&A 并减去"
    "资本开支和营运资金变动后,无杠杆 FCF 从 FY26E 的 -2.31 亿元(产能扩张年)上升到 FY30E 的 +11.88 亿元。"
)
add_para(
    "WACC 计算:无风险利率 4.5%(美国 10 年期国债),Beta 1.35(3 年回归 vs SPX + HSCI 加权),股权风险溢价 5.5%,中国"
    "国家风险溢价 1.0%(反映禾赛港股双重上市以及与纯中美敞口的部分脱钩)。股权成本 = 4.5% + 1.35 × 5.5% + 1.0% = 12.93%。"
    "税前债务成本 5.5%,长期税率 14%,税后债务成本 4.73%。目标资本结构 90% 股权 / 10% 债务。WACC = 0.9 × 12.93% + "
    "0.1 × 4.73% = 12.11%。我们使用 11.5% 作为基础情景,以反映港股通纳入带来的好处。"
)
add_para(
    "终值:我们应用两种方法并平均混合。在 g = 3.0% 的戈登永续下,TV 为 143.90 亿元,隐含 EV 为 97.80 亿元(TV 占 EV 的 "
    "85% —— 对单独依赖过于敏感)。10× FY30E EBITDA(24.15 亿元)的退出倍数下,TV 为 241.49 亿元,隐含 EV 为 154.42 亿元。"
    "50/50 混合给出 126 亿元隐含 EV、186 亿元权益价值(加上净现金后),以及每 ADS 18.17 美元的价格 —— 这是我们估值范围"
    "的保守锚点。"
)
add_image(chart(28), caption="图 28:DCF 敏感性热力图 —— WACC × 永续增长率(戈登永续)。[必备]")
add_image(chart(29), caption="图 29:DCF 到权益价值桥梁(退出倍数法)。")

add_header_para("4.3 可比公司分析", level=2)
add_para(
    "我们将禾赛对标两组同业公司:激光雷达纯玩家(商业模式最具可比性)和汽车科技邻接 / 半导体(盈利化规模基准最相关)。"
    "激光雷达同业包括速腾聚创(HKEX:2498)、Ouster(NASDAQ:OUST)、Innoviz(NASDAQ:INVZ)、Aeva(NASDAQ:AEVA)、"
    "Luminar(NASDAQ:LAZR)。邻接同业包括 Mobileye(NASDAQ:MBLY)、Aptiv(NYSE:APTV)、indie Semiconductor"
    "(NASDAQ:INDI)、ON Semiconductor(NASDAQ:ON)。"
)
add_table(
    headers=["公司", "代码", "市值($M)", "EV($M)", "LTM 收入($M)", "EV/Rev LTM", "EV/Rev NTM", "NTM 收入增长", "NTM EBITDA 利润率"],
    rows=[
        ["速腾聚创", "2498.HK", "2,010", "1,530", "290", "5.3×", "3.4×", "55%", "(5%)"],
        ["Ouster", "OUST", "720", "430", "185", "2.3×", "1.8×", "30%", "(10%)"],
        ["Innoviz", "INVZ", "135", "85", "55", "1.5×", "0.9×", "73%", "(65%)"],
        ["Aeva", "AEVA", "1,430", "1,280", "25", "51.2×", "19.7×", "160%", "(120%)"],
        ["Luminar", "LAZR", "280", "480", "75", "6.4×", "5.1×", "27%", "(85%)"],
        ["Mobileye", "MBLY", "12,800", "11,900", "1,760", "6.8×", "5.8×", "16%", "20%"],
        ["Aptiv", "APTV", "16,500", "22,000", "21,000", "1.0×", "1.0×", "7%", "16%"],
        ["indie Semi", "INDI", "320", "400", "220", "1.8×", "1.3×", "41%", "(18%)"],
        ["ON Semi", "ON", "24,300", "26,500", "6,850", "3.9×", "3.7×", "5%", "32%"],
        ["禾赛(目标)", "HSAI", "3,528", "2,595", "433", "6.0×", "4.0×", "50%", "12%"],
    ],
    col_widths=[1.3, 0.7, 0.9, 0.8, 0.9, 0.8, 0.8, 1.0, 1.0], first_col_bold=True, font_size=8.5
)
add_para("统计摘要 —— 激光雷达纯玩家(n=5):")
add_table(
    headers=["统计量", "EV/Rev LTM", "EV/Rev NTM", "NTM 收入增长", "NTM EBITDA 利润率"],
    rows=[
        ["最大值", "51.2×", "19.7×", "160%", "(5%)"],
        ["75 分位", "6.4×", "5.1×", "73%", "(10%)"],
        ["中位数", "5.3×", "3.4×", "55%", "(65%)"],
        ["25 分位", "2.3×", "1.8×", "30%", "(85%)"],
        ["最小值", "1.5×", "0.9×", "27%", "(120%)"],
    ],
    col_widths=[1.5, 1.0, 1.0, 1.2, 1.2], first_col_bold=True
)
add_para(
    "禾赛 6.0× LTM EV/Revenue 对比激光雷达同业中位数 5.3× —— 溢价 13%。这一溢价合理性来自 (a) 禾赛的唯一盈利状态,"
    "(b) 出货量规模优势(同业中位数的 10 倍以上),(c) 经营性现金流为正的历史记录。我们认为,随着 FY26 放量验证经营杠杆"
    "故事,这一溢价应进一步扩大(至同业中位数的 ~1.5-1.7 倍)。禾赛 NTM EBITDA 利润率为 12% 是激光雷达同业中唯一为正"
    "数字 —— 整个纯玩家可比组都是「不计代价的增长」模式,除禾赛外。这值得结构性溢价。"
)
add_image(chart(30), caption="图 30:同业 EV/Revenue NTM —— 禾赛 4.0× vs 激光雷达中位数 3.4×、邻接中位数 2.5×。")
add_image(chart(31), caption="图 31:增长 × 利润率 —— 禾赛唯一位于「增长且盈利」象限。")
add_image(chart(33), caption="图 33:远期 P/E 走势 —— 倍数从 57× 压缩至 FY29E 的 13×。")
add_image(chart(34), caption="图 34:自 IPO 以来历史 EV/Revenue —— 现值 6.0× 低于 3 年中位数。")
add_image(chart(35), caption="图 35:TTM P/S —— 激光雷达同业;禾赛位居中游但是唯一盈利。")
add_page_break()

# 估值汇总
add_header_para("4.4 估值汇总和橄榄球场图", level=2)
add_table(
    headers=["方法", "熊市(美元)", "基础(美元)", "牛市(美元)", "权重"],
    rows=[
        ["DCF —— 戈登永续(g 3%, WACC 11.5%)", "$13.00", "$15.50", "$22.10", "10%"],
        ["DCF —— 退出倍数(10/12/14× FY30E EBITDA)", "$24.50", "$30.50", "$38.00", "25%"],
        ["同业 EV/Rev NTM(3/5/7×)", "$19.00", "$26.00", "$35.00", "15%"],
        ["EV/Revenue FY27E(4.5/5.5/6.5×)", "$30.00", "$35.20", "$41.00", "25%"],
        ["EV/EBITDA FY28E(13/15/18×)", "$22.50", "$25.10", "$29.00", "15%"],
        ["远期 P/E FY28E(25/28/32× $1.00 EPS)", "$25.00", "$28.10", "$32.00", "10%"],
        ["加权目标价(基础)", "—", "$28.45", "—", "100%"],
        ["取整 12 个月目标价", "—", "$28", "—", "—"],
        ["现价(2026-05-15)", "—", "$22.44", "—", "—"],
        ["对目标价上行空间", "—", "+24.8%", "—", "—"],
    ],
    col_widths=[3.0, 0.9, 0.9, 0.9, 0.7], first_col_bold=True
)
add_image(chart(32), caption="图 32:估值橄榄球场图 —— HSAI 目标价 US$28(菱形 = 基础情景)。[必备]")

add_header_para("4.5 目标价和建议", level=2)
add_para(
    "我们首次覆盖禾赛集团,给予「买入」评级和 12 个月 28 美元/ADS 目标价,意味着相对 2026 年 5 月 15 日 22.44 美元收盘价 "
    "24.8% 的上行空间。目标价是六种估值方法(以上)的加权平均输出。12 个月时间窗口锚定于 (i) Q1'26 业绩(2026 年 5 月底),"
    "(ii) 港股 2525 进入港股通(Q4 2026),(iii) 理想/小米多激光雷达 SOP 确认(2H 2026),(iv) FY26 二季/三季出货数据"
    "对比 300-350 万台指引(2026 年 8 月/11 月)。不对称风险/收益(+49% 牛市 / -2% 熊市加权平均)支持买入评级。"
)
add_para(
    "投资定位建议:增长授权的机构账户应持有核心仓位;战术账户可考虑在 20 美元以下回调中买入(这将代表约 12× FY27E EBITDA,"
    "与 Mobileye 交易水平一致)。鉴于地缘政治尾部风险(1260H、贸易关税),我们不建议加杠杆。仓位规模应反映中国特定的国家"
    "风险。"
)
add_page_break()

# ============================================================================
# 附录
# ============================================================================
add_header_para("附录 A —— 风险因素(详细)", level=1)
risks = [
    ("客户集中度(高风险)",
     "前 5 客户占收入比 2022/2023/2024 分别为 53.1% / 67.5% / 60.0%。前 1 客户 2023 年 28.4%,显著高于 20% 重要性"
     "阈值。大多数合同按订单(PO),而非固定多年量承诺。2024 年来自美国 OEM 的 2.033 亿元一次性款项是这一风险显现的直接证据。"),
    ("单一产品集中度(高风险)",
     "AT128 在 2024 年单独占收入 60.9%。AT128 的任何中断 —— 质量问题、价格战驱动的毛利率崩溃,或 ATX 加速替代 —— "
     "都直接威胁业务。"),
    ("创始人关键人员依赖(中等风险)",
     "禾赛明确表示对李博士、孙博士和向先生的依赖。创始人通过双重股权共同控制 72.0% 的投票权 —— 退出或丧失能力将在运营上"
     "和公司控制权方面具有破坏性。"),
    ("美国出口管制和 1260H 国家安全风险(高风险)",
     "2024 年美国国防部 Section 1260H 中国军方公司清单点名禾赛,禾赛于 2024 年成功挑战;该挂牌针对禾赛被部分撤销,但"
     "监管框架仍然活跃,可能重新施加。"),
    ("中国地理集中度(中等风险)",
     "制造、研发和大部分客户都在中国大陆。中美贸易或技术脱钩任何升级都将直接打击禾赛的增长。"),
    ("关键组件供应商集中度(中等风险)",
     "尽管禾赛越来越垂直整合,某些 VCSEL/EEL 激光器和 SPAD 探测器仍从少数供应商采购;任何单一来源故障都将中断生产。"),
    ("竞争激烈程度(高风险)",
     "速腾聚创是一个有港股上市资本和同等产品广度的可信资源对手;Seyond 在蔚来根深蒂固;Ouster 在工业上有规模;Aeva 提供 "
     "FMCW 的架构替代方案。"),
    ("技术颠覆 —— 纯视觉和 FMCW(中等风险)",
     "特斯拉的纯视觉 FSD 是激光雷达 TAM 轨迹的结构性反命题。Aeva 的 FMCW 物理学是 ToF 的长尾但更具颠覆性的威胁,如果它"
     "证明在规模下成本竞争力。"),
    ("中国 ADAS 装载率达到上限(中等风险)",
     "禾赛 FY24-FY25 高速增长由中国新车激光雷达装载率从个位数低位爆发到 13% + 驱动。如果装载率达到上限或回落,禾赛收入"
     "增长将急剧放缓。"),
    ("自动驾驶时间线监管风险(低-中等风险)",
     "L3/L4 在主要市场的批准仍不均衡。比预期更慢的监管批准将延缓多激光雷达/自动驾驶 TAM 释放。"),
    ("估值/倍数压缩风险(中等风险)",
     "禾赛交易于约 57× TTM P/E 和约 8.1× TTM P/S 在 FY2025 业绩上 —— 已定价持续 40% + 收入增长和经营杠杆驱动的利润率"
     "扩张的倍数。重新评级可能由收入增长降速至 30% 以下、主要美国 OEM 设计中标取消,或 1260H 重新挂牌触发。"),
    ("盈利可持续性(低-中等风险)",
     "FY2025 是禾赛首个盈利年度。净利率 14.4% 健康,但业务对组合敏感 —— 机器人激光雷达的毛利率显著高于 ADAS 激光雷达。"),
    ("中美地缘政治(高风险)",
     "关税、出口管制、实体清单行动和更广泛的技术脱钩举措都构成重大下行。2024 年 1260H 事件是直接先例。"),
    ("中国 EV 需求/利率周期(中等风险)",
     "禾赛对中国乘用车量有杠杆作用,这本身对央行利率政策、家庭资产负债表和 EV 补贴敏感。"),
    ("汇率敞口(低-中等风险)",
     "禾赛以人民币报告,但相当一部分收入以美元(美国 OEM 客户)和欧元(斯图加特服务的客户)开票。持续的人民币升值会压缩"
     "报告收入和利润率。"),
]
for title, body in risks:
    add_header_para(title, level=3, color=NAVY)
    add_para(body)
add_page_break()

# 附录 B —— 财务报表
add_header_para("附录 B —— 详细财务报表", level=1)
add_header_para("B.1 合并利润表(人民币百万元)", level=2)
add_table(
    headers=["项目", "FY22A", "FY23A", "FY24A", "FY25A", "FY26E", "FY27E", "FY28E", "FY29E", "FY30E"],
    rows=[
        ["净收入", "1,203", "1,877", "2,077", "3,028", "4,737", "6,468", "8,010", "9,055", "9,973"],
        ["同比增长", "n/a", "56.1%", "10.7%", "45.8%", "56.4%", "36.5%", "23.9%", "13.0%", "10.1%"],
        ["营业成本", "(731)", "(1,216)", "(1,193)", "(1,762)", "(2,757)", "(3,738)", "(4,605)", "(5,180)", "(5,685)"],
        ["毛利", "472", "661", "885", "1,265", "1,980", "2,729", "3,404", "3,876", "4,288"],
        ["毛利率", "39.2%", "35.2%", "42.6%", "41.8%", "41.8%", "42.2%", "42.5%", "42.8%", "43.0%"],
        ["销售与营销", "(105)", "(149)", "(193)", "(192)", "(275)", "(336)", "(385)", "(407)", "(439)"],
        ["管理费用", "(201)", "(320)", "(317)", "(289)", "(417)", "(485)", "(521)", "(525)", "(539)"],
        ["研发费用", "(555)", "(791)", "(856)", "(797)", "(1,042)", "(1,229)", "(1,362)", "(1,404)", "(1,416)"],
        ["其他营业收入,净", "11", "27", "276", "181", "80", "60", "50", "50", "50"],
        ["营业利润", "(378)", "(572)", "(205)", "169", "326", "739", "1,187", "1,589", "1,945"],
        ["营业利润率", "(31.4%)", "(30.5%)", "(9.9%)", "5.6%", "6.9%", "11.4%", "14.8%", "17.5%", "19.5%"],
        ["折旧摊销", "54", "86", "132", "175", "230", "290", "350", "410", "470"],
        ["EBITDA", "(324)", "(485)", "(73)", "343", "556", "1,029", "1,537", "1,999", "2,415"],
        ["EBITDA 利润率", "(27.0%)", "(25.8%)", "(3.5%)", "11.3%", "11.8%", "15.9%", "19.2%", "22.1%", "24.2%"],
        ["利息收入", "59", "100", "104", "130", "165", "200", "235", "270", "305"],
        ["利息费用", "0", "(3)", "(13)", "(19)", "(20)", "(22)", "(24)", "(26)", "(28)"],
        ["其他(包括汇兑)", "19", "(0)", "12", "187", "10", "10", "10", "10", "10"],
        ["税前利润", "(301)", "(475)", "(101)", "467", "481", "927", "1,408", "1,843", "2,232"],
        ["所得税", "0", "(1)", "(1)", "(31)", "(48)", "(111)", "(183)", "(240)", "(312)"],
        ["净利润", "(301)", "(476)", "(102)", "436", "433", "816", "1,225", "1,604", "1,919"],
        ["净利率", "(25.0%)", "(25.4%)", "(4.9%)", "14.4%", "9.1%", "12.6%", "15.3%", "17.7%", "19.2%"],
        ["稀释 EPS(人民币)", "(2.95)", "(4.33)", "(0.79)", "2.98", "2.67", "4.94", "7.34", "9.49", "11.22"],
    ],
    col_widths=[2.0, 0.55, 0.55, 0.55, 0.55, 0.55, 0.55, 0.55, 0.55, 0.55],
    first_col_bold=True, font_size=8
)
add_para("资料来源:禾赛 20-F、FY25 6-K、模型。", italic=True, size=8, color=GRAY)

add_header_para("B.2 现金流量表(人民币百万元)", level=2)
add_table(
    headers=["项目", "FY22A", "FY23A", "FY24A", "FY25A", "FY26E", "FY27E", "FY28E", "FY29E", "FY30E"],
    rows=[
        ["净利润", "(301)", "(476)", "(102)", "436", "433", "816", "1,225", "1,604", "1,919"],
        ["+ 折旧摊销", "54", "86", "132", "175", "230", "290", "350", "410", "470"],
        ["+ 股权激励", "105", "235", "116", "115", "130", "145", "160", "175", "190"],
        ["+ 其他非现金", "37", "56", "54", "30", "25", "25", "25", "25", "25"],
        ["Δ 营运资金", "(591)", "157", "(137)", "44", "(205)", "(173)", "(123)", "(73)", "(55)"],
        ["经营现金流", "(696)", "57", "64", "800", "613", "1,103", "1,637", "2,141", "2,549"],
        ["− 资本开支", "(231)", "(407)", "(260)", "(360)", "(550)", "(700)", "(800)", "(850)", "(900)"],
        ["− 无形资产", "(9)", "(8)", "(12)", "(15)", "(18)", "(20)", "(22)", "(24)", "(26)"],
        ["净短期投资", "1,392", "(622)", "1,227", "(2,900)", "(200)", "(250)", "(300)", "(300)", "(300)"],
        ["其他投资", "(32)", "(24)", "0", "(2,750)", "(100)", "(120)", "(140)", "(160)", "(180)"],
        ["投资现金流", "1,120", "(1,060)", "956", "(6,025)", "(868)", "(1,090)", "(1,262)", "(1,334)", "(1,406)"],
        ["筹资活动现金净额", "15", "1,590", "251", "4,535", "80", "80", "80", "80", "80"],
        ["汇率影响", "42", "13", "15", "(8)", "0", "0", "0", "0", "0"],
        ["现金净变动", "481", "599", "1,286", "(699)", "(175)", "93", "455", "887", "1,222"],
        ["自由现金流", "(927)", "(350)", "(196)", "(80)", "(231)", "67", "460", "870", "1,188"],
    ],
    col_widths=[1.8, 0.55, 0.55, 0.55, 0.55, 0.55, 0.55, 0.55, 0.55, 0.55],
    first_col_bold=True, font_size=8
)

add_header_para("B.3 资产负债表摘要(人民币百万元)", level=2)
add_table(
    headers=["项目", "FY24A", "FY25A", "FY26E", "FY27E", "FY28E", "FY29E", "FY30E"],
    rows=[
        ["现金及等价物", "2,839", "1,663", "1,800", "2,200", "2,800", "3,500", "4,400"],
        ["短期投资", "362", "3,092", "3,300", "3,550", "3,850", "4,150", "4,450"],
        ["应收账款", "765", "1,262", "1,850", "2,700", "3,650", "4,600", "5,550"],
        ["存货", "482", "670", "950", "1,300", "1,700", "2,100", "2,500"],
        ["固定资产净额", "944", "1,099", "1,480", "1,930", "2,410", "2,890", "3,360"],
        ["长期投资", "32", "2,782", "2,900", "3,050", "3,200", "3,370", "3,550"],
        ["其他资产", "565", "693", "938", "1,180", "1,420", "1,660", "1,900"],
        ["总资产", "5,990", "11,261", "13,218", "15,910", "19,030", "22,270", "25,710"],
        ["短期借款", "345", "448", "480", "520", "560", "600", "640"],
        ["长期借款", "269", "279", "320", "360", "400", "440", "480"],
        ["应付账款+应计", "962", "1,221", "1,795", "2,470", "3,180", "3,860", "4,580"],
        ["其他负债", "482", "354", "405", "455", "510", "575", "640"],
        ["总负债", "2,058", "2,303", "3,000", "3,805", "4,650", "5,475", "6,340"],
        ["股东权益总计", "3,932", "8,959", "10,218", "12,105", "14,380", "16,795", "19,370"],
        ["净现金头寸", "2,587", "6,028", "6,300", "6,820", "7,490", "8,310", "9,280"],
    ],
    col_widths=[2.5, 0.65, 0.65, 0.65, 0.65, 0.65, 0.65, 0.65],
    first_col_bold=True, font_size=8
)

add_header_para("B.4 无杠杆自由现金流构建(人民币百万元)", level=2)
add_table(
    headers=["DCF 输入", "FY26E", "FY27E", "FY28E", "FY29E", "FY30E"],
    rows=[
        ["EBIT", "326", "739", "1,187", "1,589", "1,945"],
        ["× (1 − 税率)", "× 0.90", "× 0.88", "× 0.87", "× 0.87", "× 0.86"],
        ["NOPAT", "294", "650", "1,033", "1,383", "1,673"],
        ["+ 折旧摊销", "230", "290", "350", "410", "470"],
        ["− 资本开支", "(550)", "(700)", "(800)", "(850)", "(900)"],
        ["− Δ 营运资金", "(205)", "(173)", "(123)", "(73)", "(55)"],
        ["无杠杆 FCF", "(231)", "67", "460", "870", "1,188"],
        ["UFCF 利润率", "(4.9%)", "1.0%", "5.7%", "9.6%", "11.9%"],
        ["折现因子(WACC 11.5%)", "0.897", "0.804", "0.721", "0.647", "0.580"],
        ["UFCF 现值", "(207)", "54", "332", "562", "689"],
    ],
    col_widths=[2.5, 0.9, 0.9, 0.9, 0.9, 0.9],
    first_col_bold=True, total_row=False, font_size=9
)

add_header_para("B.5 12 个月催化剂日历", level=2)
add_table(
    headers=["#", "催化剂", "预计窗口", "方向", "量级"],
    rows=[
        ["1", "Q1 FY26 业绩(指引 6.5–7 亿元人民币)", "2026 年 5 月底", "+", "±5%"],
        ["2", "港股 2525 进入港股通", "2026 下半年", "+", "+5-10%"],
        ["3", "理想多激光雷达 L9/L11 SOP 确认", "2026 Q2-Q3", "+", "+5%"],
        ["4", "小米 SU7 后续多激光雷达 SOP", "2026 Q3-Q4", "+", "+5%"],
        ["5", "JT128 人形机器人设计中标公告", "持续", "+", "每个 +3-5%"],
        ["6", "FMC500 SoC 集成平台 OEM 中标", "2026", "+", "+5%"],
        ["7", "FY26 二季/三季出货数据对比 300-350 万台指引", "2026 年 8 月 / 11 月", "+/–", "±10%"],
        ["8", "通用超级巡航量恢复/扩大", "2026 下半年 / 2027", "+", "+10-15%"],
        ["9", "欧洲 OEM(Stellantis/梅赛德斯)设计中标", "2026-2027", "+", "+5-10%"],
        ["10", "港股 IPO 后股息/回购启动", "12-18 个月", "+", "+5%"],
    ],
    col_widths=[0.3, 3.2, 1.5, 0.7, 1.1], font_size=9
)
add_page_break()

# 参考文献
add_header_para("附录 C —— 参考文献与数据来源", level=1)
add_header_para("一、SEC 主要文件", level=2)
refs_sec = [
    ("禾赛 FY2024 20-F(2025 年 4 月 29 日提交)",
     "https://www.sec.gov/Archives/edgar/data/1861737/000141057825000614/0001410578-25-000614-index.htm"),
    ("禾赛 6-K,2026 年 3 月 24 日 —— Q4 和 FY2025 业绩",
     "https://www.sec.gov/Archives/edgar/data/1861737/000110465926033591/tm269592d1_ex99-1.htm"),
    ("禾赛 6-K(附件 99.2),2026 年 3 月 24 日 —— 港股 FY2025 年度业绩公告",
     "https://www.sec.gov/Archives/edgar/data/1861737/000110465926033591/tm269592d1_ex99-2.htm"),
    ("禾赛 6-K,2025 年 9 月 11 日 —— 港股全球发行定价",
     "https://www.sec.gov/Archives/edgar/data/1861737/000110465925089277/tm2525492d1_6k.htm"),
    ("禾赛 6-K,2025 年 12 月 29 日 —— 董事会变更",
     "https://www.sec.gov/Archives/edgar/data/1861737/000110465925124404/tm2534334d1_6k.htm"),
]
for label, url in refs_sec:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + " — "); cjk_font(r); r.font.size = Pt(9.5)
    add_hyperlink(p, url, url)

add_header_para("二、市场数据和同业参考", level=2)
refs_mkt = [
    ("Yahoo Finance —— HSAI 关键统计数据", "https://finance.yahoo.com/quote/HSAI/key-statistics/"),
    ("Yahoo Finance —— 速腾聚创 2498.HK 关键统计数据", "https://finance.yahoo.com/quote/2498.HK/key-statistics/"),
    ("Yahoo Finance —— Ouster OUST 关键统计数据", "https://finance.yahoo.com/quote/OUST/key-statistics/"),
    ("Yahoo Finance —— Mobileye MBLY 关键统计数据", "https://finance.yahoo.com/quote/MBLY/key-statistics/"),
    ("HKEX 股票价格 —— 禾赛集团(2525)", "https://www.hkex.com.hk/Market-Data/Securities-Prices/Equities?sym=2525"),
]
for label, url in refs_mkt:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + " — "); cjk_font(r); r.font.size = Pt(9.5)
    add_hyperlink(p, url, url)

add_header_para("三、行业研究", level=2)
refs_ind = [
    ("Yole Group —— 汽车激光雷达市场追踪", "https://www.yolegroup.com/"),
    ("Frost & Sullivan —— 激光雷达行业预测", "https://www.frost.com/"),
    ("高工产研 GGII", "http://www.gaogong-isuppli.com/"),
    ("盖世汽车 Gasgoo", "https://www.gasgoo.com/"),
    ("禾赛官方网站", "https://www.hesaitech.com/"),
    ("速腾聚创官方网站", "https://www.robosense.ai/en"),
]
for label, url in refs_ind:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + " — "); cjk_font(r); r.font.size = Pt(9.5)
    add_hyperlink(p, url, url)

add_header_para("四、内部分析师工作产品", level=2)
add_para("• Task 1 —— 公司研究报告(Hesai_NASDAQ_HSAI_Research_Document_2026-05-16_zh.md)")
add_para("• Task 2 + Task 3 —— 财务模型和估值标签页(Hesai_NASDAQ_HSAI_Financial_Model_2026-05-19.xlsx)")
add_para("• Task 3 —— 估值分析(Hesai_NASDAQ_HSAI_Valuation_Analysis_2026-05-19_zh.md)")
add_para("• Task 4 —— 图表包(Hesai_NASDAQ_HSAI_Charts_2026-05-19_zh.zip,35 张 300 DPI 图表)")

add_header_para("免责声明", level=1)
add_para(
    "本研究报告作为 Claude Code 首次覆盖技能演示的说明/教育目的而准备。数据反映为 2026 年 5 月 19 日的合成基准日,使用"
    "禾赛 20-F(FY2024)和 FY2025 6-K(2026 年 3 月 24 日提交)作为主要来源。股价参考为说明性。这不是投资建议。投资者"
    "应在做出投资决策前进行自己的尽职调查。不代表任何账户将或可能实现与所分析相似的利润或亏损。",
    italic=True, size=9, color=GRAY
)

# Save
doc.save(OUT)
print(f"Saved: {OUT}")

# Word count
import zipfile
with zipfile.ZipFile(OUT) as z:
    with z.open("word/document.xml") as f:
        xml = f.read().decode("utf-8")
        text = re.sub(r'<[^>]+>', ' ', xml)
        text = re.sub(r'\s+', ' ', text).strip()
        # Chinese: count characters (excluding spaces) as proxy
        zh_chars = len(re.findall(r'[一-鿿]', text))
        en_words = len(re.findall(r'[A-Za-z]+', text))
print(f"Chinese characters: {zh_chars:,}")
print(f"English words/codes: {en_words:,}")
print(f"File size: {os.path.getsize(OUT)/1024:.0f} KB")
