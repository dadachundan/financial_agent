"""
Append Q1 2026 results addendum to both EN and ZH DOCX reports.
Q1 2026 actual numbers from Hesai press release May 19, 2026.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x00, 0x33, 0x66)
ACCENT = RGBColor(0xFF, 0xA5, 0x00)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xC6, 0x28, 0x28)
GRAY = RGBColor(0x66, 0x66, 0x66)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def cjk_font(run, name=None):
    if name is None:
        return
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)


def append_addendum(docx_path, lang='en'):
    doc = Document(docx_path)
    font_name = "Times New Roman" if lang == 'en' else "Songti SC"

    def set_font(run, size=10.5, bold=False, italic=False, color=None):
        run.font.name = font_name
        run.font.size = Pt(size)
        if bold: run.bold = True
        if italic: run.italic = True
        if color: run.font.color.rgb = color
        if lang == 'zh':
            cjk_font(run, font_name)

    # Page break to start addendum
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    # Title
    if lang == 'en':
        title_text = "ADDENDUM — Q1 2026 RESULTS UPDATE"
        subtitle_text = "Added May 19, 2026 (after market) following Hesai's Q1 2026 press release."
    else:
        title_text = "附录 —— Q1 2026 业绩更新"
        subtitle_text = "2026 年 5 月 19 日(盘后)追加;禾赛 Q1 2026 业绩公告后更新。"

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title_text); set_font(r, size=18, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(subtitle_text); set_font(r, size=10, italic=True, color=GRAY)

    # PT reaffirm box
    pt_t = doc.add_table(rows=2, cols=4)
    pt_t.alignment = WD_TABLE_ALIGNMENT.LEFT
    if lang == 'en':
        hdrs = ["RATING (REAFFIRMED)", "12M PRICE TARGET", "Q1 2026 ACTUAL", "Q1 vs GUIDE"]
        vals = ["BUY", "US$28 (unchanged)", "RMB 680.6M (+29.6%)", "Mid of 650-700M"]
    else:
        hdrs = ["评级(维持)", "12 个月目标价", "Q1 2026 实际", "Q1 vs 指引"]
        vals = ["买入", "US$28(不变)", "RMB 680.6M(+29.6%)", "650-700M 中位数"]

    for i, h in enumerate(hdrs):
        cell = pt_t.rows[0].cells[i]; cell.text = ""
        para = cell.paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(h); set_font(r, size=8, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, "002060")
    for i, v in enumerate(vals):
        cell = pt_t.rows[1].cells[i]; cell.text = ""
        para = cell.paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(v); set_font(r, size=11, bold=True)
        if i == 0: r.font.color.rgb = GREEN
        shade_cell(cell, "F2F2F2")
    for row in pt_t.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for bn in ['top','left','bottom','right']:
                b = OxmlElement(f'w:{bn}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'8'); b.set(qn('w:color'),'002060')
                tcBorders.append(b)
            tcPr.append(tcBorders)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)

    # Section 1: Actuals vs model
    h1_text = "Q1 2026 Actuals vs. Our Model" if lang == 'en' else "Q1 2026 实际数 vs. 我们的模型"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(h1_text); set_font(r, size=13, bold=True, color=NAVY)

    # Comparison table
    if lang == 'en':
        cmp_hdr = ["Metric", "Q1 2026 Actual", "Mgmt Guide", "Our FY26E Base", "Read"]
        cmp_rows = [
            ["Net revenue (RMB M)", "680.6 (+29.6%)", "650-700", "n/a quarterly", "Mid of guide"],
            ["Gross margin", "39.1%", "n/a", "41.8%", "270bps BELOW model"],
            ["GAAP net income (RMB M)", "18.3", "n/a", "433 FY", "4 consec. profit Qs"],
            ["Non-GAAP NI (RMB M)", "47.7 (+452.9%)", "n/a", "n/a", "Strong"],
            ["Total units shipped", "471,723 (+140.9%)", "n/a", "3,300K FY", "Q1 = 14% of FY26E"],
            ["ADAS units", "353,441 (+141.9%)", "n/a", "2,700K FY", "On pace"],
            ["Robotics units", "118,282 (+137.8%)", "n/a", "600K FY", "Ahead — higher GM mix"],
            ["Robotics % of units", "25%", "n/a", "18% FY", "Mix shift = GM tailwind in 2H"],
        ]
    else:
        cmp_hdr = ["指标", "Q1 2026 实际", "管理层指引", "我们 FY26E 基础", "解读"]
        cmp_rows = [
            ["净收入(人民币百万元)", "680.6(+29.6%)", "650-700", "不适用(全年)", "指引中位数"],
            ["毛利率", "39.1%", "不适用", "41.8%", "较模型低 270 bps"],
            ["GAAP 净利润(人民币百万元)", "18.3", "不适用", "433(全年)", "连续 4 个盈利季"],
            ["非 GAAP 净利润(人民币百万元)", "47.7(+452.9%)", "不适用", "不适用", "强劲"],
            ["总出货量", "471,723(+140.9%)", "不适用", "3,300K 全年", "Q1 = FY26E 的 14%"],
            ["ADAS 出货", "353,441(+141.9%)", "不适用", "2,700K 全年", "按节奏推进"],
            ["Robotics 出货", "118,282(+137.8%)", "不适用", "600K 全年", "超前 —— 毛利率组合更佳"],
            ["Robotics 占单位比", "25%", "不适用", "18%(全年)", "组合转移 = 2H 毛利率顺风"],
        ]

    t = doc.add_table(rows=1 + len(cmp_rows), cols=5)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [1.8, 1.4, 1.0, 1.3, 1.5]
    for i, w in enumerate(widths):
        for cell in t.columns[i].cells:
            cell.width = Inches(w)
    for i, h in enumerate(cmp_hdr):
        cell = t.rows[0].cells[i]; cell.text = ""
        para = cell.paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(h); set_font(r, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, "002060")
    for ri, row in enumerate(cmp_rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]; cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = para.add_run(str(val)); set_font(r, size=9, bold=(ci == 0))
            if ri % 2 == 1: shade_cell(cell, "F2F2F2")
    for row in t.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for bn in ['top','left','bottom','right']:
                b = OxmlElement(f'w:{bn}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:color'),'888888')
                tcBorders.append(b)
            tcPr.append(tcBorders)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 2: New disclosures
    h2_text = "New Disclosures (Catalysts Triggered)" if lang == 'en' else "新披露(催化剂触发)"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(h2_text); set_font(r, size=13, bold=True, color=NAVY)

    if lang == 'en':
        bullets = [
            ("Mercedes L3 design-in CONFIRMED.",
             " Hesai officially announced as strategic LiDAR supplier for Mercedes L3 autonomous models in both Europe AND China. This was speculation in our initiation (Section 2.5) and is now hard news — supports our base-case European revenue trajectory (RMB 1,050M FY30E)."),
            ("Galileo manufacturing facility in Thailand.",
             " Hesai's first major offshore production capacity. This materially de-risks risk #2 (Section 1260H national-security risk) from our Section 1.2 risk matrix, because it provides a non-China-origin lidar supply path for US OEMs."),
            ("New \"Strategic Growth Initiatives\" (SGI) segment.",
             " Hesai introduced a new financial reporting segment covering the Picasso 6D full-color chip and Kosmo spatial intelligence device (launched April 2026). FY26 SGI contribution guided at ~RMB 100M, beginning Q2. Not currently in our model — represents incremental upside."),
            ("ETX LiDAR mass production expected H2 2026.",
             " Additional new SKU not in our model."),
            ("Strategic pivot from \"spatial perception\" to \"spatial intelligence.\"",
             " Language pivot suggests Hesai is positioning as a higher-margin systems play rather than commodity sensor vendor."),
        ]
    else:
        bullets = [
            ("梅赛德斯 L3 设计中标已正式确认。",
             " 禾赛正式宣布为梅赛德斯欧洲与中国 L3 自动驾驶车型的战略激光雷达供应商。这在我们首次覆盖(2.5 节)中是猜测,现在是硬新闻 —— 支持我们基础情景下的欧洲收入轨迹(FY30E 10.50 亿元)。"),
            ("泰国 Galileo 制造基地。",
             " 禾赛首个大型海外生产基地。这实质性降低我们 1.2 节风险矩阵中的风险 #2(Section 1260H 国家安全风险),因为它为美国 OEM 提供了非中国原产地的激光雷达供应路径。"),
            ("新「战略增长计划」(SGI)分部。",
             " 禾赛新设的财务报告分部,涵盖 Picasso 6D 全彩芯片和 Kosmo 空间智能设备(2026 年 4 月发布)。指引 FY26 SGI 贡献约 1 亿元,从 Q2 开始。目前未纳入我们的模型 —— 代表增量上行空间。"),
            ("ETX 激光雷达预计 2H 2026 量产。",
             " 模型中未纳入的额外新 SKU。"),
            ("战略转型从「空间感知」到「空间智能」。",
             " 语言转变表明禾赛正在定位为更高利润率的系统玩家,而非商品化传感器供应商。"),
        ]
    for lead, body in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(lead); set_font(r, size=10.5, bold=True)
        r = p.add_run(body); set_font(r, size=10.5)

    # Section 3: PT implications
    h3_text = "Impact on Price Target" if lang == 'en' else "对目标价的影响"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(h3_text); set_font(r, size=13, bold=True, color=NAVY)

    if lang == 'en':
        impact_para = (
            "PT REAFFIRMED at US$28 / ADS · BUY. The Q1 print is broadly in line with our base-case "
            "trajectory: revenue at the guide midpoint, units +141% YoY validating the operating leverage thesis. "
            "Mercedes L3 + Thailand facility = positive (catalysts triggered; risk #2 partially de-risked). "
            "Q1 gross margin at 39.1% = small negative — we revise our FY26E gross margin assumption from 41.8% "
            "to ~40.5% (-130bps). Pass-through impact: ~-RMB 60M to FY26E EBIT, ~-2 cents to FY26E EPS, ~-$1 "
            "to per-ADS DCF value. Net: PT stays at US$28 (within model precision). "
            "Robotics units +137.8% with 25% share of units is ahead of our model (we had 18% for full year FY26). "
            "If sustained, this is a tailwind to blended GM in 2H."
        )
    else:
        impact_para = (
            "目标价维持 US$28 / ADS · 买入。Q1 业绩与我们基础情景轨迹大致一致:收入处于指引中位数,出货量同比 +141% "
            "验证经营杠杆论点。梅赛德斯 L3 + 泰国基地 = 正面(催化剂已触发;风险 #2 部分去风险化)。"
            "Q1 毛利率 39.1% = 小负面 —— 我们将 FY26E 毛利率假设从 41.8% 下调至约 40.5%(-130 bps)。"
            "透过影响:FY26E EBIT 约 -6,000 万元,FY26E EPS 约 -2 美分,每 ADS DCF 价值约 -1 美元。"
            "净:目标价保持 US$28(在模型精度范围内)。Robotics 出货 +137.8% 占单位 25% 领先于我们模型"
            "(我们 FY26 全年为 18%)。如果持续,这是 2H 混合毛利率的顺风。"
        )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(impact_para); set_font(r, size=10.5)

    # Section 4: Q2 2026 guidance
    h4_text = "Q2 2026 Guidance: RMB 850-900M (+20-27% YoY)" if lang == 'en' else "Q2 2026 指引:RMB 850-900M(+20-27% YoY)"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(h4_text); set_font(r, size=13, bold=True, color=NAVY)

    if lang == 'en':
        q2_para = (
            "Q1'26 actual RMB 681M + Q2'26 guide midpoint RMB 875M = 1H 2026 ~RMB 1,556M (~33% of our FY26E "
            "RMB 4,737M). Implies 2H 2026 must ramp to ~RMB 3,180M (~2× 1H) — heavily H2-loaded. This is "
            "plausible given Li Auto / Xiaomi / Mercedes multi-lidar L3+ SOPs all scheduled for H2 2026, but "
            "creates execution risk for the full-year number."
        )
    else:
        q2_para = (
            "Q1'26 实际 6.81 亿 + Q2'26 指引中位数 8.75 亿 = 1H 2026 约 15.56 亿元(约占我们 FY26E 47.37 亿的 33%)。"
            "意味着 2H 2026 必须放量至约 31.80 亿元(约 1H 的 2 倍)—— 严重 H2 加重。鉴于理想/小米/梅赛德斯多激光"
            "雷达 L3+ SOPs 均安排在 2H 2026,这是可信的,但给全年数字带来执行风险。"
        )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(q2_para); set_font(r, size=10.5)

    # Section 5: Updated catalyst tracker
    h5_text = "Updated Catalyst Tracker" if lang == 'en' else "更新的催化剂跟踪"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(h5_text); set_font(r, size=13, bold=True, color=NAVY)

    if lang == 'en':
        cat_hdr = ["#", "Catalyst", "Status (May 19, 2026)"]
        cat_rows = [
            ["1", "Q1 FY26 earnings — guided RMB 650-700M", "✅ TRIGGERED — RMB 680.6M, units +141%"],
            ["2", "Stock Connect inclusion of 2525.HK", "Pending Q4 2026"],
            ["3", "Li Auto / Xiaomi multi-lidar SOP", "Pending H2 2026"],
            ["4", "JT128 humanoid design-wins", "Ongoing"],
            ["5", "GM Super Cruise volume resumption", "Pending"],
            ["6", "NEW: Mercedes L3 design-in", "✅ TRIGGERED — official Q1'26 disclosure"],
            ["7", "NEW: Thailand manufacturing capacity", "✅ TRIGGERED — Galileo facility online"],
            ["8", "NEW: SGI / Picasso / Kosmo segment", "Materialising Q2 2026 onwards"],
        ]
    else:
        cat_hdr = ["#", "催化剂", "状态(2026 年 5 月 19 日)"]
        cat_rows = [
            ["1", "Q1 FY26 业绩 —— 指引 6.5-7 亿元", "✅ 已触发 —— 6.806 亿元,出货 +141%"],
            ["2", "港股 2525 进入港股通", "待定 Q4 2026"],
            ["3", "理想/小米多激光雷达 SOP 确认", "待定 2H 2026"],
            ["4", "JT128 人形机器人设计中标", "持续"],
            ["5", "通用超级巡航量恢复", "待定"],
            ["6", "新:梅赛德斯 L3 设计中标", "✅ 已触发 —— Q1'26 正式披露"],
            ["7", "新:泰国制造产能", "✅ 已触发 —— Galileo 基地上线"],
            ["8", "新:SGI / Picasso / Kosmo 分部", "从 Q2 2026 开始物化"],
        ]

    t = doc.add_table(rows=1 + len(cat_rows), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [0.4, 3.5, 3.0]
    for i, w in enumerate(widths):
        for cell in t.columns[i].cells:
            cell.width = Inches(w)
    for i, h in enumerate(cat_hdr):
        cell = t.rows[0].cells[i]; cell.text = ""
        para = cell.paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(h); set_font(r, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, "002060")
    for ri, row in enumerate(cat_rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]; cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = para.add_run(str(val)); set_font(r, size=9, bold=("TRIGGERED" in val or "已触发" in val))
            if "TRIGGERED" in val or "已触发" in val:
                r.font.color.rgb = GREEN
            if ri % 2 == 1: shade_cell(cell, "F2F2F2")
    for row in t.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for bn in ['top','left','bottom','right']:
                b = OxmlElement(f'w:{bn}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:color'),'888888')
                tcBorders.append(b)
            tcPr.append(tcBorders)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 6: new risks
    h6_text = "New Risks to Monitor" if lang == 'en' else "新增需监测风险"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(h6_text); set_font(r, size=13, bold=True, color=NAVY)

    if lang == 'en':
        new_risks = [
            ("Gross margin trajectory.",
             " Q1 GM at 39.1% is meaningfully below our model. If 2H 2026 GM stays sub-40%, our FY28E EPS could be 5-10% below model."),
            ("Heavily H2-loaded 2H ramp.",
             " Roughly half the year's revenue must come from H2 to meet our FY26E number. Execution risk on simultaneous Li Auto / Xiaomi / Mercedes SOPs."),
            ("SGI segment dilution.",
             " New Strategic Growth Initiatives segment may run at lower margins during incubation phase — watch FY26 SGI contribution vs the RMB 100M guide."),
        ]
    else:
        new_risks = [
            ("毛利率轨迹。",
             " Q1 毛利率 39.1% 明显低于我们模型。若 2H 2026 毛利率维持 40% 以下,我们 FY28E EPS 可能比模型低 5-10%。"),
            ("严重 H2 加重的放量。",
             " 大约半年收入需来自 2H 才能达到我们 FY26E 数字。同时执行理想/小米/梅赛德斯 SOPs 存在风险。"),
            ("SGI 分部稀释。",
             " 新「战略增长计划」分部可能在孵化阶段以较低利润率运行 —— 关注 FY26 SGI 实际贡献 vs 1 亿元指引。"),
        ]
    for lead, body in new_risks:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(lead); set_font(r, size=10.5, bold=True, color=RED)
        r = p.add_run(body); set_font(r, size=10.5)

    # Source line
    src_text = ("Source: Hesai Q1 2026 unaudited financial results press release, May 19, 2026 "
                "(https://investor.hesaitech.com/). Numbers as disclosed.") if lang == 'en' else (
                "资料来源:禾赛 Q1 2026 未经审计业绩公告,2026 年 5 月 19 日(https://investor.hesaitech.com/)。"
                "数据按披露引用。")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(src_text); set_font(r, size=8, italic=True, color=GRAY)

    doc.save(docx_path)
    return doc


ROOT = "/Users/x/projects/financial_agent/reports/company/Hesai_NASDAQ_HSAI"
for fname, lang in [("Hesai_NASDAQ_HSAI_Initiation_Report_2026-05-19.docx", "en"),
                    ("Hesai_NASDAQ_HSAI_Initiation_Report_2026-05-19_zh.docx", "zh")]:
    path = os.path.join(ROOT, fname)
    append_addendum(path, lang=lang)
    print(f"Appended Q1 2026 addendum to: {path}")
    # Verify
    d = Document(path)
    imgs = sum(1 for r in d.part.rels.values() if 'image' in r.target_ref)
    print(f"  Paragraphs: {len(d.paragraphs)}, Tables: {len(d.tables)}, Images: {imgs}")
